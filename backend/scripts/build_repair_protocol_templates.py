"""Build the v5 repair package without modifying approved transfer templates.

The first body table is retained from each published transfer_issue v3 source,
including its image relationships and formatting. Only the body below that
controlled company header is replaced.
"""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml.ns import qn
from docx.shared import Mm, Pt

ROOT = Path(__file__).resolve().parents[1]
TEMPLATES = ROOT / "resources" / "templates"

LABELS = {
    "bg": {
        "accept_title": "ПРОТОКОЛ ЗА ПРИЕМАНЕ НА ОБОРУДВАНЕ ЗА РЕМОНТ",
        "completed_title": "ПРОТОКОЛ ЗА ИЗВЪРШЕН РЕМОНТ",
        "number": "Протокол №", "date": "Дата", "equipment": "Оборудване",
        "ownership": "Собственост", "condition": "Състояние при приемане",
        "required_repair": "Описание на необходимия ремонт",
        "disassembly": "А) Демонтаж и подготовка",
        "findings": "Б) Констатирано след разглобяване, почистване и дефектация",
        "needed_parts": "В) Нужни части за ремонт",
        "diagnosis_time": "Реално време за демонтаж, почистване, диагностика и дефектация",
        "participants": "Извършили разглобяването / диагностиката",
        "handed": "Предал оборудването", "accepted": "Приел оборудването",
        "work": "Описание на извършените ремонтни дейности",
        "repair_time": "Реално време за сглобяване и ремонт",
        "testing_time": "Реално време за тестове",
        "parts": "Описание на вложените резервни части",
        "removed": "Демонтирани части",
        "hours": "Описание на реално отработеното време",
        "start": "Начало на ремонта", "end": "Край на ремонта",
        "total": "Общо реално време", "tests": "Тестове и резултат",
        "final_condition": "Състояние след ремонта", "final_result": "Краен резултат",
        "attachments": "Приложения към ремонтната карта",
        "performed": "Извършили ремонта", "approved": "Приел ремонта",
        "reference": "Ремонтна референция", "source": "Свързано връщане",
        "signature": "Подпис", "signature_status": "Статус на подписите", "prepared": "Съставил",
    },
    "en": {
        "accept_title": "EQUIPMENT ACCEPTANCE FOR REPAIR PROTOCOL",
        "completed_title": "COMPLETED REPAIR PROTOCOL",
        "number": "Protocol No.", "date": "Date", "equipment": "Equipment",
        "ownership": "Ownership", "condition": "Condition on acceptance",
        "required_repair": "Required repair",
        "disassembly": "A) Disassembly and preparation",
        "findings": "B) Findings after disassembly, cleaning and inspection",
        "needed_parts": "C) Parts required for repair",
        "diagnosis_time": "Actual disassembly, cleaning, diagnosis and inspection time",
        "participants": "Disassembly / diagnosis performed by",
        "handed": "Equipment handed over by", "accepted": "Equipment accepted by",
        "work": "Repair work performed", "repair_time": "Actual assembly and repair time",
        "testing_time": "Actual testing time", "parts": "Spare parts used",
        "removed": "Removed parts", "hours": "Actual labour time",
        "start": "Repair start", "end": "Repair end", "total": "Total actual time",
        "tests": "Tests and result", "final_condition": "Condition after repair",
        "final_result": "Final result", "attachments": "Repair case attachments",
        "performed": "Repair performed by", "approved": "Repair accepted by",
        "reference": "Repair reference", "source": "Linked return",
        "signature": "Signature", "signature_status": "Signature status", "prepared": "Prepared by",
    },
    "ru": {
        "accept_title": "ПРОТОКОЛ ПРИЕМА ОБОРУДОВАНИЯ В РЕМОНТ",
        "completed_title": "ПРОТОКОЛ ВЫПОЛНЕННОГО РЕМОНТА",
        "number": "Протокол №", "date": "Дата", "equipment": "Оборудование",
        "ownership": "Собственность", "condition": "Состояние при приеме",
        "required_repair": "Необходимый ремонт",
        "disassembly": "А) Демонтаж и подготовка",
        "findings": "Б) Результаты разборки, очистки и дефектации",
        "needed_parts": "В) Необходимые запасные части",
        "diagnosis_time": "Фактическое время демонтажа, очистки, диагностики и дефектации",
        "participants": "Демонтаж / диагностику выполнили",
        "handed": "Оборудование передал", "accepted": "Оборудование принял",
        "work": "Выполненные ремонтные работы", "repair_time": "Фактическое время сборки и ремонта",
        "testing_time": "Фактическое время испытаний", "parts": "Использованные запасные части",
        "removed": "Демонтированные части", "hours": "Фактическое рабочее время",
        "start": "Начало ремонта", "end": "Окончание ремонта", "total": "Общее фактическое время",
        "tests": "Испытания и результат", "final_condition": "Состояние после ремонта",
        "final_result": "Итоговый результат", "attachments": "Приложения к ремонту",
        "performed": "Ремонт выполнили", "approved": "Ремонт принял",
        "reference": "Референция ремонта", "source": "Связанный возврат",
        "signature": "Подпись", "signature_status": "Статус подписей", "prepared": "Составил",
    },
}


def _retain_approved_header(document: Document) -> None:
    body = document._element.body
    header_table_seen = False
    for element in list(body):
        if element.tag == qn("w:sectPr"):
            continue
        if element.tag == qn("w:tbl") and not header_table_seen:
            header_table_seen = True
            continue
        body.remove(element)
    if not header_table_seen:
        raise RuntimeError("The approved transfer template has no body header table.")


def _set_cell_text(cell, text: str, *, bold: bool = False) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.name = "Arial"
    run.font.size = Pt(8.5)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def _title(document: Document, text: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(5)
    paragraph.paragraph_format.space_after = Pt(5)
    run = paragraph.add_run(text)
    run.bold = True
    run.font.name = "Arial"
    run.font.size = Pt(12)


def _grid(document: Document, rows: list[tuple[str, str]], widths=(44, 142)) -> None:
    table = document.add_table(rows=0, cols=2)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for label, value in rows:
        cells = table.add_row().cells
        cells[0].width = Mm(widths[0])
        cells[1].width = Mm(widths[1])
        _set_cell_text(cells[0], label, bold=True)
        _set_cell_text(cells[1], value)


def _section(document: Document, label: str, placeholder: str, minimum_lines: int = 2) -> None:
    table = document.add_table(rows=2, cols=1)
    table.style = "Table Grid"
    _set_cell_text(table.cell(0, 0), label, bold=True)
    _set_cell_text(table.cell(1, 0), placeholder + ("\n" * max(0, minimum_lines - 1)))


def _dynamic_table(document: Document, label: str, code: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(4)
    paragraph.paragraph_format.space_after = Pt(2)
    run = paragraph.add_run(label)
    run.bold = True
    run.font.name = "Arial"
    run.font.size = Pt(9)
    document.add_paragraph(f"{{{{TABLE:{code}}}}}")


def _page_break(document: Document) -> None:
    document.add_paragraph().add_run().add_break(WD_BREAK.PAGE)


def build(language: str) -> Document:
    source = TEMPLATES / f"transfer_issue-{language}-v3.docx"
    document = Document(source)
    _retain_approved_header(document)
    section = document.sections[0]
    section.left_margin = section.right_margin = Mm(11)
    section.top_margin = Mm(8)
    section.bottom_margin = Mm(9)
    normal = document.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(8.5)
    t = LABELS[language]

    _title(document, t["accept_title"])
    document.add_paragraph(f"{{{{TEMPLATE_LANGUAGE:{language}}}}}")
    _grid(document, [
        (t["number"], "{{DOCUMENT_NUMBER}}"), (t["date"], "{{ACCEPTANCE_DATE}}"),
        (t["reference"], "{{REPAIR_REFERENCE}}"), (t["source"], "{{SOURCE_RETURN_REFERENCE}}"),
        (t["equipment"], "{{MACHINE_NAME}} · №{{MACHINE_NUMBER}} · {{BRAND}} {{MODEL}} · S/N {{SERIAL_NUMBER}}"),
        (t["ownership"], "{{OWNERSHIP}}"),
    ])
    _section(document, t["condition"], "{{CONDITION_BEFORE}}")
    _section(document, t["required_repair"], "{{REQUIRED_WORK}}")
    _section(document, t["disassembly"], "{{REMOVED_PARTS}}")
    _section(document, t["findings"], "{{DIAGNOSIS}}")
    _section(document, t["needed_parts"], "{{REQUIRED_PARTS}}")
    _grid(document, [(t["diagnosis_time"], "{{DIAGNOSIS_DURATION}}")])
    _dynamic_table(document, t["participants"], "REPAIR_PARTICIPANTS")
    _grid(document, [(t["handed"], "{{HANDED_OVER_NAME}}"), (t["accepted"], "{{ACCEPTED_BY_NAME}}")])

    _page_break(document)
    _title(document, t["completed_title"])
    _grid(document, [
        (t["number"], "{{DOCUMENT_NUMBER}}"), (t["date"], "{{COMPLETION_DATE}}"),
        (t["reference"], "{{REPAIR_REFERENCE}}"),
        (t["equipment"], "{{MACHINE_NAME}} · №{{MACHINE_NUMBER}} · {{BRAND}} {{MODEL}} · S/N {{SERIAL_NUMBER}}"),
        (t["ownership"], "{{OWNERSHIP}}"),
    ])
    _section(document, t["work"], "{{WORK_PERFORMED}}", 5)
    _section(document, t["removed"], "{{REMOVED_PARTS}}")
    _grid(document, [
        (t["repair_time"], "{{REPAIR_DURATION}}"),
        (t["testing_time"], "{{TESTING_DURATION}}"),
    ])
    _dynamic_table(document, t["performed"], "REPAIR_PARTICIPANTS")
    _dynamic_table(document, t["parts"], "PARTS_USED")

    _page_break(document)
    _title(document, t["hours"])
    _grid(document, [
        (t["start"], "{{REPAIR_START}}"), (t["end"], "{{REPAIR_END}}"),
        (t["total"], "{{TOTAL_WORK_DURATION}}"),
    ])
    _section(document, t["tests"], "{{TEST_RESULT}}", 4)
    _section(document, t["final_condition"], "{{CONDITION_AFTER}}")
    _section(document, t["final_result"], "{{FINAL_RESULT}}")
    _dynamic_table(document, t["attachments"], "REPAIR_ATTACHMENTS")
    _grid(document, [
        (t["performed"], "{{REPAIRER_NAMES}}"),
        (t["approved"], "{{APPROVED_BY_NAME}} · {{APPROVED_BY_JOB_TITLE}}"),
        (t["date"], "{{COMPLETION_DATE}}"),
        (t["signature"], "{{LEFT_SIGNATURE}}"),
        (t["signature_status"], "{{SIGNATURE_STATUS}}"),
        (t["prepared"], "{{PREPARER_NAME}} · {{PREPARER_JOB_TITLE}}"),
    ])
    return document


def main() -> None:
    for language in LABELS:
        output = TEMPLATES / f"repair_protocol-{language}-v5.docx"
        build(language).save(output)


if __name__ == "__main__":
    main()
