"""Verify the production image's non-root and LibreOffice runtime contract."""

from __future__ import annotations

import io
import os
import tempfile
from pathlib import Path

from app.template_engine import convert_docx_to_pdf
from docx import Document


def main() -> None:
    if hasattr(os, "geteuid") and os.geteuid() == 0:
        raise RuntimeError("The AssetCore runtime process must not run as root.")

    document = Document()
    document.add_heading("AssetCore container document smoke", level=1)
    document.add_paragraph("Проверка на DOCX към PDF в изолиран временен каталог.")
    output = io.BytesIO()
    document.save(output)
    pdf = convert_docx_to_pdf(output.getvalue())
    if not pdf or not pdf.startswith(b"%PDF"):
        raise RuntimeError("LibreOffice did not produce a valid PDF.")

    with tempfile.TemporaryDirectory(prefix="assetcore-runtime-smoke-") as name:
        target = Path(name) / "probe.bin"
        target.write_bytes(b"ok")
        if target.read_bytes() != b"ok":
            raise RuntimeError("The configured temporary directory is not writable.")
    print("container_runtime_status=healthy_non_root_document_generation")


if __name__ == "__main__":
    main()
