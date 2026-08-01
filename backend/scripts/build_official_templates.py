"""Build the controlled placeholder DOCX sources used by AssetCore.

The script is deterministic at the document-structure level. Resulting files are
versioned assets; historical business documents are never read or overwritten.
"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Mm, Pt

ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "resources" / "templates"
LOGO = ROOT / "resources" / "assets" / "odessos_logo.png"

LABELS = {
    "bg": {
        "issue": "ПРОТОКОЛ ПРЕДАВАНЕ НА МИЕЩА ТЕХНИКА",
        "return": "ПРОТОКОЛ ПРИЕМАНЕ НА МИЕЩА ТЕХНИКА СЛЕД ИЗПОЛЗВАНЕ",
        "repair": "ПРОТОКОЛ ПРЕДИ / СЛЕД РЕМОНТ",
        "parts": "ТЕХНИЧЕСКА СПЕЦИФИКАЦИЯ ЗА ДОСТАВКА НА РЕЗЕРВНИ ЧАСТИ",
        "date": "Дата", "number": "Документ №", "machine": "Машина",
        "brand": "Марка", "model": "Модел", "inventory": "Инвентарен №",
        "serial": "Сериен №", "pressure": "Налягане", "batch": "Партида",
        "condition": "Състояние", "usage": "Използване / местоназначение",
        "remarks": "Забележки", "component": "Елемент", "result": "Резултат",
        "preparer": "Съставил", "job": "Длъжност", "signature": "Подпис",
        "status": "Статус на подписите", "handover": "Предал", "acceptance": "Приел",
        "returned": "Върнал", "return_acceptance": "Приел връщането",
        "problem": "Регистриран проблем", "diagnosis": "Диагноза",
        "work": "Извършена работа", "test": "Резултат от теста",
        "before": "Състояние преди", "after": "Състояние след",
        "requester": "Заявител", "decision": "Решение",
        "repairer": "Извършил ремонта",
    },
    "en": {
        "issue": "HIGH-PRESSURE WASHING EQUIPMENT ISSUE PROTOCOL",
        "return": "HIGH-PRESSURE WASHING EQUIPMENT RETURN PROTOCOL",
        "repair": "BEFORE / AFTER REPAIR PROTOCOL",
        "parts": "TECHNICAL SPECIFICATION FOR SPARE-PARTS SUPPLY",
        "date": "Date", "number": "Document No.", "machine": "Machine",
        "brand": "Brand", "model": "Model", "inventory": "Inventory No.",
        "serial": "Serial No.", "pressure": "Pressure", "batch": "Batch",
        "condition": "Condition", "usage": "Use / destination", "remarks": "Remarks",
        "component": "Component", "result": "Result", "preparer": "Prepared by",
        "job": "Job title", "signature": "Signature", "status": "Signature status",
        "handover": "Handed over by", "acceptance": "Accepted by",
        "returned": "Returned by", "return_acceptance": "Return accepted by",
        "problem": "Reported problem", "diagnosis": "Diagnosis", "work": "Work performed",
        "test": "Test result", "before": "Condition before", "after": "Condition after",
        "requester": "Requester", "decision": "Decision",
        "repairer": "Repair performed by",
    },
    "ru": {
        "issue": "ПРОТОКОЛ ВЫДАЧИ МОЕЧНОЙ ТЕХНИКИ",
        "return": "ПРОТОКОЛ ПРИЁМА МОЕЧНОЙ ТЕХНИКИ ПОСЛЕ ИСПОЛЬЗОВАНИЯ",
        "repair": "ПРОТОКОЛ ДО / ПОСЛЕ РЕМОНТА",
        "parts": "ТЕХНИЧЕСКАЯ СПЕЦИФИКАЦИЯ НА ПОСТАВКУ ЗАПАСНЫХ ЧАСТЕЙ",
        "date": "Дата", "number": "Документ №", "machine": "Машина",
        "brand": "Марка", "model": "Модель", "inventory": "Инвентарный №",
        "serial": "Серийный №", "pressure": "Давление", "batch": "Партия",
        "condition": "Состояние", "usage": "Использование / назначение", "remarks": "Примечания",
        "component": "Элемент", "result": "Результат", "preparer": "Составил",
        "job": "Должность", "signature": "Подпись", "status": "Статус подписей",
        "handover": "Передал", "acceptance": "Принял", "returned": "Вернул",
        "return_acceptance": "Принял возврат", "problem": "Заявленная проблема",
        "diagnosis": "Диагноз", "work": "Выполненные работы", "test": "Результат теста",
        "before": "Состояние до", "after": "Состояние после", "requester": "Заявитель",
        "decision": "Решение",
        "repairer": "Выполнил ремонт",
    },
}

CHECKLIST_ITEMS = {
    "bg": ("Рама и корпус", "Двигател", "Помпа", "Маркучи", "Пистолет", "Дюзи", "Кабели", "Защити", "Принадлежности", "Общо състояние"),
    "en": ("Frame and body", "Engine", "Pump", "Hoses", "Gun", "Nozzles", "Cables", "Guards", "Accessories", "Overall condition"),
    "ru": ("Рама и корпус", "Двигатель", "Насос", "Шланги", "Пистолет", "Сопла", "Кабели", "Защита", "Принадлежности", "Общее состояние"),
}


def _base(title: str, language: str) -> Document:
    doc = Document()
    section = doc.sections[0]
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.left_margin = section.right_margin = Mm(11)
    section.top_margin = Mm(8)
    section.bottom_margin = Mm(9)
    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(9)
    header = doc.add_table(rows=1, cols=2)
    header.alignment = WD_TABLE_ALIGNMENT.CENTER
    header.columns[0].width = Mm(45)
    if LOGO.is_file():
        header.cell(0, 0).paragraphs[0].add_run().add_picture(str(LOGO), width=Mm(36))
    company = header.cell(0, 1).paragraphs[0]
    company.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = company.add_run("КРЗ ОДЕСОС АД\nODESSOS SHIPREPAIR YARD S.A.\nAssetCore")
    run.bold = True
    run.font.size = Pt(9)
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(title)
    run.bold = True
    run.font.size = Pt(13)
    # Machine-readable declaration, removed during rendering.
    doc.add_paragraph(f"{{{{TEMPLATE_LANGUAGE:{language}}}}}").style = doc.styles["Normal"]
    top = doc.add_table(rows=1, cols=2)
    top.alignment = WD_TABLE_ALIGNMENT.CENTER
    t = LABELS[language]
    top.cell(0, 0).text = f'{t["number"]}: {{{{DOCUMENT_NUMBER}}}}'
    top.cell(0, 1).text = f'{t["date"]}: {{{{CREATION_DATE}}}}'
    for cell in top.rows[0].cells:
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        cell.paragraphs[0].runs[0].bold = True
    return doc


def _machine_table(doc: Document, language: str) -> None:
    t = LABELS[language]
    rows = (
        (t["machine"], "{{MACHINE_NAME}}", t["inventory"], "{{MACHINE_NUMBER}}"),
        (t["brand"], "{{BRAND}}", t["model"], "{{MODEL}}"),
        (t["serial"], "{{SERIAL_NUMBER}}", t["pressure"], "{{PRESSURE_BAR}} bar"),
        (t["batch"], "{{BATCH_REFERENCE}}", "", ""),
    )
    table = doc.add_table(rows=0, cols=4)
    table.style = "Table Grid"
    for values in rows:
        cells = table.add_row().cells
        for index, value in enumerate(values):
            cells[index].text = value
            if index in (0, 2):
                cells[index].paragraphs[0].runs[0].bold = True


def _label_row(doc: Document, label: str, placeholder: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    table.cell(0, 0).text = f"{label}:\n{placeholder}"
    table.cell(0, 0).paragraphs[0].runs[0].bold = True


def _signature_table(doc: Document, language: str, left: str, right: str) -> None:
    t = LABELS[language]
    table = doc.add_table(rows=3, cols=2)
    table.style = "Table Grid"
    table.cell(0, 0).text = left
    table.cell(0, 1).text = right
    table.cell(1, 0).text = "{{LEFT_SIGNER_NAME}}\n{{LEFT_SIGNER_JOB_TITLE}}"
    table.cell(1, 1).text = "{{RIGHT_SIGNER_NAME}}\n{{RIGHT_SIGNER_JOB_TITLE}}"
    table.cell(2, 0).text = f'{t["signature"]}: {{{{LEFT_SIGNATURE}}}}'
    table.cell(2, 1).text = f'{t["signature"]}: {{{{RIGHT_SIGNATURE}}}}'
    for cell in table.rows[0].cells:
        cell.paragraphs[0].runs[0].bold = True
    _label_row(doc, t["status"], "{{SIGNATURE_STATUS}}")
    doc.add_paragraph(f'{t["preparer"]}: {{{{PREPARER_NAME}}}} | {t["job"]}: {{{{PREPARER_JOB_TITLE}}}}')


def _repair_signature_table(doc: Document, language: str) -> None:
    """Keep only the executor line from the controlled repair template."""
    t = LABELS[language]
    table = doc.add_table(rows=3, cols=1)
    table.style = "Table Grid"
    table.cell(0, 0).text = t["repairer"]
    table.cell(1, 0).text = "{{LEFT_SIGNER_NAME}}\n{{LEFT_SIGNER_JOB_TITLE}}"
    table.cell(2, 0).text = f'{t["signature"]}: {{{{LEFT_SIGNATURE}}}}'
    table.cell(0, 0).paragraphs[0].runs[0].bold = True
    _label_row(doc, t["status"], "{{SIGNATURE_STATUS}}")
    doc.add_paragraph(
        f'{t["preparer"]}: {{{{PREPARER_NAME}}}} | '
        f'{t["job"]}: {{{{PREPARER_JOB_TITLE}}}}'
    )


def build_transfer(language: str, mode: str) -> Document:
    t = LABELS[language]
    doc = _base(t[mode], language)
    _machine_table(doc, language)
    checklist = doc.add_table(rows=1, cols=3)
    checklist.style = "Table Grid"
    checklist.rows[0].cells[0].text = "№"
    checklist.rows[0].cells[1].text = t["component"]
    checklist.rows[0].cells[2].text = t["result"]
    for index, component in enumerate(CHECKLIST_ITEMS[language], 1):
        cells = checklist.add_row().cells
        cells[0].text = str(index)
        cells[1].text = component
        cells[2].text = f"{{{{CHECK_{index}}}}}"
    _label_row(doc, t["condition"], "{{CONDITION_TEXT}}")
    _label_row(doc, t["usage"], "{{USAGE_TEXT}}")
    _label_row(doc, t["remarks"], "{{REMARKS}}")
    if mode == "issue":
        _signature_table(doc, language, t["handover"], t["acceptance"])
    else:
        _signature_table(doc, language, t["returned"], t["return_acceptance"])
    return doc


def build_repair(language: str) -> Document:
    t = LABELS[language]
    doc = _base(t["repair"], language)
    _machine_table(doc, language)
    for label, placeholder in (
        (t["problem"], "{{REPORTED_PROBLEM}}"), (t["before"], "{{CONDITION_BEFORE}}"),
        (t["diagnosis"], "{{DIAGNOSIS}}"), (t["work"], "{{WORK_PERFORMED}}"),
        (t["test"], "{{TEST_RESULT}}"), (t["after"], "{{CONDITION_AFTER}}"),
    ):
        _label_row(doc, label, placeholder)
    doc.add_paragraph("{{TABLE:REPAIR_EVENTS}}")
    doc.add_paragraph("{{TABLE:PARTS_USED}}")
    _repair_signature_table(doc, language)
    return doc


def build_parts(language: str) -> Document:
    t = LABELS[language]
    doc = _base(t["parts"], language)
    _machine_table(doc, language)
    doc.add_paragraph("{{TABLE:REQUEST_LINES}}")
    _label_row(doc, t["remarks"], "{{REMARKS}}")
    _label_row(doc, t["decision"], "{{DECISION}}")
    _signature_table(doc, language, t["requester"], t["acceptance"])
    return doc


def main() -> None:
    OUTPUT.mkdir(parents=True, exist_ok=True)
    builders = {
        "transfer_issue": lambda language: build_transfer(language, "issue"),
        "transfer_return": lambda language: build_transfer(language, "return"),
        "repair_protocol": build_repair,
        "part_request": build_parts,
    }
    for code, builder in builders.items():
        version = 3 if code == "repair_protocol" else 2
        for language in LABELS:
            builder(language).save(OUTPUT / f"{code}-{language}-v{version}.docx")


if __name__ == "__main__":
    main()
