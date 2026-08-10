"""Generate isolated document samples and structural QA evidence.

The script uses a temporary in-memory SQLite database. It never writes sample
records to the configured AssetCore database and never changes source templates.
"""

from __future__ import annotations

import argparse
import hashlib
import io
import json
import re
import secrets
import sys
import zipfile
from pathlib import Path

from docx import Document
from pypdf import PdfReader
from sqlalchemy import create_engine, select
from sqlalchemy.orm import Session

BACKEND = Path(__file__).resolve().parents[1]
if str(BACKEND) not in sys.path:
    sys.path.insert(0, str(BACKEND))

from app.database import Base  # noqa: E402
from app.document_generation import (  # noqa: E402
    PARTS_REFERENCE,
    REPAIR_REFERENCE,
    make_part_request_documents,
    make_protocol_documents,
    make_repair_documents,
    make_return_documents,
)
from app.models import (  # noqa: E402
    DocumentTemplateVersion,
    Machine,
    PartCatalog,
    PartRequest,
    PartRequestLine,
    PartRequestPriority,
    PartRequestStatus,
    Repair,
    RepairEvent,
    RepairEventType,
    RepairParticipant,
    RepairStatus,
    TransferBatch,
    TransferBatchStatus,
    TransferProtocol,
    User,
    utcnow,
)
from app.seed import seed_database  # noqa: E402
from app.settings import settings  # noqa: E402
from app.template_engine import TOKEN_RE, validate_template  # noqa: E402


def _sha(path: Path) -> str:
    return hashlib.sha256(path.read_bytes()).hexdigest()


def _docx_audit(path: Path) -> dict:
    document = Document(path)
    visible_text = "\n".join(
        [paragraph.text for paragraph in document.paragraphs]
        + [paragraph.text for table in document.tables for row in table.rows for cell in row.cells for paragraph in cell.paragraphs]
    )
    with zipfile.ZipFile(path) as archive:
        names = set(archive.namelist())
        preserved = {}
        for name in ("word/header1.xml", "word/styles.xml", "word/numbering.xml", "word/settings.xml"):
            if name in names:
                preserved[name] = hashlib.sha256(archive.read(name)).hexdigest()
        media = {
            name: hashlib.sha256(archive.read(name)).hexdigest()
            for name in sorted(names)
            if name.startswith("word/media/")
        }
    section = document.sections[0]
    return {
        "sha256": _sha(path),
        "sections": len(document.sections),
        "tables": len(document.tables),
        "paragraphs": len(document.paragraphs),
        "page_width_twips": section.page_width.twips,
        "page_height_twips": section.page_height.twips,
        "margins_twips": {
            "left": section.left_margin.twips,
            "right": section.right_margin.twips,
            "top": section.top_margin.twips,
            "bottom": section.bottom_margin.twips,
        },
        "package_parts": len(names),
        "preserved_parts": preserved,
        "media": media,
        "has_cyrillic": bool(re.search(r"[А-Яа-я]", visible_text)),
        "unresolved_placeholders": sorted(set(TOKEN_RE.findall(visible_text))),
        "text": visible_text,
    }


def _pdf_audit(path: Path) -> dict:
    content = path.read_bytes()
    page_count = len(PdfReader(io.BytesIO(content)).pages)
    media_boxes = [
        [float(value) for value in match]
        for match in re.findall(
            rb"/MediaBox\s*\[\s*([0-9.]+)\s+([0-9.]+)\s+([0-9.]+)\s+([0-9.]+)\s*\]",
            content,
        )
    ]
    extracted_text = "\n".join(page.extract_text() or "" for page in PdfReader(path).pages)
    return {
        "sha256": _sha(path),
        "pages": page_count,
        "media_boxes": media_boxes,
        "has_cyrillic": bool(re.search(r"[А-Яа-я]", extracted_text)),
        "text_length": len(extracted_text),
        "text": extracted_text,
    }


def _write_pair(output: Path, stem: str, docx_bytes: bytes, pdf_bytes: bytes) -> dict:
    docx_path = output / f"{stem}.docx"
    pdf_path = output / f"{stem}.pdf"
    docx_path.write_bytes(docx_bytes)
    pdf_path.write_bytes(pdf_bytes)
    return {"docx": _docx_audit(docx_path), "pdf": _pdf_audit(pdf_path)}


def generate(output: Path) -> dict:
    output.mkdir(parents=True, exist_ok=True)
    engine = create_engine("sqlite+pysqlite:///:memory:")
    Base.metadata.create_all(engine)
    with Session(engine) as db:
        settings.owner_email = settings.owner_email or "qa-only@assetcore.invalid"
        settings.owner_job_title = settings.owner_job_title or "QA оператор"
        settings.owner_initial_password = settings.owner_initial_password or secrets.token_urlsafe(32)
        seed_database(db)
        machine = db.scalar(select(Machine).where(Machine.inventory_number == "4"))
        user = db.scalar(select(User).order_by(User.id))
        part = db.scalar(
            select(PartCatalog)
            .where(PartCatalog.is_verified.is_(True))
            .order_by(PartCatalog.id)
        )
        if machine is None or user is None or part is None:
            raise RuntimeError("Verified seed prerequisites are missing.")

        user.first_name = "Тест"
        user.middle_name = "Само"
        user.last_name = "Проверка"
        user.full_name = "Тест Само Проверка"
        user.job_title = "QA оператор"
        user.profile_status = "PROFILE_COMPLETE"

        now = utcnow()
        batch = TransferBatch(
            batch_reference="QA-ONLY-BATCH",
            status=TransferBatchStatus.ACTIVE.value,
            created_by_id=user.id,
        )
        db.add(batch)
        db.flush()
        transfer = TransferProtocol(
            machine_id=machine.id,
            batch_id=batch.id,
            protocol_type="Предаване",
            protocol_number="QA-ONLY-PROTOCOL",
            is_active=True,
            issued_by_id=user.id,
            issued_at=now,
            previous_status=machine.status,
            previous_location_id=machine.location_id,
            issue_location_id=machine.location_id,
            handed_over_by=user.full_name,
            handed_over_job_title=user.job_title,
            accepted_by="Иван Петров Иванов",
            condition_text="Добро",
            location_text="Сух док - QA проверка",
            remarks="Комплектна и подготвена за работа.",
            issue_checklist=[
                {"code": "PUMP", "label": "Помпа", "condition": "GOOD"},
                {"code": "SUPPLY_HOSE", "label": "Шланг захранващ", "condition": "GOOD", "length_m": 25},
                {"code": "HP_HOSE", "label": "Шланг изходящ ВН", "condition": "GOOD", "length_m": 40},
                {"code": "GUN", "label": "Пистолет", "condition": "GOOD"},
                {"code": "NOZZLE", "label": "Дюза метла / ротационна", "condition": "NA"},
                {"code": "TIPS", "label": "Накрайници", "condition": "GOOD"},
                {"code": "CABLE", "label": "Кабел", "condition": "GOOD", "length_m": 30},
                {"code": "PLUG", "label": "Куплунг / Еврощек", "condition": "GOOD"},
                {"code": "CHASSIS", "label": "Ходова част", "condition": "GOOD"},
                {"code": "BODY", "label": "Корпус", "condition": "GOOD"},
            ],
        )
        db.add(transfer)
        db.flush()

        repair = Repair(
            machine_id=machine.id,
            repair_reference="QA-ONLY-REPAIR",
            reported_problem="QA тестов запис за проверка на оформлението",
            diagnosis="QA тестова диагностика",
            required_work="QA тестова необходима работа",
            required_parts_text="QA тестова бележка за нужни части",
            removed_parts_text="QA тестов демонтаж и подготовка",
            diagnostic_cleaning="QA тестово почистване при диагностиката",
            work_performed="QA тестово описание на извършена работа",
            result="QA тестов резултат",
            condition_before="QA тестово състояние преди ремонта",
            condition_after="QA тестово състояние след ремонта",
            inspection_completed_at=now,
            test_passed=True,
            test_method="QA функционален тест под налягане",
            test_pressure_bar=500,
            leaks_detected=False,
            test_details="QA тестът е отчетен като успешен само в изолираната проверка",
            diagnosis_minutes=30,
            repair_minutes=75,
            testing_minutes=20,
            status=RepairStatus.COMPLETED.value,
            responsible_user_id=user.id,
            accepted_by_id=user.id,
            approved_by_id=user.id,
            approved_at=now,
            started_at=now,
            closed_at=now,
        )
        db.add(repair)
        db.flush()
        db.add(
            RepairParticipant(
                repair_id=repair.id,
                full_name_snapshot="QA участник без производствен запис",
                job_title_snapshot="QA роля",
                contribution="Визуална проверка на participant реда",
                minutes_worked=55,
                identity_key="qa-only-participant",
                created_by_id=user.id,
            )
        )
        db.add(
            RepairEvent(
                repair_id=repair.id,
                event_type=RepairEventType.TEST.value,
                status_before=RepairStatus.REPAIRING.value,
                status_after=RepairStatus.COMPLETED.value,
                description="QA тестово хронологично събитие",
                user_id=user.id,
            )
        )

        request = PartRequest(
            machine_id=machine.id,
            request_reference="QA-ONLY-PART-REQUEST",
            part_name=part.description,
            part_number=part.part_number,
            quantity=1,
            reason="QA проверка на оформлението",
            priority=PartRequestPriority.NORMAL.value,
            status=PartRequestStatus.APPROVED.value,
            language="bg",
            requested_by_id=user.id,
            submitted_at=now,
            decided_at=now,
            decided_by_id=user.id,
            decision_note="QA тестово одобрение",
        )
        db.add(request)
        db.flush()
        db.add(
            PartRequestLine(
                request_id=request.id,
                catalog_part_id=part.id,
                position=part.position,
                part_number=part.part_number,
                description=part.description,
                quantity=1,
                unit=part.unit,
                source_document=part.source_document,
                source_page=part.source_page,
            )
        )
        db.flush()
        db.refresh(transfer)
        db.refresh(repair)
        db.refresh(request)

        issue_documents = make_protocol_documents(db, transfer, batch, user.id, "bg")
        transfer.is_active = False
        transfer.returned_at = now
        transfer.returned_by_name = "Иван Петров Иванов"
        transfer.return_accepted_by = user.full_name
        transfer.return_accepted_job_title = user.job_title
        transfer.return_condition_text = "Добро"
        transfer.return_result_text = "Приета след проверка"
        transfer.return_notes = "QA бележка"
        transfer.return_checklist = list(transfer.issue_checklist or [])
        return_documents = make_return_documents(db, transfer, batch, user.id, "bg")
        repair_documents = make_repair_documents(db, repair, user.id, "bg")
        request_documents = make_part_request_documents(db, request, user.id, "bg")

        def pair(records):
            by_format = {record.format: record.content for record in records}
            return by_format["docx"], by_format["pdf"]

        issue_docx, issue_pdf = pair(issue_documents)
        return_docx, return_pdf = pair(return_documents)
        repair_docx, repair_pdf = pair(repair_documents)
        request_docx, request_pdf = pair(request_documents)

        results = {
            "issue": _write_pair(
                output,
                "issue-protocol-bg",
                issue_docx,
                issue_pdf,
            ),
            "return": _write_pair(
                output,
                "return-protocol-bg",
                return_docx,
                return_pdf,
            ),
            "repair": _write_pair(
                output,
                "repair-protocol-bg",
                repair_docx,
                repair_pdf,
            ),
            "part_request": _write_pair(
                output,
                "part-request-bg",
                request_docx,
                request_pdf,
            ),
        }
        results["template_validation"] = {
            str(version.id): validate_template(version)
            for version in db.scalars(
                select(DocumentTemplateVersion).where(DocumentTemplateVersion.is_published.is_(True))
            )
        }

    source_hashes_after = {
        "repair_reference": _sha(REPAIR_REFERENCE),
        "parts_reference": _sha(PARTS_REFERENCE),
    }
    source_structures = {
        "repair_reference": _docx_audit(REPAIR_REFERENCE),
        "parts_reference": _docx_audit(PARTS_REFERENCE),
    }
    results["source_hashes_after"] = source_hashes_after
    results["source_structures"] = source_structures
    results["source_preservation"] = {
        "repair_unchanged": source_hashes_after["repair_reference"]
        == "39337dfc445d61b4d5144259ca35624c2049d266378e326780176de3104784c1",
        "parts_unchanged": source_hashes_after["parts_reference"]
        == "3ba8e43102ae044b02b6aa7a4cd3b06ff00bce444a56fc8da6ba737c42bbc7a7",
    }
    results["release_checks"] = {
        "all_templates_valid": all(item["valid"] for item in results["template_validation"].values()),
        "no_unresolved_placeholders": all(not results[name]["docx"]["unresolved_placeholders"] for name in ("issue", "return", "repair", "part_request")),
        "cyrillic_in_docx": all(results[name]["docx"]["has_cyrillic"] for name in ("issue", "return", "repair", "part_request")),
        "cyrillic_in_pdf": all(results[name]["pdf"]["has_cyrillic"] for name in ("issue", "return", "repair", "part_request")),
        "source_hashes_unchanged": all(results["source_preservation"].values()),
        "document_numbers_present": all("QA-ONLY" in results[name]["docx"]["text"] for name in ("issue", "return", "repair", "part_request")),
        "signature_status_present": all(
            "НЕПЪЛНО ПОДПИСАН" in results[name]["docx"]["text"]
            for name in ("issue", "return", "part_request")
        )
        and "ОКОНЧАТЕЛЕН ВЪТРЕШЕН ПРОТОКОЛ"
        in results["repair"]["docx"]["text"],
        "transfer_protocols_single_page": all(
            results[name]["pdf"]["pages"] == 1 for name in ("issue", "return")
        ),
        "transfer_protocol_header_media_present": all(
            len(results[name]["docx"]["media"]) >= 3 for name in ("issue", "return")
        ),
        "transfer_protocol_original_rows_present": all(
            all(label in results[name]["docx"]["text"] for label in (
                "Помпа", "Шланг захранващ", "Шланг изходящ ВН",
                "Пистолет", "Дюза метла / ротационна", "Накрайници",
                "Кабел", "Куплунг / Еврощек", "Ходова част", "Корпус",
            ))
            for name in ("issue", "return")
        ),
        "repair_content_sections_present": all(
            all(
                label in results["repair"][format_name]["text"]
                for label in (
                    "ПРОТОКОЛ ЗА ПРИЕМАНЕ НА ОБОРУДВАНЕ ЗА РЕМОНТ",
                    "ПРОТОКОЛ ЗА ИЗВЪРШЕН РЕМОНТ",
                    "Реално време",
                )
            )
            for format_name in ("docx", "pdf")
        ),
        "repair_pdf_has_three_pages": results["repair"]["pdf"]["pages"] == 3,
        "repair_header_media_present": len(results["repair"]["docx"]["media"]) >= 3,
    }
    (output / "qa-manifest.json").write_text(
        json.dumps(results, ensure_ascii=False, indent=2), encoding="utf-8"
    )
    engine.dispose()
    return results


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("output", type=Path)
    args = parser.parse_args()
    result = generate(args.output.resolve())
    print(json.dumps(result, ensure_ascii=False, indent=2))
    if not all(result["release_checks"].values()):
        raise SystemExit(1)


if __name__ == "__main__":
    main()
