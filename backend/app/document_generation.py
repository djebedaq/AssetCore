from __future__ import annotations

import hashlib
import io
import json
import re
import zipfile
from datetime import UTC, datetime
from pathlib import Path
from xml.sax.saxutils import escape

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Mm, Pt
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import mm
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.platypus import (
    Image as ReportLabImage,
)
from reportlab.platypus import (
    KeepTogether,
    PageBreak,
    Paragraph,
    SimpleDocTemplate,
    Spacer,
    Table,
    TableStyle,
)
from sqlalchemy import or_, select
from sqlalchemy.orm import Session

from .localization import translate
from .models import (
    DocumentTemplate,
    DocumentTemplateVersion,
    DocumentType,
    GeneratedDocument,
    LanguageCode,
    OfficialDocument,
    OfficialDocumentStatus,
    OfficialDocumentVersion,
    PartRequest,
    PartRequestLine,
    ProtocolDocument,
    Repair,
    TransferBatch,
    TransferProtocol,
    User,
    utcnow,
)
from .official_documents.integrity import (
    move_current_version,
    require_current_version,
    set_current_version,
)
from .template_engine import TemplateValidationError, convert_docx_to_pdf, render_docx

DOCX_MEDIA_TYPE = (
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
)
PDF_MEDIA_TYPE = "application/pdf"


class ConfirmedTemplateUnavailableError(RuntimeError):
    def __init__(self, document_type: str, language: str):
        self.document_type = document_type
        self.language = _language(language)
        self.message = translate("document.template_unavailable", self.language)
        super().__init__(self.message)

RESOURCES = Path(__file__).resolve().parents[1] / "resources"
ASSETS = RESOURCES / "assets"


def _reference_by_sha256(folder: Path, expected_sha256: str) -> Path:
    """Resolve a controlled reference by content, never by a locale-sensitive name."""
    for candidate in sorted(folder.glob("*.docx")):
        if hashlib.sha256(candidate.read_bytes()).hexdigest() == expected_sha256:
            return candidate
    raise FileNotFoundError(
        f"Controlled DOCX reference {expected_sha256} is missing from {folder.name}."
    )


REPAIR_REFERENCE = _reference_by_sha256(
    RESOURCES / "reference_protocols",
    "39337dfc445d61b4d5144259ca35624c2049d266378e326780176de3104784c1",
)
PARTS_REFERENCE = _reference_by_sha256(
    RESOURCES / "reference_protocols",
    "3ba8e43102ae044b02b6aa7a4cd3b06ff00bce444a56fc8da6ba737c42bbc7a7",
)

TEXT = {
    "bg": {
        "protocol": "ПРОТОКОЛ",
        "issue_title": "ПРОТОКОЛ ПРЕДАВАНЕ НА МИЕЩА ТЕХНИКА",
        "return_title": "ПРОТОКОЛ ПРИЕМАНЕ НА МИЕЩА ТЕХНИКА СЛЕД ИЗПОЛЗВАНЕ",
        "date": "ДАТА",
        "equipment": "ОБОРУДВАНЕ",
        "model": "МОДЕЛ",
        "inventory": "ЗАВОДСКИ НОМЕР",
        "serial": "СЕРИЕН НОМЕР",
        "number": "№",
        "element": "Елемент",
        "condition": "Състояние",
        "overall_condition": "Общо състояние",
        "usage_issue": "Оборудването ще се използва за",
        "usage_return": "Оборудването е било използвано за",
        "remarks": "Забележки",
        "handed": "Предал оборудването от страна на ДИРП",
        "accepted": "Приел оборудването",
        "returned": "Върнал оборудването",
        "name": "Три имена / фирма / цех",
        "signature": "Подпис",
        "batch": "Партида",
        "repair_protocol": "ПРОТОКОЛ ПРЕДИ / СЛЕД РЕМОНТ",
        "machine": "Машина",
        "reported_problem": "Регистриран проблем",
        "symptoms": "Наблюдавани симптоми",
        "required_work": "Необходима работа",
        "removed_parts": "Демонтирани части",
        "cleaning": "Почистване при диагностиката",
        "duration": "Време",
        "condition_before": "Състояние преди ремонта",
        "diagnosis": "Преглед и диагноза",
        "work": "Извършени ремонтни дейности",
        "events": "Хронология на ремонта",
        "parts_used": "Използвани части",
        "test": "Почистване и тестване",
        "test_method": "Метод",
        "test_pressure": "Достигнато налягане",
        "leaks": "Установени течове",
        "electrical_test": "Електрически тест",
        "functional_test": "Функционален тест",
        "yes": "Да",
        "no": "Не",
        "condition_after": "Състояние и резултат след ремонта",
        "responsible": "Извършил ремонта",
        "approver": "Одобрил",
        "part_request_title": "ТЕХНИЧЕСКА СПЕЦИФИКАЦИЯ\nЗА ДОСТАВКА НА РЕЗЕРВНИ ЧАСТИ",
        "position": "Поз.",
        "part_number": "PART №",
        "description": "ОПИСАНИЕ",
        "quantity": "КОЛИЧЕСТВО",
        "source": "Източник",
        "request_number": "ЗАЯВКА №",
        "requester": "Заявител",
        "decision": "Решение",
    },
    "en": {
        "protocol": "PROTOCOL",
        "issue_title": "HIGH-PRESSURE WASHING EQUIPMENT ISSUE PROTOCOL",
        "return_title": "HIGH-PRESSURE WASHING EQUIPMENT RETURN PROTOCOL",
        "date": "DATE",
        "equipment": "EQUIPMENT",
        "model": "MODEL",
        "inventory": "FACTORY / INVENTORY NUMBER",
        "serial": "SERIAL NUMBER",
        "number": "No.",
        "element": "Component",
        "condition": "Condition",
        "overall_condition": "Overall condition",
        "usage_issue": "The equipment will be used for",
        "usage_return": "The equipment was used for",
        "remarks": "Remarks",
        "handed": "Handed over by DIRP",
        "accepted": "Accepted by",
        "returned": "Returned by",
        "name": "Full name / company / department",
        "signature": "Signature",
        "batch": "Batch",
        "repair_protocol": "BEFORE / AFTER REPAIR PROTOCOL",
        "machine": "Machine",
        "reported_problem": "Reported problem",
        "symptoms": "Observed symptoms",
        "required_work": "Required work",
        "removed_parts": "Removed parts",
        "cleaning": "Cleaning during diagnosis",
        "duration": "Time",
        "condition_before": "Condition before repair",
        "diagnosis": "Inspection and diagnosis",
        "work": "Repair actions performed",
        "events": "Repair timeline",
        "parts_used": "Parts used",
        "test": "Cleaning and testing",
        "test_method": "Method",
        "test_pressure": "Pressure reached",
        "leaks": "Leaks detected",
        "electrical_test": "Electrical test",
        "functional_test": "Functional test",
        "yes": "Yes",
        "no": "No",
        "condition_after": "Condition and result after repair",
        "responsible": "Repair performed by",
        "approver": "Approved by",
        "part_request_title": "TECHNICAL SPECIFICATION\nFOR SPARE PARTS SUPPLY",
        "position": "Pos.",
        "part_number": "PART No.",
        "description": "DESCRIPTION",
        "quantity": "QUANTITY",
        "source": "Source",
        "request_number": "REQUEST No.",
        "requester": "Requested by",
        "decision": "Decision",
    },
    "ru": {
        "protocol": "ПРОТОКОЛ",
        "issue_title": "ПРОТОКОЛ ВЫДАЧИ МОЕЧНОЙ ТЕХНИКИ",
        "return_title": "ПРОТОКОЛ ПРИЕМА МОЕЧНОЙ ТЕХНИКИ ПОСЛЕ ИСПОЛЬЗОВАНИЯ",
        "date": "ДАТА",
        "equipment": "ОБОРУДОВАНИЕ",
        "model": "МОДЕЛЬ",
        "inventory": "ЗАВОДСКОЙ / ИНВЕНТАРНЫЙ НОМЕР",
        "serial": "СЕРИЙНЫЙ НОМЕР",
        "number": "№",
        "element": "Элемент",
        "condition": "Состояние",
        "overall_condition": "Общее состояние",
        "usage_issue": "Оборудование будет использовано для",
        "usage_return": "Оборудование использовалось для",
        "remarks": "Примечания",
        "handed": "Передал оборудование со стороны ДИРП",
        "accepted": "Принял оборудование",
        "returned": "Вернул оборудование",
        "name": "ФИО / компания / подразделение",
        "signature": "Подпись",
        "batch": "Партия",
        "repair_protocol": "ПРОТОКОЛ ДО / ПОСЛЕ РЕМОНТА",
        "machine": "Машина",
        "reported_problem": "Заявленная неисправность",
        "symptoms": "Наблюдаемые симптомы",
        "required_work": "Необходимые работы",
        "removed_parts": "Демонтированные детали",
        "cleaning": "Очистка при диагностике",
        "duration": "Время",
        "condition_before": "Состояние до ремонта",
        "diagnosis": "Осмотр и диагностика",
        "work": "Выполненные ремонтные работы",
        "events": "Хронология ремонта",
        "parts_used": "Использованные детали",
        "test": "Очистка и испытание",
        "test_method": "Метод",
        "test_pressure": "Достигнутое давление",
        "leaks": "Обнаружены утечки",
        "electrical_test": "Электрическое испытание",
        "functional_test": "Функциональное испытание",
        "yes": "Да",
        "no": "Нет",
        "condition_after": "Состояние и результат после ремонта",
        "responsible": "Ремонт выполнил",
        "approver": "Утвердил",
        "part_request_title": "ТЕХНИЧЕСКАЯ СПЕЦИФИКАЦИЯ\nНА ПОСТАВКУ ЗАПАСНЫХ ЧАСТЕЙ",
        "position": "Поз.",
        "part_number": "PART №",
        "description": "ОПИСАНИЕ",
        "quantity": "КОЛИЧЕСТВО",
        "source": "Источник",
        "request_number": "ЗАЯВКА №",
        "requester": "Заявитель",
        "decision": "Решение",
    },
}

CHECKLIST_CODES = [
    "pump",
    "supply_hose",
    "hp_hose",
    "gun",
    "nozzle",
    "tips",
    "cable",
    "plug",
    "chassis",
    "body",
]

CHECKLIST = {
    "bg": [
        "Помпа",
        "Шланг захранващ",
        "Шланг изходящ ВН",
        "Пистолет",
        "Дюза метла / ротационна",
        "Накрайници",
        "Кабел - метри",
        "Куплунг / Еврокак",
        "Ходова част",
        "Корпус",
    ],
    "en": [
        "Pump",
        "Supply hose",
        "High-pressure outlet hose",
        "Gun",
        "Broom / rotary nozzle",
        "Nozzles",
        "Cable - metres",
        "Coupling / Euro coupling",
        "Running gear",
        "Body",
    ],
    "ru": [
        "Насос",
        "Подающий шланг",
        "Выходной шланг ВД",
        "Пистолет",
        "Щелевая / ротационная форсунка",
        "Наконечники",
        "Кабель - метры",
        "Муфта / Евросоединение",
        "Ходовая часть",
        "Корпус",
    ],
}


def safe_filename(value: str) -> str:
    """Return a stable ASCII filename stem without path components."""
    cleaned = re.sub(r"[^A-Za-z0-9._-]+", "-", value.strip())
    return cleaned.strip(".-") or "assetcore-document"


def _language(value: str | None) -> str:
    return value if value in TEXT else LanguageCode.BG.value


def _set_run_font(run, size: float = 9, bold: bool = False) -> None:
    run.font.name = "Arial"
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Arial")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    run.font.size = Pt(size)
    run.bold = bold


def _keep_with_next(paragraph) -> None:
    paragraph.paragraph_format.keep_with_next = True


def _set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    repeat = OxmlElement("w:tblHeader")
    repeat.set(qn("w:val"), "true")
    tr_pr.append(repeat)


def _set_cell(cell, value: str, *, bold: bool = False, size: float = 9) -> None:
    cell.text = ""
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run(value)
    _set_run_font(run, size, bold)


def _clear_body(document: Document) -> None:
    body = document._element.body
    for element in list(body):
        if element.tag != qn("w:sectPr"):
            body.remove(element)


def _prepare_document(reference: Path, *, a4: bool = True) -> Document:
    document = Document(reference) if reference.is_file() else Document()
    _clear_body(document)
    section = document.sections[0]
    if a4:
        section.page_width = Mm(210)
        section.page_height = Mm(297)
    section.top_margin = Mm(8)
    section.bottom_margin = Mm(9)
    section.left_margin = Mm(10)
    section.right_margin = Mm(8)
    section.header_distance = Mm(2)
    section.footer_distance = Mm(5)
    normal = document.styles["Normal"]
    normal.font.name = "Arial"
    normal._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Arial")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    normal.font.size = Pt(9)
    return document


def _add_centered(document: Document, text: str, size: float, bold: bool = False):
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(2)
    run = paragraph.add_run(text)
    _set_run_font(run, size, bold)
    return paragraph


def _machine_model(transfer: TransferProtocol) -> str:
    machine = transfer.machine
    brand = (machine.brand or machine.manufacturer or "").strip()
    pressure = f"{machine.pressure_bar}bar" if machine.pressure_bar else ""
    family = " - ".join(value for value in (brand, pressure) if value)
    return "; ".join(value for value in (family, machine.model or "") if value)


def _usage_text(transfer: TransferProtocol) -> str:
    return "; ".join(
        value
        for value in (
            transfer.company_unit,
            transfer.department,
            transfer.vessel,
            transfer.dock,
            transfer.pier,
            transfer.work_area,
            transfer.location_text,
        )
        if value
    )


def _equipment_text(transfer: TransferProtocol, language: str) -> str:
    labels = {
        "bg": ("Оборудване", "Шлангове", "Дюзи", "Пистолети", "Принадлежности"),
        "en": ("Equipment", "Hoses", "Nozzles", "Guns", "Accessories"),
        "ru": ("Оборудование", "Шланги", "Сопла", "Пистолеты", "Принадлежности"),
    }[language]
    values = (
        transfer.equipment,
        transfer.hoses,
        transfer.nozzles,
        transfer.guns,
        transfer.accessories,
    )
    return "; ".join(
        f"{label}: {value}" for label, value in zip(labels, values, strict=True) if value
    )


def _return_findings(transfer: TransferProtocol, language: str) -> str:
    labels = {
        "bg": {
            "missing": "Липсващо оборудване",
            "damage": "Повреди",
            "contamination": "Замърсяване",
            "cleaning": "Необходимо почистване",
            "inspection": "Необходим преглед",
            "repair": "Необходим ремонт",
            "yes": "да",
            "no": "не",
        },
        "en": {
            "missing": "Missing equipment", "damage": "Damage",
            "contamination": "Contamination", "cleaning": "Cleaning required",
            "inspection": "Inspection required", "repair": "Repair required",
            "yes": "yes", "no": "no",
        },
        "ru": {
            "missing": "Недостающее оборудование", "damage": "Повреждения",
            "contamination": "Загрязнение", "cleaning": "Требуется очистка",
            "inspection": "Требуется осмотр", "repair": "Требуется ремонт",
            "yes": "да", "no": "нет",
        },
    }[language]
    values = [
        (labels["missing"], transfer.return_missing_equipment),
        (labels["damage"], transfer.return_damage),
        (labels["contamination"], transfer.return_contamination),
        (labels["cleaning"], labels["yes"] if transfer.return_cleaning_required else labels["no"]),
        (labels["inspection"], labels["yes"] if transfer.return_inspection_required else labels["no"]),
        (labels["repair"], labels["yes"] if transfer.return_repair_required else labels["no"]),
    ]
    return "; ".join(f"{label}: {value}" for label, value in values if value)





def _protocol_remarks(transfer: TransferProtocol, operation: str, language: str) -> str:
    if operation == "return":
        return "; ".join(
            value
            for value in (
                transfer.return_result_text,
                _return_findings(transfer, language),
                transfer.return_notes,
            )
            if value
        )
    return "; ".join(
        value
        for value in (_equipment_text(transfer, language), transfer.remarks)
        if value
    )

def _checklist_rows(transfer: TransferProtocol, operation: str, language: str) -> list[tuple[str, str]]:
    stored = transfer.return_checklist if operation == "return" else transfer.issue_checklist
    condition_labels = {
        "bg": {"GOOD": "Добро", "SATISFACTORY": "Задоволително", "REPAIR": "За ремонт", "FAULTY": "Неизправно", "MISSING": "Липсва", "NA": "Не се прилага"},
        "en": {"GOOD": "Good", "SATISFACTORY": "Satisfactory", "REPAIR": "For repair", "FAULTY": "Faulty", "MISSING": "Missing", "NA": "N/A"},
        "ru": {"GOOD": "Хорошее", "SATISFACTORY": "Удовлетворительное", "REPAIR": "В ремонт", "FAULTY": "Неисправно", "MISSING": "Отсутствует", "NA": "Не применяется"},
    }[language]
    if not stored:
        return [(component, "") for component in CHECKLIST[language]]
    label_by_code = dict(zip(CHECKLIST_CODES, CHECKLIST[language], strict=True))
    rows = []
    for item in stored:
        code = str(item.get("code") or "")
        label = label_by_code.get(code) or str(item.get("label") or code)
        value = condition_labels.get(str(item.get("condition") or ""), str(item.get("condition") or ""))
        if item.get("length_m") is not None:
            value += f" · {item['length_m']} m"
        if item.get("note"):
            value += f" · {item['note']}"
        rows.append((label, value))
    return rows


def _identity_rows(transfer: TransferProtocol, operation: str, language: str) -> list[tuple[str, str]]:
    t = TEXT[language]
    date_value = (
        (transfer.returned_at or transfer.return_requested_at)
        if operation == "return"
        else transfer.issued_at
    ) or transfer.created_at
    return [
        (t["date"], date_value.strftime("%d.%m.%Y")),
        (t["equipment"], transfer.machine.category or "HPWJ"),
        (t["model"], _machine_model(transfer)),
        (t["inventory"], f"№ {transfer.machine.inventory_number}"),
        (t["serial"], transfer.machine.serial_number or ""),
    ]


def _protocol_snapshot(
    transfer: TransferProtocol, batch_reference: str, operation: str, language: str
) -> dict:
    return {
        "operation": operation,
        "language": language,
        "protocol_number": transfer.protocol_number,
        "batch_reference": batch_reference,
        "machine_id": transfer.machine_id,
        "machine_number": transfer.machine.inventory_number,
        "machine_name": transfer.machine.name,
        "brand": transfer.machine.brand,
        "model": transfer.machine.model,
        "pressure_bar": transfer.machine.pressure_bar,
        "serial_number": transfer.machine.serial_number,
        "company_unit": transfer.company_unit,
        "department": transfer.department,
        "vessel": transfer.vessel,
        "dock": transfer.dock,
        "pier": transfer.pier,
        "work_area": transfer.work_area,
        "location_text": transfer.location_text,
        "equipment": transfer.equipment,
        "hoses": transfer.hoses,
        "nozzles": transfer.nozzles,
        "guns": transfer.guns,
        "accessories": transfer.accessories,
        "condition": (
            transfer.return_condition_text
            if operation == "return"
            else transfer.condition_text
        ),
        "result": transfer.return_result_text if operation == "return" else None,
        "missing_equipment": transfer.return_missing_equipment if operation == "return" else None,
        "damage": transfer.return_damage if operation == "return" else None,
        "contamination": transfer.return_contamination if operation == "return" else None,
        "cleaning_required": transfer.return_cleaning_required if operation == "return" else None,
        "inspection_required": transfer.return_inspection_required if operation == "return" else None,
        "repair_required": transfer.return_repair_required if operation == "return" else None,
        "remarks": transfer.return_notes if operation == "return" else transfer.remarks,
        "handed_over_by": (
            transfer.returned_by_name if operation == "return" else transfer.handed_over_by
        ),
        "handed_over_job_title": (
            transfer.returned_by_job_title
            if operation == "return"
            else transfer.handed_over_job_title
        ),
        "handed_over_organization": (
            transfer.returned_by_company
            if operation == "return"
            else transfer.handed_over_department
        ),
        "accepted_by": (
            transfer.return_accepted_by if operation == "return" else transfer.accepted_by
        ),
        "accepted_by_job_title": (
            transfer.return_accepted_job_title
            if operation == "return"
            else transfer.accepted_by_job_title
        ),
        "accepted_by_organization": (
            transfer.return_accepted_department
            if operation == "return"
            else transfer.accepted_by_company
        ),
        "document_date": (
            (
                (transfer.returned_at or transfer.return_requested_at)
                if operation == "return"
                else transfer.issued_at
            )
            or transfer.created_at
        ).isoformat(),
    }


def _build_protocol_docx(
    transfer: TransferProtocol,
    batch_reference: str,
    *,
    operation: str,
    language: str,
) -> bytes:
    language = _language(language)
    t = TEXT[language]
    document = _prepare_document(REPAIR_REFERENCE)
    number = transfer.protocol_number + ("-R" if operation == "return" else "")
    _add_centered(document, f'{t["protocol"]} № {number}', 10.5, True)
    _add_centered(
        document,
        t["return_title"] if operation == "return" else t["issue_title"],
        10.5,
        True,
    )

    identity = document.add_table(rows=0, cols=2)
    identity.style = "Table Grid"
    identity.alignment = WD_TABLE_ALIGNMENT.CENTER
    identity.autofit = False
    for label, value in _identity_rows(transfer, operation, language):
        cells = identity.add_row().cells
        cells[0].width = Mm(46)
        cells[1].width = Mm(146)
        _set_cell(cells[0], f"{label}:", bold=True, size=8.5)
        _set_cell(cells[1], value, size=8.5)

    document.add_paragraph().paragraph_format.space_after = Pt(0)
    checklist = document.add_table(rows=1, cols=3)
    checklist.style = "Table Grid"
    checklist.alignment = WD_TABLE_ALIGNMENT.CENTER
    checklist.autofit = False
    widths = (Mm(10), Mm(66), Mm(116))
    for index, value in enumerate((t["number"], t["element"], t["condition"])):
        checklist.rows[0].cells[index].width = widths[index]
        _set_cell(checklist.rows[0].cells[index], value, bold=True, size=8.5)
        checklist.rows[0].cells[index].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_repeat_table_header(checklist.rows[0])
    for index, label in enumerate(CHECKLIST[language], start=1):
        cells = checklist.add_row().cells
        for cell_index, width in enumerate(widths):
            cells[cell_index].width = width
        _set_cell(cells[0], f"{index}.", size=8.5)
        _set_cell(cells[1], label, size=8.5)
        _set_cell(cells[2], "", size=8.5)

    overall_condition = (
        transfer.return_condition_text if operation == "return" else transfer.condition_text
    ) or ""
    condition = document.add_paragraph()
    condition.paragraph_format.space_before = Pt(2)
    condition.paragraph_format.space_after = Pt(1)
    run = condition.add_run(f'{t["overall_condition"]}: ')
    _set_run_font(run, 8.5, True)
    _set_run_font(condition.add_run(overall_condition), 8.5)

    usage = document.add_paragraph()
    usage.paragraph_format.space_before = Pt(1)
    usage.paragraph_format.space_after = Pt(1)
    usage_label = t["usage_return"] if operation == "return" else t["usage_issue"]
    _set_run_font(usage.add_run(f"{usage_label}: "), 8.5, True)
    _set_run_font(usage.add_run(_usage_text(transfer)), 8.5)

    remarks_value = _protocol_remarks(transfer, operation, language)
    remarks = document.add_paragraph()
    remarks.paragraph_format.space_before = Pt(1)
    remarks.paragraph_format.space_after = Pt(2)
    _set_run_font(remarks.add_run(f'{t["remarks"]}: '), 8.5, True)
    _set_run_font(remarks.add_run(remarks_value or ""), 8.5)

    signatures = document.add_table(rows=2, cols=3)
    signatures.style = "Table Grid"
    signatures.alignment = WD_TABLE_ALIGNMENT.CENTER
    signatures.autofit = False
    handed_label = t["returned"] if operation == "return" else t["handed"]
    handed_name = (
        transfer.returned_by_name if operation == "return" else transfer.handed_over_by
    ) or ""
    accepted_name = (
        transfer.return_accepted_by if operation == "return" else transfer.accepted_by
    ) or ""
    for row_index, (label, name) in enumerate(
        ((handed_label, handed_name), (t["accepted"], accepted_name))
    ):
        cells = signatures.rows[row_index].cells
        cells[0].width = Mm(38)
        cells[1].width = Mm(118)
        cells[2].width = Mm(36)
        _set_cell(cells[0], label, bold=True, size=8)
        _set_cell(cells[1], f"{name}\n({t['name']})", size=8)
        cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        _set_cell(cells[2], f"\n({t['signature']})", size=8)
        cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    footer = document.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer.paragraph_format.space_before = Pt(2)
    _set_run_font(
        footer.add_run(f'{t["batch"]}: {batch_reference} · {number}'), 7.5
    )
    output = io.BytesIO()
    document.save(output)
    return output.getvalue()


def build_protocol_docx(
    transfer: TransferProtocol, batch_reference: str, language: str = "bg"
) -> bytes:
    return _build_protocol_docx(
        transfer, batch_reference, operation="issue", language=language
    )


def build_return_protocol_docx(
    transfer: TransferProtocol, batch_reference: str, language: str = "bg"
) -> bytes:
    return _build_protocol_docx(
        transfer, batch_reference, operation="return", language=language
    )


def _register_pdf_fonts() -> tuple[str, str]:
    normal_candidates = [
        Path("C:/Windows/Fonts/arial.ttf"),
        Path("C:/Windows/Fonts/dejavusans.ttf"),
        Path("/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf"),
    ]
    bold_candidates = [
        Path("C:/Windows/Fonts/arialbd.ttf"),
        Path("C:/Windows/Fonts/dejavusans-bold.ttf"),
        Path("/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf"),
    ]
    normal_path = next((path for path in normal_candidates if path.is_file()), None)
    bold_path = next((path for path in bold_candidates if path.is_file()), None)
    if normal_path is None or bold_path is None:
        raise RuntimeError(
            "Липсва Unicode шрифт за генериране на документи на кирилица."
        )
    if "AssetCoreSans" not in pdfmetrics.getRegisteredFontNames():
        pdfmetrics.registerFont(TTFont("AssetCoreSans", str(normal_path)))
    if "AssetCoreSans-Bold" not in pdfmetrics.getRegisteredFontNames():
        pdfmetrics.registerFont(TTFont("AssetCoreSans-Bold", str(bold_path)))
    return "AssetCoreSans", "AssetCoreSans-Bold"


def _rina_image() -> io.BytesIO | None:
    if not REPAIR_REFERENCE.is_file():
        return None
    with zipfile.ZipFile(REPAIR_REFERENCE) as archive:
        candidates = [
            name
            for name in archive.namelist()
            if name.lower().endswith((".jpeg", ".jpg")) and name.startswith("word/media/")
        ]
        if not candidates:
            return None
        return io.BytesIO(archive.read(candidates[-1]))


def _pdf_styles():
    normal_font, bold_font = _register_pdf_fonts()
    styles = getSampleStyleSheet()
    body = ParagraphStyle(
        "AssetCoreBody",
        parent=styles["BodyText"],
        fontName=normal_font,
        fontSize=8.2,
        leading=10,
        textColor=colors.black,
    )
    label = ParagraphStyle(
        "AssetCoreLabel", parent=body, fontName=bold_font, fontSize=8.2, leading=10
    )
    title = ParagraphStyle(
        "AssetCoreTitle",
        parent=body,
        fontName=bold_font,
        fontSize=10.5,
        leading=12,
        alignment=TA_CENTER,
    )
    small = ParagraphStyle(
        "AssetCoreSmall", parent=body, fontSize=7.2, leading=8.5
    )
    return normal_font, bold_font, body, label, title, small


def _pdf_header() -> Table:
    _, _, body, _, _, _ = _pdf_styles()
    krz = ASSETS / "krz_logo.png"
    odessos = ASSETS / "odessos_logo.png"
    rina = _rina_image()
    left = ReportLabImage(str(krz), width=18 * mm, height=21 * mm) if krz.is_file() else Paragraph("KRZ", body)
    center = ReportLabImage(str(odessos), width=91 * mm, height=21 * mm) if odessos.is_file() else Paragraph("ODESSOS SHIPREPAIR &amp; CONVERSION", body)
    right = ReportLabImage(rina, width=23 * mm, height=20 * mm) if rina else Paragraph("RINA / AQAP 2110", body)
    table = Table([[left, center, right]], colWidths=[27 * mm, 119 * mm, 35 * mm], rowHeights=[23 * mm])
    table.setStyle(TableStyle([("GRID", (0, 0), (-1, -1), 0.65, colors.black), ("VALIGN", (0, 0), (-1, -1), "MIDDLE"), ("ALIGN", (0, 0), (-1, -1), "CENTER"), ("LEFTPADDING", (0, 0), (-1, -1), 2), ("RIGHTPADDING", (0, 0), (-1, -1), 2), ("TOPPADDING", (0, 0), (-1, -1), 1), ("BOTTOMPADDING", (0, 0), (-1, -1), 1)]))
    return table


def _pdf_table_style(*, header_rows: int = 0) -> TableStyle:
    commands = [
        ("GRID", (0, 0), (-1, -1), 0.55, colors.black),
        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
        ("LEFTPADDING", (0, 0), (-1, -1), 3),
        ("RIGHTPADDING", (0, 0), (-1, -1), 3),
        ("TOPPADDING", (0, 0), (-1, -1), 2),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 2),
    ]
    if header_rows:
        commands.extend([("BACKGROUND", (0, 0), (-1, header_rows - 1), colors.HexColor("#eeeeee")), ("ALIGN", (0, 0), (-1, header_rows - 1), "CENTER")])
    return TableStyle(commands)


def _build_protocol_pdf(
    transfer: TransferProtocol,
    batch_reference: str,
    *,
    operation: str,
    language: str,
) -> bytes:
    language = _language(language)
    t = TEXT[language]
    _, _, body, label, title, small = _pdf_styles()
    output = io.BytesIO()
    number = transfer.protocol_number + ("-R" if operation == "return" else "")
    pdf = SimpleDocTemplate(output, pagesize=A4, leftMargin=10 * mm, rightMargin=8 * mm, topMargin=7 * mm, bottomMargin=8 * mm, title=number, author="AssetCore")
    story = [
        _pdf_header(),
        Spacer(1, 2 * mm),
        Paragraph(f'{escape(t["protocol"])} № {escape(number)}', title),
        Paragraph(escape(t["return_title"] if operation == "return" else t["issue_title"]), title),
        Spacer(1, 2 * mm),
    ]
    identity = Table([[Paragraph(escape(label_text) + ":", label), Paragraph(escape(value), body)] for label_text, value in _identity_rows(transfer, operation, language)], colWidths=[46 * mm, 146 * mm])
    identity.setStyle(_pdf_table_style())
    story.extend([identity, Spacer(1, 3 * mm)])
    checklist_data = [[Paragraph(escape(t["number"]), label), Paragraph(escape(t["element"]), label), Paragraph(escape(t["condition"]), label)]]
    checklist_data.extend([[Paragraph(f"{index}.", body), Paragraph(escape(component), body), Paragraph(escape(value), body)] for index, (component, value) in enumerate(_checklist_rows(transfer, operation, language), start=1)])
    checklist = Table(checklist_data, colWidths=[10 * mm, 66 * mm, 116 * mm], repeatRows=1)
    checklist.setStyle(_pdf_table_style(header_rows=1))
    condition_value = (transfer.return_condition_text if operation == "return" else transfer.condition_text) or ""
    usage_label = t["usage_return"] if operation == "return" else t["usage_issue"]
    remarks_value = _protocol_remarks(transfer, operation, language)
    story.extend([checklist, Spacer(1, 2 * mm), Paragraph(f'<b>{escape(t["overall_condition"])}:</b> {escape(condition_value)}', body), Paragraph(f"<b>{escape(usage_label)}:</b> {escape(_usage_text(transfer))}", body), Paragraph(f'<b>{escape(t["remarks"])}:</b> {escape(remarks_value or "")}', body), Spacer(1, 3 * mm)])
    handed_label = t["returned"] if operation == "return" else t["handed"]
    handed_name = (transfer.returned_by_name if operation == "return" else transfer.handed_over_by) or ""
    accepted_name = (transfer.return_accepted_by if operation == "return" else transfer.accepted_by) or ""
    signature_data = [[Paragraph(escape(handed_label), label), Paragraph(f'{escape(handed_name)}<br/><font size="7">({escape(t["name"])})</font>', body), Paragraph(f'<br/><font size="7">({escape(t["signature"])})</font>', small)], [Paragraph(escape(t["accepted"]), label), Paragraph(f'{escape(accepted_name)}<br/><font size="7">({escape(t["name"])})</font>', body), Paragraph(f'<br/><font size="7">({escape(t["signature"])})</font>', small)]]
    signatures = Table(signature_data, colWidths=[38 * mm, 118 * mm, 36 * mm], rowHeights=[18 * mm, 18 * mm])
    signatures.setStyle(_pdf_table_style())
    story.extend([signatures, Spacer(1, 1.5 * mm), Paragraph(f'{escape(t["batch"])}: {escape(batch_reference)} · {escape(number)}', small)])
    pdf.build(story)
    return output.getvalue()


def build_protocol_pdf(
    transfer: TransferProtocol, batch_reference: str, language: str = "bg"
) -> bytes:
    return _build_protocol_pdf(
        transfer, batch_reference, operation="issue", language=language
    )


def build_return_protocol_pdf(
    transfer: TransferProtocol, batch_reference: str, language: str = "bg"
) -> bytes:
    return _build_protocol_pdf(
        transfer, batch_reference, operation="return", language=language
    )


def _add_section_title(document: Document, title: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(4)
    paragraph.paragraph_format.space_after = Pt(2)
    _keep_with_next(paragraph)
    _set_run_font(paragraph.add_run(title), 9.5, True)


def build_repair_protocol_docx(repair: Repair, language: str = "bg") -> bytes:
    language = _language(language)
    t = TEXT[language]
    document = _prepare_document(REPAIR_REFERENCE)
    _add_centered(document, f'{t["repair_protocol"]} № {repair.repair_reference or repair.id}', 11, True)
    machine_table = document.add_table(rows=0, cols=2)
    machine_table.style = "Table Grid"
    values = [
        (t["machine"], f"{repair.machine.name}; {repair.machine.brand} {repair.machine.model or ''}".strip()),
        (t["inventory"], repair.machine.inventory_number),
        (t["serial"], repair.machine.serial_number or ""),
        (t["date"], repair.opened_at.strftime("%d.%m.%Y")),
    ]
    for label_text, value in values:
        cells = machine_table.add_row().cells
        _set_cell(cells[0], label_text, bold=True)
        _set_cell(cells[1], value)
    for title_key, value in (
        ("reported_problem", repair.reported_problem),
        ("symptoms", repair.symptoms),
        ("condition_before", repair.condition_before),
        ("diagnosis", repair.diagnosis),
        ("required_work", repair.required_work),
        ("removed_parts", repair.removed_parts_text),
        ("cleaning", repair.diagnostic_cleaning),
        ("work", repair.work_performed),
    ):
        _add_section_title(document, t[title_key])
        paragraph = document.add_paragraph(value or "")
        paragraph.paragraph_format.space_after = Pt(1)
    if repair.events:
        _add_section_title(document, t["events"])
        event_table = document.add_table(rows=1, cols=3)
        event_table.style = "Table Grid"
        for cell, value in zip(event_table.rows[0].cells, (t["date"], t["condition"], t["description"]), strict=True):
            _set_cell(cell, value, bold=True, size=8)
        _set_repeat_table_header(event_table.rows[0])
        for event in sorted(repair.events, key=lambda item: (item.created_at, item.id)):
            cells = event_table.add_row().cells
            _set_cell(cells[0], event.created_at.strftime("%d.%m.%Y %H:%M"), size=8)
            _set_cell(cells[1], event.event_type, size=8)
            _set_cell(cells[2], event.description or "", size=8)
    if repair.parts_used:
        _add_section_title(document, t["parts_used"])
        parts = document.add_table(rows=1, cols=4)
        parts.style = "Table Grid"
        for cell, value in zip(parts.rows[0].cells, (t["position"], t["part_number"], t["description"], t["quantity"]), strict=True):
            _set_cell(cell, value, bold=True, size=8)
        _set_repeat_table_header(parts.rows[0])
        for index, part in enumerate(repair.parts_used, start=1):
            cells = parts.add_row().cells
            for cell, value in zip(cells, (str(index), part.part_number or "", part.description, f"{part.quantity:g} {part.unit or ''}".strip()), strict=True):
                _set_cell(cell, value, size=8)
    if repair.participants:
        _add_section_title(document, t.get("participants", "Допълнителни участници"))
        participants = document.add_table(rows=1, cols=4)
        participants.style = "Table Grid"
        for cell, value in zip(participants.rows[0].cells, (t.get("full_name", "Три имена"), t.get("job_title", "Длъжност"), t.get("contribution", "Участие"), t.get("duration", "Време")), strict=True):
            _set_cell(cell, value, bold=True, size=8)
        for participant in repair.participants:
            cells = participants.add_row().cells
            for cell, value in zip(cells, (participant.full_name_snapshot, participant.job_title_snapshot or "", participant.contribution or "", _repair_duration(participant.minutes_worked, language)), strict=True):
                _set_cell(cell, value, size=8)
    _add_section_title(document, t["test"])
    test_text = "; ".join(
        value
        for value in (
            f"cleaning={repair.cleaning_completed_at.strftime('%d.%m.%Y %H:%M')}" if repair.cleaning_completed_at else None,
            f"test={'PASS' if repair.test_passed else 'FAIL'}" if repair.test_passed is not None else None,
            f'{t["test_method"]}: {repair.test_method}' if repair.test_method else None,
            f'{t["test_pressure"]}: {repair.test_pressure_bar} bar' if repair.test_pressure_bar is not None else None,
            f'{t["leaks"]}: {t["yes"] if repair.leaks_detected else t["no"]}' if repair.leaks_detected is not None else None,
            f'{t["electrical_test"]}: {repair.electrical_test_result}' if repair.electrical_test_result else None,
            f'{t["functional_test"]}: {repair.functional_test_result}' if repair.functional_test_result else None,
            repair.test_details,
        )
        if value
    )
    document.add_paragraph(test_text)
    _add_section_title(document, t["condition_after"])
    document.add_paragraph("; ".join(value for value in (repair.condition_after, repair.result) if value))
    if repair.attachments:
        image_items = [item for item in repair.attachments if item.media_type in {"image/jpeg", "image/png"}]
        if image_items:
            document.add_page_break()
            for item in image_items:
                paragraph = document.add_paragraph()
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                try:
                    paragraph.add_run().add_picture(io.BytesIO(item.content), width=Mm(170))
                    caption = document.add_paragraph(item.caption or item.filename)
                    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
                except (ValueError, TypeError):
                    document.add_paragraph(item.filename)
    signatures = document.add_table(rows=1, cols=2)
    signatures.style = "Table Grid"
    responsible_identity = " · ".join(value for value in ((repair.responsible_user.full_name if repair.responsible_user else ""), (repair.responsible_user.job_title if repair.responsible_user else "")) if value)
    _set_cell(signatures.cell(0, 0), f'{t["responsible"]}: {responsible_identity}', bold=True)
    _set_cell(signatures.cell(0, 1), t["signature"])
    status = document.add_paragraph(_signature_status(language, finalized_internal=True))
    _set_run_font(status.runs[0], 8, True)
    output = io.BytesIO()
    document.save(output)
    return output.getvalue()


def _build_repair_protocol_pdf_legacy(repair: Repair, language: str = "bg") -> bytes:
    language = _language(language)
    t = TEXT[language]
    _, _, body, label, title, small = _pdf_styles()
    output = io.BytesIO()
    number = repair.repair_reference or f"REP-{repair.id:06d}"
    pdf = SimpleDocTemplate(output, pagesize=A4, leftMargin=10 * mm, rightMargin=8 * mm, topMargin=7 * mm, bottomMargin=8 * mm, title=number, author="AssetCore")
    story = [_pdf_header(), Spacer(1, 2 * mm), Paragraph(f'{escape(t["repair_protocol"])} № {escape(number)}', title), Spacer(1, 2 * mm)]
    identity_values = [(t["machine"], f"{repair.machine.name}; {repair.machine.brand} {repair.machine.model or ''}".strip()), (t["inventory"], repair.machine.inventory_number), (t["serial"], repair.machine.serial_number or ""), (t["date"], repair.opened_at.strftime("%d.%m.%Y"))]
    identity = Table([[Paragraph(escape(key), label), Paragraph(escape(value), body)] for key, value in identity_values], colWidths=[45 * mm, 147 * mm])
    identity.setStyle(_pdf_table_style())
    story.append(identity)
    for key, value in (("reported_problem", repair.reported_problem), ("symptoms", repair.symptoms), ("condition_before", repair.condition_before), ("diagnosis", repair.diagnosis), ("required_work", repair.required_work), ("removed_parts", repair.removed_parts_text), ("cleaning", repair.diagnostic_cleaning), ("work", repair.work_performed)):
        story.extend([Spacer(1, 2 * mm), Paragraph(escape(t[key]), label), Paragraph(escape(value or ""), body)])
    if repair.events:
        data = [[Paragraph(escape(t["date"]), label), Paragraph(escape(t["condition"]), label), Paragraph(escape(t["description"]), label)]]
        data.extend([[Paragraph(event.created_at.strftime("%d.%m.%Y %H:%M"), small), Paragraph(escape(event.event_type), small), Paragraph(escape(event.description or ""), small)] for event in sorted(repair.events, key=lambda item: (item.created_at, item.id))])
        table = Table(data, colWidths=[34 * mm, 40 * mm, 118 * mm], repeatRows=1)
        table.setStyle(_pdf_table_style(header_rows=1))
        story.extend([Spacer(1, 2 * mm), Paragraph(escape(t["events"]), label), table])
    if repair.parts_used:
        data = [[Paragraph(escape(value), label) for value in (t["position"], t["part_number"], t["description"], t["quantity"])]]
        data.extend([[Paragraph(str(index), small), Paragraph(escape(part.part_number or ""), small), Paragraph(escape(part.description), small), Paragraph(escape(f"{part.quantity:g} {part.unit or ''}".strip()), small)] for index, part in enumerate(repair.parts_used, start=1)])
        table = Table(data, colWidths=[14 * mm, 38 * mm, 110 * mm, 30 * mm], repeatRows=1)
        table.setStyle(_pdf_table_style(header_rows=1))
        story.extend([Spacer(1, 2 * mm), Paragraph(escape(t["parts_used"]), label), table])
    if repair.participants:
        data = [[Paragraph(escape(value), label) for value in (t.get("full_name", "Три имена"), t.get("job_title", "Длъжност"), t.get("contribution", "Участие"), t.get("duration", "Време"))]]
        data.extend([
            [Paragraph(escape(participant.full_name_snapshot), small), Paragraph(escape(participant.job_title_snapshot or ""), small), Paragraph(escape(participant.contribution or ""), small), Paragraph(escape(_repair_duration(participant.minutes_worked, language)), small)]
            for participant in repair.participants
        ])
        table = Table(data, colWidths=[55 * mm, 42 * mm, 70 * mm, 25 * mm], repeatRows=1)
        table.setStyle(_pdf_table_style(header_rows=1))
        story.extend([Spacer(1, 2 * mm), Paragraph(escape(t.get("participants", "Допълнителни участници")), label), table])
    test_text = _repair_test_summary(repair, language)
    story.extend([Spacer(1, 2 * mm), Paragraph(escape(t["test"]), label), Paragraph(escape(test_text), body), Spacer(1, 2 * mm), Paragraph(escape(t["condition_after"]), label), Paragraph(escape("; ".join(value for value in (repair.condition_after, repair.result) if value)), body)])
    if repair.attachments:
        for attachment in [item for item in repair.attachments if item.media_type in {"image/jpeg", "image/png"}]:
            story.extend([PageBreak(), ReportLabImage(io.BytesIO(attachment.content), width=170 * mm, height=220 * mm, kind="proportional"), Paragraph(escape(attachment.caption or attachment.filename), small)])
    responsible_identity = " · ".join(value for value in ((repair.responsible_user.full_name if repair.responsible_user else ""), (repair.responsible_user.job_title if repair.responsible_user else "")) if value)
    signature_data = [[Paragraph(f'<b>{escape(t["responsible"])}:</b> {escape(responsible_identity)}', body), Paragraph(escape(t["signature"]), body)]]
    signatures = Table(signature_data, colWidths=[156 * mm, 36 * mm], rowHeights=[14 * mm])
    signatures.setStyle(_pdf_table_style())
    story.extend([
        Spacer(1, 3 * mm),
        KeepTogether(signatures),
        Spacer(1, 2 * mm),
        Paragraph(escape(_signature_status(language, finalized_internal=True)), label),
    ])
    pdf.build(story)
    return output.getvalue()


def build_repair_protocol_pdf(
    repair: Repair,
    language: str = "bg",
    source_reference: str = "",
) -> bytes:
    """Three-part PDF fallback mirroring the controlled v6 DOCX package."""
    language = _language(language)
    _, _, body, label, title, small = _pdf_styles()
    labels = {
        "bg": {
            "accept": "ПРОТОКОЛ ЗА ПРИЕМАНЕ НА ОБОРУДВАНЕ ЗА РЕМОНТ",
            "done": "ПРОТОКОЛ ЗА ИЗВЪРШЕН РЕМОНТ",
            "number": "Протокол №", "date": "Дата", "equipment": "Оборудване",
            "ownership": "Собственост", "condition": "Състояние при приемане",
            "problem": "Описание на проблема", "cleaning": "Почистване при диагностиката",
            "required": "Описание на необходимия ремонт", "removed": "А) Демонтаж и подготовка",
            "diagnosis": "Б) Констатирано след разглобяване, почистване и дефектация",
            "needed": "В) Нужни части за ремонт", "diagnosis_time": "Реално време за диагностика",
            "participants": "Извършили ремонта / диагностиката", "handed": "Предал оборудването",
            "accepted": "Приел оборудването", "work": "Извършени ремонтни дейности",
            "repair_time": "Реално време за ремонт", "testing_time": "Реално време за тестове",
            "parts": "Вложени резервни части", "hours": "РЕАЛНО ОТРАБОТЕНО ВРЕМЕ",
            "start": "Начало", "end": "Край", "total": "Общо реално време",
            "participant_total": "Общо време на участниците",
            "tests": "Тестове и резултат", "after": "Състояние след ремонта",
            "result": "Краен резултат", "attachments": "Приложения",
            "approved": "Приел ремонта", "reference": "Ремонтна референция",
            "source": "Свързано връщане",
        },
        "en": {
            "accept": "EQUIPMENT ACCEPTANCE FOR REPAIR PROTOCOL", "done": "COMPLETED REPAIR PROTOCOL",
            "number": "Protocol No.", "date": "Date", "equipment": "Equipment", "ownership": "Ownership",
            "condition": "Condition on acceptance", "problem": "Reported problem",
            "cleaning": "Cleaning during diagnosis", "required": "Required repair", "removed": "A) Disassembly and preparation",
            "diagnosis": "B) Findings after inspection", "needed": "C) Parts required", "diagnosis_time": "Actual diagnosis time",
            "participants": "Repair / diagnosis participants", "handed": "Handed over by", "accepted": "Accepted by",
            "work": "Repair work performed", "repair_time": "Actual repair time", "testing_time": "Actual testing time",
            "parts": "Spare parts used", "hours": "ACTUAL LABOUR TIME", "start": "Start", "end": "End", "total": "Total actual time",
            "participant_total": "Total participant labour time",
            "tests": "Tests and result", "after": "Condition after repair", "result": "Final result", "attachments": "Attachments",
            "approved": "Repair accepted by", "reference": "Repair reference", "source": "Linked return",
        },
        "ru": {
            "accept": "ПРОТОКОЛ ПРИЕМА ОБОРУДОВАНИЯ В РЕМОНТ", "done": "ПРОТОКОЛ ВЫПОЛНЕННОГО РЕМОНТА",
            "number": "Протокол №", "date": "Дата", "equipment": "Оборудование", "ownership": "Собственность",
            "condition": "Состояние при приеме", "problem": "Описание проблемы",
            "cleaning": "Очистка при диагностике", "required": "Необходимый ремонт", "removed": "А) Демонтаж и подготовка",
            "diagnosis": "Б) Результаты дефектации", "needed": "В) Необходимые детали", "diagnosis_time": "Фактическое время диагностики",
            "participants": "Участники ремонта / диагностики", "handed": "Передал", "accepted": "Принял",
            "work": "Выполненные работы", "repair_time": "Фактическое время ремонта", "testing_time": "Фактическое время испытаний",
            "parts": "Использованные детали", "hours": "ФАКТИЧЕСКОЕ РАБОЧЕЕ ВРЕМЯ", "start": "Начало", "end": "Окончание", "total": "Общее время",
            "participant_total": "Общее время участников",
            "tests": "Испытания и результат", "after": "Состояние после ремонта", "result": "Итоговый результат", "attachments": "Приложения",
            "approved": "Ремонт принял", "reference": "Референция ремонта", "source": "Связанный возврат",
        },
    }[language]
    output = io.BytesIO()
    number = repair.repair_reference or f"REP-{repair.id:06d}"
    pdf = SimpleDocTemplate(
        output, pagesize=A4, leftMargin=10 * mm, rightMargin=8 * mm,
        topMargin=7 * mm, bottomMargin=8 * mm, title=number, author="AssetCore",
    )

    def field_table(rows: list[tuple[str, str]]) -> Table:
        table = Table(
            [[Paragraph(escape(key), label), Paragraph(escape(value or ""), body)] for key, value in rows],
            colWidths=[48 * mm, 144 * mm],
        )
        table.setStyle(_pdf_table_style())
        return table

    def box(heading: str, value: str | None) -> Table:
        table = Table(
            [[Paragraph(escape(heading), label)], [Paragraph(escape(value or ""), body)]],
            colWidths=[192 * mm],
        )
        table.setStyle(_pdf_table_style(header_rows=1))
        return table

    participant_text = "; ".join(
        " — ".join(
            value for value in (
                item.full_name_snapshot, item.job_title_snapshot, item.contribution,
                _repair_duration(item.minutes_worked, language),
            ) if value
        )
        for item in repair.participants
    )
    machine_text = (
        f"{repair.machine.name} · №{repair.machine.inventory_number} · "
        f"{repair.machine.brand} {repair.machine.model or ''} · S/N {repair.machine.serial_number or ''}"
    )
    total_minutes = sum(
        value or 0 for value in (
            repair.diagnosis_minutes, repair.repair_minutes, repair.testing_minutes
        )
    )
    participant_total_minutes = sum(
        item.minutes_worked or 0 for item in repair.participants
    )
    story: list[object] = [
        _pdf_header(), Spacer(1, 2 * mm), Paragraph(escape(labels["accept"]), title),
        field_table([
            (labels["number"], number), (labels["date"], repair.opened_at.strftime("%d.%m.%Y")),
            (labels["reference"], number), (labels["source"], source_reference),
            (labels["equipment"], machine_text), (labels["ownership"], repair.machine.ownership or ""),
        ]), Spacer(1, 2 * mm), box(labels["problem"], repair.reported_problem),
        Spacer(1, 2 * mm), box(labels["condition"], repair.condition_before),
        Spacer(1, 2 * mm), box(labels["required"], repair.required_work),
        Spacer(1, 2 * mm), box(labels["removed"], repair.removed_parts_text),
        Spacer(1, 2 * mm), box(labels["cleaning"], repair.diagnostic_cleaning),
        Spacer(1, 2 * mm), box(labels["diagnosis"], repair.diagnosis),
        Spacer(1, 2 * mm), box(labels["needed"], repair.required_parts_text),
        Spacer(1, 2 * mm), field_table([
            (labels["diagnosis_time"], _repair_duration(repair.diagnosis_minutes, language)),
            (labels["participants"], participant_text),
            (labels["handed"], repair.reported_by_name or ""),
            (labels["accepted"], repair.accepted_by.full_name if repair.accepted_by else ""),
        ]), PageBreak(), _pdf_header(), Spacer(1, 2 * mm),
        Paragraph(escape(labels["done"]), title),
        field_table([
            (labels["number"], number),
            (labels["date"], (repair.closed_at or repair.opened_at).strftime("%d.%m.%Y")),
            (labels["equipment"], machine_text), (labels["ownership"], repair.machine.ownership or ""),
        ]), Spacer(1, 2 * mm), box(labels["work"], repair.work_performed),
        Spacer(1, 2 * mm), field_table([
            (labels["repair_time"], _repair_duration(repair.repair_minutes, language)),
            (labels["testing_time"], _repair_duration(repair.testing_minutes, language)),
            (labels["participants"], participant_text),
        ]), Spacer(1, 2 * mm), Paragraph(escape(labels["parts"]), label),
    ]
    part_rows = [[Paragraph(escape(value), label) for value in ("№", "Part No.", "Описание", "Количество")]]
    part_rows.extend([
        [Paragraph(str(index), small), Paragraph(escape(part.part_number or ""), small),
         Paragraph(escape(part.description), small), Paragraph(escape(f"{part.quantity:g} {part.unit or ''}".strip()), small)]
        for index, part in enumerate(repair.parts_used, start=1)
    ])
    if len(part_rows) == 1:
        part_rows.append([Paragraph("—", small), Paragraph("", small), Paragraph("", small), Paragraph("", small)])
    parts_table = Table(part_rows, colWidths=[12 * mm, 40 * mm, 108 * mm, 32 * mm], repeatRows=1)
    parts_table.setStyle(_pdf_table_style(header_rows=1))
    story.extend([
        parts_table, PageBreak(), _pdf_header(), Spacer(1, 2 * mm),
        Paragraph(escape(labels["hours"]), title),
        field_table([
            (labels["start"], repair.started_at.strftime("%d.%m.%Y %H:%M") if repair.started_at else ""),
            (labels["end"], repair.closed_at.strftime("%d.%m.%Y %H:%M") if repair.closed_at else ""),
            (labels["total"], _repair_duration(total_minutes, language)),
            (labels["participant_total"], _repair_duration(participant_total_minutes, language)),
        ]), Spacer(1, 2 * mm), box(labels["tests"], _repair_test_summary(repair, language)),
        Spacer(1, 2 * mm), box(labels["after"], repair.condition_after),
        Spacer(1, 2 * mm), box(labels["result"], repair.result),
        Spacer(1, 2 * mm), box(labels["attachments"], "; ".join(item.filename for item in repair.attachments)),
        Spacer(1, 2 * mm), field_table([
            (labels["participants"], participant_text),
            (labels["approved"], " — ".join(value for value in (
                repair.approved_by.full_name if repair.approved_by else "",
                repair.approved_by.job_title if repair.approved_by else "",
            ) if value)),
            (labels["date"], (repair.closed_at or repair.opened_at).strftime("%d.%m.%Y")),
        ]),
    ])
    pdf.build(story)
    return output.getvalue()


def _request_snapshot(request: PartRequest) -> dict:
    return {
        "request_id": request.id,
        "request_reference": request.request_reference,
        "machine_id": request.machine_id,
        "machine_number": request.machine.inventory_number if request.machine else None,
        "repair_id": request.repair_id,
        "repair_reference": request.repair.repair_reference if request.repair else None,
        "priority": request.priority,
        "status": request.status,
        "language": request.language,
        "reason": request.reason,
        "department": request.department,
        "supplier": request.supplier,
        "delivery_note": request.delivery_note,
        "ordered_at": request.ordered_at.isoformat() if request.ordered_at else None,
        "delivered_at": request.delivered_at.isoformat() if request.delivered_at else None,
        "requested_by_id": request.requested_by_id,
        "submitted_at": request.submitted_at.isoformat() if request.submitted_at else None,
        "decided_at": request.decided_at.isoformat() if request.decided_at else None,
        "decision_note": request.decision_note,
        "lines": [
            {
                "catalog_part_id": line.catalog_part_id,
                "position": line.position,
                "part_number": line.part_number,
                "description": line.description,
                "quantity": line.quantity,
                "unit": line.unit,
                "source_document": line.source_document,
                "source_page": line.source_page,
                "delivered_quantity": line.delivered_quantity,
                "is_unknown_part": line.is_unknown_part,
                "assembly": line.assembly,
                "note": line.note,
                "linked_catalog_part_id": line.linked_catalog_part_id,
                "linked_part_number": line.linked_catalog_part.part_number if line.linked_catalog_part else None,
                "linked_at": line.linked_at.isoformat() if line.linked_at else None,
            }
            for line in request.lines
        ],
    }


def _part_request_line_description(line: PartRequestLine, language: str) -> str:
    if not line.is_unknown_part:
        return line.description
    label = {
        "bg": "Част без потвърден part number",
        "en": "Part without a confirmed part number",
        "ru": "Деталь без подтверждённого part number",
    }[_language(language)]
    assembly = {"bg": "Възел", "en": "Assembly", "ru": "Узел"}[_language(language)]
    linked = ""
    if line.linked_catalog_part:
        linked_label = {"bg": "Свързана с", "en": "Linked to", "ru": "Связана с"}[_language(language)]
        linked = f"; {linked_label}: {line.linked_catalog_part.part_number}"
    return f"[{label}] {assembly}: {line.assembly or '-'}; {line.description}{linked}"


def build_part_request_docx(request: PartRequest, language: str = "bg") -> bytes:
    language = _language(language)
    t = TEXT[language]
    document = _prepare_document(PARTS_REFERENCE)
    for line in t["part_request_title"].splitlines():
        _add_centered(document, line, 10.5, True)
    if request.machine:
        machine = request.machine
        for label_text, value in ((t["machine"], f"{machine.name}; {machine.brand} {machine.model or ''}".strip()), (t["inventory"], machine.inventory_number), (t["serial"], machine.serial_number or "")):
            paragraph = document.add_paragraph()
            paragraph.paragraph_format.space_after = Pt(0)
            _set_run_font(paragraph.add_run(f"{label_text}: "), 8.5, True)
            _set_run_font(paragraph.add_run(value), 8.5)
    document.add_paragraph()
    table = document.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = (t["position"], t["part_number"], t["description"], t["quantity"])
    widths = (Mm(14), Mm(31), Mm(127), Mm(20))
    for index, value in enumerate(headers):
        table.rows[0].cells[index].width = widths[index]
        _set_cell(table.rows[0].cells[index], value, bold=True, size=8)
        table.rows[0].cells[index].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_repeat_table_header(table.rows[0])
    for line in request.lines:
        cells = table.add_row().cells
        values = (line.position or "", line.part_number or "", _part_request_line_description(line, language), f"{line.quantity:g} {line.unit or ''}".strip())
        for index, value in enumerate(values):
            cells[index].width = widths[index]
            _set_cell(cells[index], value, size=8)
    if request.reason:
        paragraph = document.add_paragraph()
        _set_run_font(paragraph.add_run(f'{t["remarks"]}: '), 8.5, True)
        _set_run_font(paragraph.add_run(request.reason), 8.5)
    provenance = [f"{line.source_document}, p. {line.source_page}" for line in request.lines if line.source_document and line.source_page]
    if provenance:
        paragraph = document.add_paragraph()
        _set_run_font(paragraph.add_run(f'{t["source"]}: '), 8, True)
        _set_run_font(paragraph.add_run("; ".join(provenance)), 8)
    footer = document.add_table(rows=1, cols=3)
    footer.style = "Table Grid"
    request_ref = request.request_reference or f"PR-{request.id:06d}"
    requester = request.requested_by.full_name if request.requested_by else ""
    decision = request.decision_note or request.status
    for cell, value in zip(footer.rows[0].cells, (f'{t["request_number"]}: {request_ref}', f'{t["date"]}: {request.created_at:%d.%m.%Y}', f'{t["requester"]}: {requester}\n{t["decision"]}: {decision}'), strict=True):
        _set_cell(cell, value, size=8)
    output = io.BytesIO()
    document.save(output)
    return output.getvalue()


def build_part_request_pdf(request: PartRequest, language: str = "bg") -> bytes:
    language = _language(language)
    t = TEXT[language]
    _, _, body, label, title, small = _pdf_styles()
    output = io.BytesIO()
    request_ref = request.request_reference or f"PR-{request.id:06d}"
    pdf = SimpleDocTemplate(output, pagesize=A4, leftMargin=10 * mm, rightMargin=8 * mm, topMargin=7 * mm, bottomMargin=8 * mm, title=request_ref, author="AssetCore")
    story = [_pdf_header(), Spacer(1, 3 * mm)]
    story.extend(Paragraph(escape(line), title) for line in t["part_request_title"].splitlines())
    if request.machine:
        machine = request.machine
        for label_text, value in ((t["machine"], f"{machine.name}; {machine.brand} {machine.model or ''}".strip()), (t["inventory"], machine.inventory_number), (t["serial"], machine.serial_number or "")):
            story.append(Paragraph(f"<b>{escape(label_text)}:</b> {escape(value)}", body))
    story.append(Spacer(1, 3 * mm))
    data = [[Paragraph(escape(value), label) for value in (t["position"], t["part_number"], t["description"], t["quantity"])]]
    data.extend([[Paragraph(escape(line.position or ""), small), Paragraph(escape(line.part_number or ""), small), Paragraph(escape(_part_request_line_description(line, language)), small), Paragraph(escape(f"{line.quantity:g} {line.unit or ''}".strip()), small)] for line in request.lines])
    table = Table(data, colWidths=[14 * mm, 31 * mm, 127 * mm, 20 * mm], repeatRows=1)
    table.setStyle(_pdf_table_style(header_rows=1))
    story.append(table)
    if request.reason:
        story.extend([Spacer(1, 2 * mm), Paragraph(f'<b>{escape(t["remarks"])}:</b> {escape(request.reason)}', body)])
    provenance = [f"{line.source_document}, p. {line.source_page}" for line in request.lines if line.source_document and line.source_page]
    if provenance:
        story.append(Paragraph(f'<b>{escape(t["source"])}:</b> {escape("; ".join(provenance))}', small))
    requester = request.requested_by.full_name if request.requested_by else ""
    decision = request.decision_note or request.status
    footer = Table([[Paragraph(f'{escape(t["request_number"])}: {escape(request_ref)}', small), Paragraph(f'{escape(t["date"])}: {request.created_at:%d.%m.%Y}', small), Paragraph(f'{escape(t["requester"])}: {escape(requester)}<br/>{escape(t["decision"])}: {escape(decision)}', small)]], colWidths=[64 * mm, 52 * mm, 76 * mm])
    footer.setStyle(_pdf_table_style())
    story.extend([Spacer(1, 4 * mm), footer])
    pdf.build(story)
    return output.getvalue()


def _template_version(
    db: Session, document_type: str, language: str
) -> DocumentTemplateVersion:
    now = datetime.now(UTC).replace(tzinfo=None)
    version = db.scalar(
        select(DocumentTemplateVersion)
        .join(DocumentTemplate)
        .where(
            DocumentTemplate.document_type == document_type,
            DocumentTemplate.is_active.is_(True),
            DocumentTemplateVersion.language == _language(language),
            DocumentTemplateVersion.is_published.is_(True),
            DocumentTemplateVersion.validation_status == "PASSED",
            or_(
                DocumentTemplateVersion.effective_from.is_(None),
                DocumentTemplateVersion.effective_from <= now,
            ),
            or_(
                DocumentTemplateVersion.effective_to.is_(None),
                DocumentTemplateVersion.effective_to > now,
            ),
        )
        .order_by(DocumentTemplateVersion.version.desc())
    )
    if version is None:
        raise ConfirmedTemplateUnavailableError(document_type, language)
    return version


def _preparer_values(db: Session, created_by_id: int) -> dict[str, str]:
    user = db.get(User, created_by_id)
    if user is None:
        raise TemplateValidationError("Съставителят на документа не е намерен.")
    complete = bool(
        user.first_name
        and (
            user.middle_name
            or (
                user.legal_name_exception
                and user.legal_name_exception_reason
                and user.legal_name_exception_approved_by_id
                and user.legal_name_exception_approved_at
            )
        )
        and user.last_name
        and user.job_title
        and user.profile_status == "PROFILE_COMPLETE"
    )
    if not complete:
        raise TemplateValidationError(
            "Профилът на съставителя трябва да съдържа потвърдени три имена и длъжност."
        )
    return {"PREPARER_NAME": user.full_name, "PREPARER_JOB_TITLE": user.job_title}


def _signature_status(language: str, *, finalized_internal: bool = False) -> str:
    if finalized_internal:
        return {
            "bg": "ОКОНЧАТЕЛЕН ВЪТРЕШЕН ПРОТОКОЛ",
            "en": "FINAL INTERNAL PROTOCOL",
            "ru": "ОКОНЧАТЕЛЬНЫЙ ВНУТРЕННИЙ ПРОТОКОЛ",
        }[_language(language)]
    return {
        "bg": "НЕПЪЛНО ПОДПИСАН",
        "en": "NOT FULLY SIGNED",
        "ru": "ПОДПИСАН НЕ ПОЛНОСТЬЮ",
    }[_language(language)]


def _protocol_template_values(
    db: Session,
    transfer: TransferProtocol,
    batch_reference: str,
    operation: str,
    language: str,
    created_by_id: int,
) -> dict[str, object]:
    language = _language(language)
    t = TEXT[language]
    date_value = (
        transfer.returned_at if operation == "return" else transfer.issued_at
    ) or transfer.created_at
    left_name = transfer.returned_by_name if operation == "return" else transfer.handed_over_by
    right_name = transfer.return_accepted_by if operation == "return" else transfer.accepted_by
    left_job = (
        transfer.returned_by_job_title
        if operation == "return"
        else transfer.handed_over_job_title
    )
    right_job = (
        transfer.return_accepted_job_title
        if operation == "return"
        else transfer.accepted_by_job_title
    )
    if operation == "return":
        left_role = t["returned"]
        right_role = " · ".join(value for value in (t["accepted"], right_job) if value)
    else:
        left_role = " · ".join(value for value in (t["handed"], left_job) if value)
        right_role = t["accepted"]
    values: dict[str, object] = {
        "DOCUMENT_NUMBER": f"{transfer.protocol_number}-R" if operation == "return" else transfer.protocol_number,
        "CREATION_DATE": date_value.strftime("%d.%m.%Y"),
        "EQUIPMENT_TYPE": transfer.machine.category or "HPWJ",
        "MACHINE_NAME": transfer.machine.name,
        "MACHINE_NUMBER": transfer.machine.inventory_number,
        "BRAND": transfer.machine.brand,
        "MODEL": transfer.machine.model or "",
        "MODEL_DISPLAY": _machine_model(transfer),
        "SERIAL_NUMBER": transfer.machine.serial_number or "",
        "PRESSURE_BAR": transfer.machine.pressure_bar,
        "BATCH_REFERENCE": batch_reference,
        "CONDITION_LABEL": t["overall_condition"],
        "CONDITION_TEXT": transfer.return_condition_text if operation == "return" else transfer.condition_text,
        "USAGE_TEXT": _usage_text(transfer),
        "REMARKS": _protocol_remarks(transfer, operation, language),
        "LEFT_SIGNER_NAME": left_name or "",
        "LEFT_SIGNER_JOB_TITLE": left_job or "",
        "LEFT_SIGNER_ROLE": left_role or "",
        "RIGHT_SIGNER_NAME": right_name or "",
        "RIGHT_SIGNER_JOB_TITLE": right_job or "",
        "RIGHT_SIGNER_ROLE": right_role or "",
        "LEFT_SIGNATURE": "[[ASSETCORE_LEFT_SIGNATURE]]",
        "RIGHT_SIGNATURE": "[[ASSETCORE_RIGHT_SIGNATURE]]",
        "SIGNATURE_STATUS": _signature_status(language),
    }
    checklist_rows = _checklist_rows(transfer, operation, language)
    for index in range(1, 11):
        values[f"CHECK_{index}"] = (
            checklist_rows[index - 1][1] if index <= len(checklist_rows) else ""
        )
    values.update(_preparer_values(db, created_by_id))
    return values

def _next_generated_number(db: Session, base: str) -> str:
    existing = db.scalars(
        select(GeneratedDocument.document_number)
        .where(GeneratedDocument.document_number.like(f"{base}%"))
        .distinct()
    ).all()
    if base not in existing:
        return base
    version = 2
    while f"{base}-V{version}" in existing:
        version += 1
    return f"{base}-V{version}"


def _generated_documents(
    *,
    number: str,
    document_type: str,
    language: str,
    template_version: DocumentTemplateVersion | None,
    docx: bytes,
    pdf: bytes,
    snapshot: dict,
    created_by_id: int,
    machine_id: int | None = None,
    repair_id: int | None = None,
    part_request_id: int | None = None,
    transfer_id: int | None = None,
    batch_id: int | None = None,
) -> list[GeneratedDocument]:
    documents = []
    stem = safe_filename(number)
    for format_name, media_type, content in (
        ("docx", DOCX_MEDIA_TYPE, docx),
        ("pdf", PDF_MEDIA_TYPE, pdf),
    ):
        documents.append(
            GeneratedDocument(
                document_number=number,
                document_type=document_type,
                format=format_name,
                language=_language(language),
                filename=f"{stem}.{format_name}",
                media_type=media_type,
                content=content,
                sha256=hashlib.sha256(content).hexdigest(),
                template_version_id=template_version.id if template_version else None,
                machine_id=machine_id,
                repair_id=repair_id,
                part_request_id=part_request_id,
                transfer_id=transfer_id,
                batch_id=batch_id,
                snapshot=snapshot,
                created_by_id=created_by_id,
            )
        )
    return documents


def _register_official_version(
    db: Session,
    *,
    number: str,
    document_type: str,
    language: str,
    docx: bytes,
    pdf: bytes,
    snapshot: dict,
    created_by_id: int,
    machine_id: int | None = None,
    transfer_id: int | None = None,
    batch_id: int | None = None,
    template_version_id: int | None = None,
    initial_status: str = OfficialDocumentStatus.DRAFT.value,
) -> OfficialDocument:
    existing = db.scalar(
        select(OfficialDocument).where(OfficialDocument.document_number == number)
    )
    if existing is not None:
        raise TemplateValidationError(
            f"Официален документ с номер {number} вече съществува и няма да бъде презаписан."
        )
    preparer = db.get(User, created_by_id)
    preparer_values = _preparer_values(db, created_by_id)
    official_snapshot = dict(snapshot)
    official_snapshot["prepared_by"] = {
        "user_id": preparer.id,
        "first_name": preparer.first_name,
        "middle_name": preparer.middle_name,
        "last_name": preparer.last_name,
        "display_name": preparer_values["PREPARER_NAME"],
        "job_title": preparer_values["PREPARER_JOB_TITLE"],
        "department_id": preparer.department_id,
        "department": preparer.profile_department.name_bg if preparer.profile_department else None,
        "operation_role": "PREPARER",
        "captured_at": utcnow().isoformat(timespec="seconds") + "Z",
    }
    document = OfficialDocument(
        document_number=number,
        document_type=document_type,
        machine_id=machine_id,
        transfer_id=transfer_id,
        batch_id=batch_id,
        created_by_id=created_by_id,
    )
    db.add(document)
    db.flush()
    snapshot_sha256 = hashlib.sha256(
        json.dumps(
            official_snapshot,
            ensure_ascii=False,
            sort_keys=True,
            separators=(",", ":"),
            default=str,
        ).encode("utf-8")
    ).hexdigest()
    docx_sha256 = hashlib.sha256(docx).hexdigest()
    pdf_sha256 = hashlib.sha256(pdf).hexdigest()
    signing_sha256 = hashlib.sha256(
        json.dumps(
            {
                "document_number": number,
                "document_type": document_type,
                "snapshot_sha256": snapshot_sha256,
                "docx_sha256": docx_sha256,
                "pdf_sha256": pdf_sha256,
                "version": 1,
            },
            sort_keys=True,
            separators=(",", ":"),
        ).encode("utf-8")
    ).hexdigest()
    version = OfficialDocumentVersion(
        document_id=document.id,
        version=1,
        status=initial_status,
        language=_language(language),
        template_version_id=template_version_id,
        snapshot=official_snapshot,
        snapshot_sha256=snapshot_sha256,
        signing_sha256=signing_sha256,
        docx_content=docx,
        docx_sha256=docx_sha256,
        pdf_content=pdf,
        pdf_sha256=pdf_sha256,
        prepared_by_id=created_by_id,
        finalized_at=(
            utcnow()
            if initial_status == OfficialDocumentStatus.FINALIZED.value
            else None
        ),
    )
    db.add(version)
    db.flush()
    set_current_version(db, document, version)
    return document


def make_protocol_documents(
    db: Session,
    transfer: TransferProtocol,
    batch: TransferBatch,
    created_by_id: int,
    language: str = "bg",
) -> list[ProtocolDocument]:
    template = _template_version(db, DocumentType.TRANSFER_ISSUE.value, language)
    snapshot = _protocol_snapshot(
        transfer, batch.batch_reference, "issue", _language(language)
    )
    stem = safe_filename(transfer.protocol_number)
    docx = render_docx(
        template,
        _protocol_template_values(
            db, transfer, batch.batch_reference, "issue", language, created_by_id
        ),
    )
    pdf = convert_docx_to_pdf(docx) or build_protocol_pdf(
        transfer, batch.batch_reference, language
    )
    _register_official_version(
        db,
        number=transfer.protocol_number,
        document_type=DocumentType.TRANSFER_ISSUE.value,
        language=language,
        docx=docx,
        pdf=pdf,
        snapshot=snapshot,
        created_by_id=created_by_id,
        machine_id=transfer.machine_id,
        transfer_id=transfer.id,
        batch_id=batch.id,
        template_version_id=template.id,
    )
    generated = [
        ("docx", DOCX_MEDIA_TYPE, docx),
        ("pdf", PDF_MEDIA_TYPE, pdf),
    ]
    return [
        ProtocolDocument(
            transfer_id=transfer.id,
            machine_id=transfer.machine_id,
            batch_id=batch.id,
            format=format_name,
            filename=f"{stem}.{format_name}",
            media_type=media_type,
            content=content,
            sha256=hashlib.sha256(content).hexdigest(),
            document_number=transfer.protocol_number,
            language=_language(language),
            template_version_id=template.id if template else None,
            snapshot=snapshot,
            created_by_id=created_by_id,
            created_at=transfer.created_at,
        )
        for format_name, media_type, content in generated
    ]


def make_return_documents(
    db: Session,
    transfer: TransferProtocol,
    batch: TransferBatch | None,
    created_by_id: int,
    language: str = "bg",
) -> list[GeneratedDocument]:
    template = _template_version(db, DocumentType.TRANSFER_RETURN.value, language)
    batch_reference = batch.batch_reference if batch else "-"
    base = f"{transfer.protocol_number}-R"
    number = _next_generated_number(db, base)
    values = _protocol_template_values(
        db, transfer, batch_reference, "return", language, created_by_id
    )
    values["DOCUMENT_NUMBER"] = number
    docx = render_docx(template, values)
    pdf = convert_docx_to_pdf(docx) or build_return_protocol_pdf(
        transfer, batch_reference, language
    )
    snapshot = _protocol_snapshot(transfer, batch_reference, "return", _language(language))
    _register_official_version(
        db,
        number=number,
        document_type=DocumentType.TRANSFER_RETURN.value,
        language=language,
        docx=docx,
        pdf=pdf,
        snapshot=snapshot,
        created_by_id=created_by_id,
        machine_id=transfer.machine_id,
        transfer_id=transfer.id,
        batch_id=batch.id if batch else None,
        template_version_id=template.id,
    )
    return _generated_documents(
        number=number,
        document_type=DocumentType.TRANSFER_RETURN.value,
        language=language,
        template_version=template,
        docx=docx,
        pdf=pdf,
        snapshot=snapshot,
        created_by_id=created_by_id,
        machine_id=transfer.machine_id,
        transfer_id=transfer.id,
        batch_id=batch.id if batch else None,
    )


def _repair_test_summary(repair: Repair, language: str = "bg") -> str:
    language = _language(language)
    t = TEXT[language]
    values = (
        f'{t["test_method"]}: {repair.test_method}' if repair.test_method else None,
        f'{t["test_pressure"]}: {repair.test_pressure_bar} bar' if repair.test_pressure_bar is not None else None,
        f'{t["leaks"]}: {t["yes"] if repair.leaks_detected else t["no"]}' if repair.leaks_detected is not None else None,
        f'{t["electrical_test"]}: {repair.electrical_test_result}' if repair.electrical_test_result else None,
        f'{t["functional_test"]}: {repair.functional_test_result}' if repair.functional_test_result else None,
        repair.test_details,
        f'{ {"bg": "Краен резултат", "en": "Final result", "ru": "Итоговый результат"}[language]}: {repair.result}' if repair.result else None,
    )
    passed = None
    if repair.test_passed is not None:
        passed = f'{t["test"]}: {t["yes"] if repair.test_passed else t["no"]}'
    return "; ".join(value for value in (passed, *values) if value)


def _repair_duration(minutes: int | None, language: str = "bg") -> str:
    if minutes is None:
        return ""
    hours, remainder = divmod(minutes, 60)
    labels = {
        "bg": ("ч", "мин"),
        "en": ("h", "min"),
        "ru": ("ч", "мин"),
    }[_language(language)]
    parts = []
    if hours:
        parts.append(f"{hours} {labels[0]}")
    if remainder or not parts:
        parts.append(f"{remainder} {labels[1]}")
    return " ".join(parts)


def make_repair_documents(
    db: Session, repair: Repair, created_by_id: int, language: str = "bg"
) -> list[GeneratedDocument]:
    template = _template_version(db, DocumentType.REPAIR_PROTOCOL.value, language)
    base = repair.repair_reference or f"REP-{repair.id:06d}"
    number = _next_generated_number(db, base)
    snapshot = {
        "repair_id": repair.id,
        "repair_reference": repair.repair_reference,
        "machine_id": repair.machine_id,
        "machine_number": repair.machine.inventory_number,
        "status": repair.status,
        "reported_problem": repair.reported_problem,
        "condition_before": repair.condition_before,
        "condition_after": repair.condition_after,
        "reported_by_name": repair.reported_by_name,
        "symptoms": repair.symptoms,
        "required_work": repair.required_work,
        "required_parts_text": repair.required_parts_text,
        "removed_parts_text": repair.removed_parts_text,
        "diagnostic_cleaning": repair.diagnostic_cleaning,
        "diagnosis": repair.diagnosis,
        "work_performed": repair.work_performed,
        "result": repair.result,
        "test_passed": repair.test_passed,
        "test_method": repair.test_method,
        "test_pressure_bar": repair.test_pressure_bar,
        "leaks_detected": repair.leaks_detected,
        "electrical_test_result": repair.electrical_test_result,
        "functional_test_result": repair.functional_test_result,
        "diagnosis_minutes": repair.diagnosis_minutes,
        "repair_minutes": repair.repair_minutes,
        "testing_minutes": repair.testing_minutes,
        "total_work_minutes": sum(
            value or 0
            for value in (
                repair.diagnosis_minutes,
                repair.repair_minutes,
                repair.testing_minutes,
            )
        ),
        "participant_total_minutes": sum(
            participant.minutes_worked or 0 for participant in repair.participants
        ),
        "source_return_transfer_id": repair.source_return_transfer_id,
        "source_return_document_id": repair.source_return_document_id,
        "source_return_batch_id": repair.source_return_batch_id,
        "opened_at": repair.opened_at.isoformat(),
        "started_at": repair.started_at.isoformat() if repair.started_at else None,
        "closed_at": repair.closed_at.isoformat() if repair.closed_at else None,
        "events": [
            {
                "id": event.id,
                "event_type": event.event_type,
                "description": event.description,
                "created_at": event.created_at.isoformat(),
            }
            for event in repair.events
        ],
        "parts_used": [
            {
                "id": part.id,
                "part_number": part.part_number,
                "description": part.description,
                "quantity": part.quantity,
                "unit": part.unit,
                "source": part.source,
            }
            for part in repair.parts_used
        ],
        "participants": [
            {
                "id": participant.id,
                "user_id": participant.user_id,
                "full_name": participant.full_name_snapshot,
                "job_title": participant.job_title_snapshot,
                "contribution": participant.contribution,
                "minutes_worked": participant.minutes_worked,
            }
            for participant in repair.participants
        ],
        "attachment_ids": [attachment.id for attachment in repair.attachments],
        "responsible_user": {
            "user_id": repair.responsible_user.id if repair.responsible_user else None,
            "display_name": (
                repair.responsible_user.full_name if repair.responsible_user else None
            ),
            "job_title": (
                repair.responsible_user.job_title if repair.responsible_user else None
            ),
        },
        "accepted_by": {
            "user_id": repair.accepted_by.id if repair.accepted_by else None,
            "display_name": repair.accepted_by.full_name if repair.accepted_by else None,
            "job_title": repair.accepted_by.job_title if repair.accepted_by else None,
        },
        "approved_by": {
            "user_id": repair.approved_by.id if repair.approved_by else None,
            "display_name": repair.approved_by.full_name if repair.approved_by else None,
            "job_title": repair.approved_by.job_title if repair.approved_by else None,
        },
    }
    machine = repair.machine
    date_value = repair.closed_at or repair.opened_at
    source_document = (
        db.get(OfficialDocument, repair.source_return_document_id)
        if repair.source_return_document_id
        else None
    )
    source_batch = (
        db.get(TransferBatch, repair.source_return_batch_id)
        if repair.source_return_batch_id
        else None
    )
    source_reference = " · ".join(
        value
        for value in (
            source_document.document_number if source_document else None,
            source_batch.batch_reference if source_batch else None,
        )
        if value
    )
    participant_names = "; ".join(
        " — ".join(
            value
            for value in (
                participant.full_name_snapshot,
                participant.job_title_snapshot,
                participant.contribution,
                _repair_duration(participant.minutes_worked, language),
            )
            if value
        )
        for participant in repair.participants
    )
    total_work_minutes = sum(
        value or 0
        for value in (
            repair.diagnosis_minutes,
            repair.repair_minutes,
            repair.testing_minutes,
        )
    )
    participant_total_minutes = sum(
        participant.minutes_worked or 0 for participant in repair.participants
    )
    values: dict[str, object] = {
        "DOCUMENT_NUMBER": number,
        "CREATION_DATE": date_value.strftime("%d.%m.%Y"),
        "MACHINE_NAME": machine.name,
        "MACHINE_NUMBER": machine.inventory_number,
        "BRAND": machine.brand,
        "MODEL": machine.model or "",
        "SERIAL_NUMBER": machine.serial_number or "",
        "PRESSURE_BAR": machine.pressure_bar,
        "BATCH_REFERENCE": "",
        "REPAIR_REFERENCE": repair.repair_reference or base,
        "SOURCE_RETURN_REFERENCE": source_reference,
        "ACCEPTANCE_DATE": repair.opened_at.strftime("%d.%m.%Y"),
        "COMPLETION_DATE": date_value.strftime("%d.%m.%Y"),
        "OWNERSHIP": machine.ownership or "",
        "REPORTED_PROBLEM": repair.reported_problem,
        "CONDITION_BEFORE": repair.condition_before or "",
        "REQUIRED_WORK": repair.required_work or "",
        "REQUIRED_PARTS": repair.required_parts_text or "",
        "REMOVED_PARTS": repair.removed_parts_text or "",
        "DIAGNOSTIC_CLEANING": repair.diagnostic_cleaning or "",
        "DIAGNOSIS": repair.diagnosis or "",
        "DIAGNOSIS_DURATION": _repair_duration(repair.diagnosis_minutes, language),
        "WORK_PERFORMED": repair.work_performed or "",
        "REPAIR_DURATION": _repair_duration(repair.repair_minutes, language),
        "TESTING_DURATION": _repair_duration(repair.testing_minutes, language),
        "TOTAL_WORK_DURATION": _repair_duration(total_work_minutes, language),
        "PARTICIPANT_TOTAL_DURATION": _repair_duration(
            participant_total_minutes, language
        ),
        "TEST_RESULT": _repair_test_summary(repair, language),
        "CONDITION_AFTER": repair.condition_after or repair.result or "",
        "FINAL_RESULT": repair.result or "",
        "REPAIR_START": repair.started_at.strftime("%d.%m.%Y %H:%M") if repair.started_at else "",
        "REPAIR_END": repair.closed_at.strftime("%d.%m.%Y %H:%M") if repair.closed_at else "",
        "HANDED_OVER_NAME": repair.reported_by_name or "",
        "ACCEPTED_BY_NAME": repair.accepted_by.full_name if repair.accepted_by else "",
        "REPAIRER_NAMES": participant_names,
        "APPROVED_BY_NAME": repair.approved_by.full_name if repair.approved_by else "",
        "APPROVED_BY_JOB_TITLE": repair.approved_by.job_title if repair.approved_by else "",
        "REPAIR_STATUS": repair.status,
        "LEFT_SIGNER_NAME": repair.responsible_user.full_name if repair.responsible_user else "",
        "LEFT_SIGNER_JOB_TITLE": repair.responsible_user.job_title if repair.responsible_user else "",
        "RIGHT_SIGNER_NAME": "",
        "RIGHT_SIGNER_JOB_TITLE": "",
        "LEFT_SIGNATURE": "",
        "RIGHT_SIGNATURE": "",
        "SIGNATURE_STATUS": _signature_status(language, finalized_internal=True),
    }
    values.update(_preparer_values(db, created_by_id))
    event_type_labels = {
        "ACCEPTED": "Създаване",
        "STATUS_CHANGE": "Промяна на статус",
        "PARTS": "Използвани части",
        "DIAGNOSIS": "Диагностика",
        "WORK": "Извършена работа",
        "TEST": "Тест",
        "COMPLETED": "Приключване",
    }
    event_rows = [["Дата", "Тип", "Описание"]] + [
        [
            event.created_at.strftime("%d.%m.%Y %H:%M"),
            event_type_labels.get(event.event_type, event.event_type.replace("_", " ").title()),
            event.description,
        ]
        for event in repair.events
    ]
    part_rows = [["Поз.", "Номер", "Описание", "Количество"]] + [
        [str(index), part.part_number or "", part.description, f"{part.quantity:g} {part.unit or ''}".strip()]
        for index, part in enumerate(repair.parts_used, start=1)
    ]
    participant_rows = [["Три имена", "Длъжност", "Участие", "Време"]] + (
        [
            [
                participant.full_name_snapshot,
                participant.job_title_snapshot or "",
                participant.contribution or "",
                _repair_duration(participant.minutes_worked, language),
            ]
            for participant in repair.participants
        ]
        or [["Няма допълнителни участници", "", "", ""]]
    )
    attachment_rows = [["Файл", "Етап", "Описание"]] + (
        [
            [attachment.filename, attachment.stage, attachment.caption or ""]
            for attachment in repair.attachments
        ]
        or [["Няма приложения", "", ""]]
    )
    docx = render_docx(
        template,
        values,
        {
            "REPAIR_EVENTS": event_rows,
            "PARTS_USED": part_rows,
            "REPAIR_PARTICIPANTS": participant_rows,
            "REPAIR_ATTACHMENTS": attachment_rows,
        },
    )
    pdf = convert_docx_to_pdf(docx) or build_repair_protocol_pdf(
        repair, language, source_reference
    )
    _register_official_version(
        db,
        number=number,
        document_type=DocumentType.REPAIR_PROTOCOL.value,
        language=language,
        docx=docx,
        pdf=pdf,
        snapshot=snapshot,
        created_by_id=created_by_id,
        machine_id=repair.machine_id,
        template_version_id=template.id,
        initial_status=OfficialDocumentStatus.FINALIZED.value,
    )
    return _generated_documents(
        number=number,
        document_type=DocumentType.REPAIR_PROTOCOL.value,
        language=language,
        template_version=template,
        docx=docx,
        pdf=pdf,
        snapshot=snapshot,
        created_by_id=created_by_id,
        machine_id=repair.machine_id,
        repair_id=repair.id,
    )


def make_repair_correction(
    db: Session,
    repair: Repair,
    created_by_id: int,
    reason: str,
    language: str = "bg",
) -> tuple[list[GeneratedDocument], OfficialDocument, OfficialDocumentVersion]:
    """Create a locked new version without rewriting the completed repair protocol."""
    existing_numbers = list(
        db.scalars(
            select(GeneratedDocument.document_number)
            .where(
                GeneratedDocument.repair_id == repair.id,
                GeneratedDocument.language == _language(language),
            )
            .order_by(GeneratedDocument.id)
        )
    )
    original = db.scalar(
        select(OfficialDocument)
        .where(
            OfficialDocument.document_type == DocumentType.REPAIR_PROTOCOL.value,
            OfficialDocument.machine_id == repair.machine_id,
            OfficialDocument.document_number.in_(existing_numbers or [""]),
        )
        .order_by(OfficialDocument.id)
    )
    if original is None or original.current_version_id is None:
        raise TemplateValidationError("Липсва заключена начална версия на ремонтния протокол.")
    previous = require_current_version(db, original)
    if previous.status not in {
        OfficialDocumentStatus.FINALIZED.value,
        OfficialDocumentStatus.SIGNED.value,
    }:
        raise TemplateValidationError("Само окончателен ремонтен протокол може да бъде коригиран.")

    documents = make_repair_documents(db, repair, created_by_id, language)
    temporary = db.scalar(
        select(OfficialDocument).where(
            OfficialDocument.document_number == documents[0].document_number
        )
    )
    if temporary is None or temporary.current_version_id is None:
        raise TemplateValidationError("Новата версия на ремонтния протокол не беше регистрирана.")
    version = require_current_version(db, temporary)
    next_version = previous.version + 1
    snapshot = dict(version.snapshot)
    snapshot["correction"] = {
        "reason": reason,
        "supersedes_version": previous.version,
        "created_by_id": created_by_id,
        "created_at": utcnow().isoformat(timespec="seconds") + "Z",
    }
    snapshot_sha256 = hashlib.sha256(
        json.dumps(
            snapshot,
            ensure_ascii=False,
            sort_keys=True,
            separators=(",", ":"),
            default=str,
        ).encode("utf-8")
    ).hexdigest()
    signing_sha256 = hashlib.sha256(
        json.dumps(
            {
                "document_number": original.document_number,
                "document_type": original.document_type,
                "snapshot_sha256": snapshot_sha256,
                "docx_sha256": version.docx_sha256,
                "pdf_sha256": version.pdf_sha256,
                "version": next_version,
            },
            sort_keys=True,
            separators=(",", ":"),
        ).encode("utf-8")
    ).hexdigest()
    version.version = next_version
    version.snapshot = snapshot
    version.snapshot_sha256 = snapshot_sha256
    version.signing_sha256 = signing_sha256
    version.correction_reason = reason
    version.supersedes_version_id = previous.id
    version.status = OfficialDocumentStatus.FINALIZED.value
    version.finalized_at = utcnow()
    previous.status = OfficialDocumentStatus.SUPERSEDED.value
    move_current_version(
        db,
        source_document=temporary,
        target_document=original,
        version=version,
    )
    db.delete(temporary)
    for document in documents:
        document.snapshot = {
            **document.snapshot,
            "official_document_id": original.id,
            "official_document_version": next_version,
            "correction_reason": reason,
        }
    db.flush()
    return documents, original, version


def make_part_request_documents(
    db: Session, request: PartRequest, created_by_id: int, language: str = "bg"
) -> list[GeneratedDocument]:
    template = _template_version(db, DocumentType.PART_REQUEST.value, language)
    number = request.request_reference or f"PR-{request.id:06d}"
    machine = request.machine
    values: dict[str, object] = {
        "DOCUMENT_NUMBER": number,
        "CREATION_DATE": request.created_at.strftime("%d.%m.%Y"),
        "MACHINE_NAME": machine.name if machine else "",
        "MACHINE_NUMBER": machine.inventory_number if machine else "",
        "BRAND": machine.brand if machine else "",
        "MODEL": machine.model or "" if machine else "",
        "SERIAL_NUMBER": machine.serial_number or "" if machine else "",
        "PRESSURE_BAR": machine.pressure_bar if machine else "",
        "BATCH_REFERENCE": "",
        "REMARKS": request.reason or "",
        "DECISION": request.decision_note or request.status,
        "LEFT_SIGNER_NAME": request.requested_by.full_name if request.requested_by else "",
        "LEFT_SIGNER_JOB_TITLE": request.requested_by.job_title if request.requested_by else "",
        "RIGHT_SIGNER_NAME": request.decided_by.full_name if request.decided_by else "",
        "RIGHT_SIGNER_JOB_TITLE": request.decided_by.job_title if request.decided_by else "",
        "LEFT_SIGNATURE": "",
        "RIGHT_SIGNATURE": "",
        "SIGNATURE_STATUS": _signature_status(language),
    }
    values.update(_preparer_values(db, created_by_id))
    line_rows = [["Поз.", "PART №", "Описание", "Количество", "Източник"]] + [
        [line.position or "", line.part_number or "", _part_request_line_description(line, language), f"{line.quantity:g} {line.unit or ''}".strip(), f"{line.source_document or ''} / {line.source_page or ''}".strip(" /")]
        for line in request.lines
    ]
    docx = render_docx(template, values, {"REQUEST_LINES": line_rows})
    pdf = convert_docx_to_pdf(docx) or build_part_request_pdf(request, language)
    snapshot = _request_snapshot(request)
    _register_official_version(
        db,
        number=number,
        document_type=DocumentType.PART_REQUEST.value,
        language=language,
        docx=docx,
        pdf=pdf,
        snapshot=snapshot,
        created_by_id=created_by_id,
        machine_id=request.machine_id,
        template_version_id=template.id,
    )
    return _generated_documents(
        number=number,
        document_type=DocumentType.PART_REQUEST.value,
        language=language,
        template_version=template,
        docx=docx,
        pdf=pdf,
        snapshot=snapshot,
        created_by_id=created_by_id,
        machine_id=request.machine_id,
        part_request_id=request.id,
    )


def build_daily_report_pdf(repairs: list[Repair]) -> bytes:
    _, _, body, label, title, _ = _pdf_styles()
    output = io.BytesIO()
    pdf = SimpleDocTemplate(output, pagesize=A4, leftMargin=10 * mm, rightMargin=8 * mm, topMargin=7 * mm, bottomMargin=8 * mm, title="AssetCore - дневен HPWJ отчет", author="AssetCore")
    story = [_pdf_header(), Spacer(1, 3 * mm), Paragraph("AssetCore - дневен HPWJ отчет", title), Paragraph(datetime.now().strftime("%d.%m.%Y %H:%M"), body), Spacer(1, 3 * mm)]
    if not repairs:
        story.append(Paragraph("Няма регистрирани ремонти.", body))
    else:
        data = [[Paragraph("Машина", label), Paragraph("Проблем", label), Paragraph("Статус", label)]]
        data.extend([[Paragraph(escape(repair.machine.name), body), Paragraph(escape(repair.reported_problem), body), Paragraph(escape(repair.status), body)] for repair in repairs])
        table = Table(data, colWidths=[38 * mm, 112 * mm, 42 * mm], repeatRows=1)
        table.setStyle(_pdf_table_style(header_rows=1))
        story.append(table)
    pdf.build(story)
    return output.getvalue()
