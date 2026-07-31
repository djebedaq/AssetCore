from __future__ import annotations
import io, json
from datetime import datetime
from pathlib import Path
import qrcode
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Mm, Pt
from fastapi import Depends, FastAPI, HTTPException, Response
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse
from fastapi.staticfiles import StaticFiles
from reportlab.lib.pagesizes import A4
from reportlab.pdfgen import canvas
from sqlalchemy import func, or_, select
from sqlalchemy.orm import Session, joinedload
from .database import Base, SessionLocal, engine, get_db
from .models import AuditLog, Location, Machine, PartCatalog, PartRequest, Repair, TechnicalDocument, TransferProtocol, User
from .schemas import *
from .security import create_access_token, get_current_user, verify_password
from .seed import seed_database
from .settings import settings

ROOT=Path(__file__).resolve().parents[1]
RESOURCES=ROOT/'resources'
DOCS_DIR=RESOURCES/'technical_docs'
app=FastAPI(title='AssetCore API',version='1.0.0-rc1-director-edition')
app.add_middleware(CORSMiddleware,allow_origins=[settings.frontend_origin,'http://localhost:4173'],allow_credentials=True,allow_methods=['*'],allow_headers=['*'])

@app.on_event('startup')
def startup():
    Base.metadata.create_all(engine)
    with SessionLocal() as db: seed_database(db)

@app.get('/api/health')
def health(): return {'status':'ok','service':'AssetCore','version':'1.0.0-rc1-director-edition'}

@app.post('/api/auth/login',response_model=TokenResponse)
def login(data:LoginRequest,db:Session=Depends(get_db)):
    user=db.scalar(select(User).where(User.email==data.email))
    if not user or not verify_password(data.password,user.password_hash): raise HTTPException(401,'Грешен имейл или парола')
    return TokenResponse(access_token=create_access_token(user),user={'id':user.id,'email':user.email,'full_name':user.full_name,'role':user.role})

def log(db:Session,user:User,entity_type:str,entity_id:int|None,action:str,details:dict|str|None=None):
    db.add(AuditLog(entity_type=entity_type,entity_id=entity_id,action=action,details=json.dumps(details,ensure_ascii=False) if isinstance(details,dict) else details,user_name=user.full_name))

@app.get('/api/dashboard')
def dashboard(_:User=Depends(get_current_user),db:Session=Depends(get_db)):
    total=db.scalar(select(func.count(Machine.id))) or 0
    by_status=dict(db.execute(select(Machine.status,func.count(Machine.id)).group_by(Machine.status)).all())
    return {'total_machines':total,'ready':by_status.get('Готова',0),'in_repair':sum(by_status.get(x,0) for x in ['В ремонт','За преглед','Чака части','Тестване']),'in_use':by_status.get('В употреба',0)+by_status.get('Издадена',0),'open_repairs':db.scalar(select(func.count(Repair.id)).where(Repair.closed_at.is_(None))) or 0,'pending_parts':db.scalar(select(func.count(PartRequest.id)).where(PartRequest.status.not_in(['Доставена','Завършена']))) or 0,'protocols':db.scalar(select(func.count(TransferProtocol.id))) or 0,'documents':db.scalar(select(func.count(TechnicalDocument.id))) or 0,'status_breakdown':by_status,'recent_repairs':[{'id':r.id,'machine':r.machine.name,'problem':r.reported_problem,'status':r.status,'opened_at':r.opened_at} for r in db.scalars(select(Repair).options(joinedload(Repair.machine)).order_by(Repair.opened_at.desc()).limit(5)).all()]}

@app.get('/api/locations')
def locations(_:User=Depends(get_current_user),db:Session=Depends(get_db)): return db.scalars(select(Location).order_by(Location.name)).all()
@app.get('/api/machines',response_model=list[MachineOut])
def machines(_:User=Depends(get_current_user),db:Session=Depends(get_db)): return db.scalars(select(Machine).options(joinedload(Machine.location)).order_by(Machine.pressure_bar.desc(),Machine.inventory_number)).all()
@app.get('/api/machines/{machine_id}',response_model=MachineOut)
def machine(machine_id:int,_:User=Depends(get_current_user),db:Session=Depends(get_db)):
    m=db.scalar(select(Machine).options(joinedload(Machine.location)).where(Machine.id==machine_id))
    if not m: raise HTTPException(404,'Машината не е намерена')
    return m
@app.post('/api/machines',response_model=MachineOut)
def create_machine(data:MachineCreate,user:User=Depends(get_current_user),db:Session=Depends(get_db)):
    if db.scalar(select(Machine).where(Machine.inventory_number==data.inventory_number)): raise HTTPException(409,'Дублиран инвентарен номер')
    m=Machine(**data.model_dump()); db.add(m); db.flush(); log(db,user,'machine',m.id,'Създадена машина',data.model_dump()); db.commit(); db.refresh(m); return m
@app.patch('/api/machines/{machine_id}',response_model=MachineOut)
def update_machine(machine_id:int,data:MachineUpdate,user:User=Depends(get_current_user),db:Session=Depends(get_db)):
    m=db.get(Machine,machine_id)
    if not m: raise HTTPException(404,'Машината не е намерена')
    before={'status':m.status,'location_id':m.location_id}
    for k,v in data.model_dump(exclude_unset=True).items(): setattr(m,k,v)
    m.updated_at=datetime.utcnow(); log(db,user,'machine',m.id,'Актуализирана машина',{'преди':before,'след':data.model_dump(exclude_unset=True)}); db.commit(); db.refresh(m); return m
@app.get('/api/machines/{machine_id}/qr')
def qr(machine_id:int,db:Session=Depends(get_db)):
    m=db.get(Machine,machine_id)
    if not m: raise HTTPException(404,'Машината не е намерена')
    img=qrcode.make(f'assetcore://machine/{m.id}'); out=io.BytesIO(); img.save(out,format='PNG'); return Response(out.getvalue(),media_type='image/png')

@app.get('/api/repairs',response_model=list[RepairOut])
def repairs(_:User=Depends(get_current_user),db:Session=Depends(get_db)): return db.scalars(select(Repair).options(joinedload(Repair.machine).joinedload(Machine.location)).order_by(Repair.opened_at.desc())).unique().all()
@app.post('/api/repairs',response_model=RepairOut)
def create_repair(data:RepairCreate,user:User=Depends(get_current_user),db:Session=Depends(get_db)):
    m=db.get(Machine,data.machine_id)
    if not m: raise HTTPException(404,'Машината не е намерена')
    r=Repair(**data.model_dump()); m.status='В ремонт'; db.add(r); db.flush(); log(db,user,'repair',r.id,'Приета машина за ремонт',{'machine':m.inventory_number,'problem':data.reported_problem}); db.commit(); return db.scalar(select(Repair).options(joinedload(Repair.machine).joinedload(Machine.location)).where(Repair.id==r.id))
@app.patch('/api/repairs/{repair_id}',response_model=RepairOut)
def update_repair(repair_id:int,data:RepairUpdate,user:User=Depends(get_current_user),db:Session=Depends(get_db)):
    r=db.get(Repair,repair_id)
    if not r: raise HTTPException(404,'Ремонтът не е намерен')
    for k,v in data.model_dump(exclude={'close'},exclude_unset=True).items(): setattr(r,k,v)
    if data.close: r.closed_at=datetime.utcnow(); r.status='Завършена'; r.machine.status='Готова'
    log(db,user,'repair',r.id,'Актуализиран ремонт',data.model_dump()); db.commit(); return db.scalar(select(Repair).options(joinedload(Repair.machine).joinedload(Machine.location)).where(Repair.id==r.id))

@app.get('/api/parts',response_model=list[PartRequestOut])
def parts(_:User=Depends(get_current_user),db:Session=Depends(get_db)): return db.scalars(select(PartRequest).options(joinedload(PartRequest.machine).joinedload(Machine.location)).order_by(PartRequest.created_at.desc())).all()
@app.post('/api/parts',response_model=PartRequestOut)
def create_part(data:PartRequestCreate,user:User=Depends(get_current_user),db:Session=Depends(get_db)):
    p=PartRequest(**data.model_dump()); db.add(p); db.flush(); log(db,user,'part_request',p.id,'Създадена заявка за части',data.model_dump()); db.commit(); return db.scalar(select(PartRequest).options(joinedload(PartRequest.machine).joinedload(Machine.location)).where(PartRequest.id==p.id))
@app.get('/api/catalog/parts',response_model=list[PartCatalogOut])
def catalog(q:str='',brand:str='',_:User=Depends(get_current_user),db:Session=Depends(get_db)):
    stmt=select(PartCatalog)
    if brand: stmt=stmt.where(PartCatalog.brand==brand)
    if q: stmt=stmt.where(or_(PartCatalog.part_number.ilike(f'%{q}%'),PartCatalog.description.ilike(f'%{q}%'),PartCatalog.assembly.ilike(f'%{q}%')))
    return db.scalars(stmt.order_by(PartCatalog.brand,PartCatalog.assembly,PartCatalog.position).limit(500)).all()

@app.get('/api/transfers',response_model=list[TransferOut])
def transfers(_:User=Depends(get_current_user),db:Session=Depends(get_db)): return db.scalars(select(TransferProtocol).options(joinedload(TransferProtocol.machine).joinedload(Machine.location)).order_by(TransferProtocol.created_at.desc())).unique().all()
@app.post('/api/transfers',response_model=TransferOut)
def create_transfer(data:TransferCreate,user:User=Depends(get_current_user),db:Session=Depends(get_db)):
    m=db.get(Machine,data.machine_id)
    if not m: raise HTTPException(404,'Машината не е намерена')
    number=f"HPWJ-{datetime.now():%Y%m%d}-{(db.scalar(select(func.count(TransferProtocol.id))) or 0)+1:04d}"
    t=TransferProtocol(**data.model_dump(),protocol_number=number); db.add(t); db.flush()
    if data.protocol_type=='Предаване': m.status='Издадена'
    elif data.protocol_type in ['Приемане','Връщане']: m.status='Върната'
    log(db,user,'transfer',t.id,f"Протокол: {data.protocol_type}",{'number':number,'machine':m.inventory_number}); db.commit(); return db.scalar(select(TransferProtocol).options(joinedload(TransferProtocol.machine).joinedload(Machine.location)).where(TransferProtocol.id==t.id))

def build_protocol_docx(t:TransferProtocol)->bytes:
    d=Document(); sec=d.sections[0]; sec.top_margin=Mm(16); sec.bottom_margin=Mm(16); sec.left_margin=Mm(18); sec.right_margin=Mm(18)
    p=d.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; r=p.add_run('ОДЕСОС ШИПРИПЕЪР ЯРД АД'); r.bold=True; r.font.size=Pt(14)
    p=d.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; r=p.add_run(f'ПРИЕМО-ПРЕДАВАТЕЛЕН ПРОТОКОЛ\n№ {t.protocol_number}'); r.bold=True; r.font.size=Pt(13)
    d.add_paragraph(f'Вид операция: {t.protocol_type}')
    table=d.add_table(rows=0,cols=2); table.style='Table Grid'
    rows=[('Дата',t.created_at.strftime('%d.%m.%Y %H:%M')),('Машина',f'{t.machine.name} — {t.machine.brand} {t.machine.model or ""}'),('Инвентарен №',t.machine.inventory_number),('Сериен №',t.machine.serial_number or '—'),('Налягане',f'{t.machine.pressure_bar} bar'),('Фирма/звено',t.company_unit or '—'),('Кораб',t.vessel or '—'),('Местоположение',t.location_text or '—'),('Комплектовка',t.equipment or '—'),('Състояние',t.condition_text or '—'),('Забележки',t.remarks or '—')]
    for a,b in rows: cells=table.add_row().cells; cells[0].text=a; cells[1].text=b
    d.add_paragraph(); sig=d.add_table(rows=2,cols=2); sig.cell(0,0).text=f'Предал: {t.handed_over_by or ""}'; sig.cell(0,1).text=f'Приел: {t.accepted_by or ""}'; sig.cell(1,0).text='Подпис: __________________'; sig.cell(1,1).text='Подпис: __________________'
    out=io.BytesIO(); d.save(out); return out.getvalue()
@app.get('/api/transfers/{transfer_id}/docx')
def protocol_docx(transfer_id:int,_:User=Depends(get_current_user),db:Session=Depends(get_db)):
    t=db.scalar(select(TransferProtocol).options(joinedload(TransferProtocol.machine)).where(TransferProtocol.id==transfer_id))
    if not t: raise HTTPException(404,'Протоколът не е намерен')
    return Response(build_protocol_docx(t),media_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document',headers={'Content-Disposition':f'attachment; filename={t.protocol_number}.docx'})
@app.get('/api/transfers/{transfer_id}/pdf')
def protocol_pdf(transfer_id:int,_:User=Depends(get_current_user),db:Session=Depends(get_db)):
    t=db.scalar(select(TransferProtocol).options(joinedload(TransferProtocol.machine)).where(TransferProtocol.id==transfer_id))
    if not t: raise HTTPException(404,'Протоколът не е намерен')
    out=io.BytesIO(); pdf=canvas.Canvas(out,pagesize=A4); w,h=A4; pdf.setFont('Helvetica-Bold',14); pdf.drawCentredString(w/2,h-45,'ASSETCORE - HPWJ TRANSFER PROTOCOL'); pdf.setFont('Helvetica',10); y=h-80
    for line in [f'No: {t.protocol_number}',f'Type: {t.protocol_type}',f'Machine: {t.machine.name}',f'Brand/model: {t.machine.brand} {t.machine.model or ""}',f'Serial: {t.machine.serial_number or "-"}',f'Company/unit: {t.company_unit or "-"}',f'Vessel: {t.vessel or "-"}',f'Location: {t.location_text or "-"}',f'Equipment: {t.equipment or "-"}',f'Condition: {t.condition_text or "-"}',f'Remarks: {t.remarks or "-"}']:
        pdf.drawString(45,y,line[:105]); y-=22
    pdf.save(); return Response(out.getvalue(),media_type='application/pdf',headers={'Content-Disposition':f'attachment; filename={t.protocol_number}.pdf'})

@app.get('/api/documents',response_model=list[TechnicalDocumentOut])
def documents(_:User=Depends(get_current_user),db:Session=Depends(get_db)): return db.scalars(select(TechnicalDocument).order_by(TechnicalDocument.brand,TechnicalDocument.category,TechnicalDocument.title)).all()
@app.get('/api/documents/{doc_id}/download')
def download_doc(doc_id:int,_:User=Depends(get_current_user),db:Session=Depends(get_db)):
    doc=db.get(TechnicalDocument,doc_id)
    if not doc: raise HTTPException(404,'Документът не е намерен')
    path=DOCS_DIR/doc.file_path
    if not path.is_file(): raise HTTPException(404,'Файлът липсва')
    return FileResponse(path,filename=path.name)
@app.get('/api/audit')
def audit(_:User=Depends(get_current_user),db:Session=Depends(get_db)): return db.scalars(select(AuditLog).order_by(AuditLog.created_at.desc()).limit(200)).all()

@app.get('/api/reports/daily.pdf')
def daily_report(_:User=Depends(get_current_user),db:Session=Depends(get_db)):
    out=io.BytesIO(); pdf=canvas.Canvas(out,pagesize=A4); w,h=A4; pdf.setFont('Helvetica-Bold',16); pdf.drawString(40,h-50,'AssetCore - Daily HPWJ report'); pdf.setFont('Helvetica',10); pdf.drawString(40,h-70,datetime.now().strftime('%d.%m.%Y %H:%M')); y=h-100
    for r in db.scalars(select(Repair).options(joinedload(Repair.machine)).order_by(Repair.opened_at.desc()).limit(30)).all(): pdf.drawString(40,y,f'{r.machine.name}: {r.reported_problem} [{r.status}]'[:110]); y-=18
    pdf.save(); return Response(out.getvalue(),media_type='application/pdf',headers={'Content-Disposition':'attachment; filename=assetcore-daily-report.pdf'})

frontend_dist=Path(__file__).resolve().parents[2]/'frontend'/'dist'
if frontend_dist.exists():
    app.mount('/assets',StaticFiles(directory=frontend_dist/'assets'),name='assets')
    @app.get('/{full_path:path}')
    def spa(full_path:str):
        path=frontend_dist/full_path
        return FileResponse(path if path.is_file() else frontend_dist/'index.html')
