"""Shared DOCX typography and PDF presentation primitives; no domain persistence."""

from __future__ import annotations

import io
import zipfile
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Mm, Pt
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import mm
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.platypus import (
    Image as ReportLabImage,
)
from reportlab.platypus import (
    Paragraph,
    Table,
    TableStyle,
)

from .common import (
    ASSETS,
    REPAIR_REFERENCE,
)


def _set_run_font(run, size: float = 9, bold: bool = False) -> None:
    run.font.name = "Arial"
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Arial")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    run.font.size = Pt(size)
    run.bold = bold


def _keep_with_next(paragraph) -> None:
    paragraph.paragraph_format.keep_with_next = True


def _set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    repeat = OxmlElement("w:tblHeader")
    repeat.set(qn("w:val"), "true")
    tr_pr.append(repeat)


def _set_cell(cell, value: str, *, bold: bool = False, size: float = 9) -> None:
    cell.text = ""
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run(value)
    _set_run_font(run, size, bold)


def _clear_body(document: Document) -> None:
    body = document._element.body
    for element in list(body):
        if element.tag != qn("w:sectPr"):
            body.remove(element)


def _prepare_document(reference: Path, *, a4: bool = True) -> Document:
    document = Document(reference) if reference.is_file() else Document()
    _clear_body(document)
    section = document.sections[0]
    if a4:
        section.page_width = Mm(210)
        section.page_height = Mm(297)
    section.top_margin = Mm(8)
    section.bottom_margin = Mm(9)
    section.left_margin = Mm(10)
    section.right_margin = Mm(8)
    section.header_distance = Mm(2)
    section.footer_distance = Mm(5)
    normal = document.styles["Normal"]
    normal.font.name = "Arial"
    normal._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Arial")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    normal.font.size = Pt(9)
    return document


def _add_centered(document: Document, text: str, size: float, bold: bool = False):
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(2)
    run = paragraph.add_run(text)
    _set_run_font(run, size, bold)
    return paragraph


def _register_pdf_fonts() -> tuple[str, str]:
    normal_candidates = [
        Path("C:/Windows/Fonts/arial.ttf"),
        Path("C:/Windows/Fonts/dejavusans.ttf"),
        Path("/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf"),
    ]
    bold_candidates = [
        Path("C:/Windows/Fonts/arialbd.ttf"),
        Path("C:/Windows/Fonts/dejavusans-bold.ttf"),
        Path("/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf"),
    ]
    normal_path = next((path for path in normal_candidates if path.is_file()), None)
    bold_path = next((path for path in bold_candidates if path.is_file()), None)
    if normal_path is None or bold_path is None:
        raise RuntimeError(
            "Липсва Unicode шрифт за генериране на документи на кирилица."
        )
    if "AssetCoreSans" not in pdfmetrics.getRegisteredFontNames():
        pdfmetrics.registerFont(TTFont("AssetCoreSans", str(normal_path)))
    if "AssetCoreSans-Bold" not in pdfmetrics.getRegisteredFontNames():
        pdfmetrics.registerFont(TTFont("AssetCoreSans-Bold", str(bold_path)))
    return "AssetCoreSans", "AssetCoreSans-Bold"


def _rina_image() -> io.BytesIO | None:
    if not REPAIR_REFERENCE.is_file():
        return None
    with zipfile.ZipFile(REPAIR_REFERENCE) as archive:
        candidates = [
            name
            for name in archive.namelist()
            if name.lower().endswith((".jpeg", ".jpg")) and name.startswith("word/media/")
        ]
        if not candidates:
            return None
        return io.BytesIO(archive.read(candidates[-1]))


def _pdf_styles():
    normal_font, bold_font = _register_pdf_fonts()
    styles = getSampleStyleSheet()
    body = ParagraphStyle(
        "AssetCoreBody",
        parent=styles["BodyText"],
        fontName=normal_font,
        fontSize=8.2,
        leading=10,
        textColor=colors.black,
    )
    label = ParagraphStyle(
        "AssetCoreLabel", parent=body, fontName=bold_font, fontSize=8.2, leading=10
    )
    title = ParagraphStyle(
        "AssetCoreTitle",
        parent=body,
        fontName=bold_font,
        fontSize=10.5,
        leading=12,
        alignment=TA_CENTER,
    )
    small = ParagraphStyle(
        "AssetCoreSmall", parent=body, fontSize=7.2, leading=8.5
    )
    return normal_font, bold_font, body, label, title, small


def _pdf_header() -> Table:
    _, _, body, _, _, _ = _pdf_styles()
    krz = ASSETS / "krz_logo.png"
    odessos = ASSETS / "odessos_logo.png"
    rina = _rina_image()
    left = ReportLabImage(str(krz), width=18 * mm, height=21 * mm) if krz.is_file() else Paragraph("KRZ", body)
    center = ReportLabImage(str(odessos), width=91 * mm, height=21 * mm) if odessos.is_file() else Paragraph("ODESSOS SHIPREPAIR &amp; CONVERSION", body)
    right = ReportLabImage(rina, width=23 * mm, height=20 * mm) if rina else Paragraph("RINA / AQAP 2110", body)
    table = Table([[left, center, right]], colWidths=[27 * mm, 119 * mm, 35 * mm], rowHeights=[23 * mm])
    table.setStyle(TableStyle([("GRID", (0, 0), (-1, -1), 0.65, colors.black), ("VALIGN", (0, 0), (-1, -1), "MIDDLE"), ("ALIGN", (0, 0), (-1, -1), "CENTER"), ("LEFTPADDING", (0, 0), (-1, -1), 2), ("RIGHTPADDING", (0, 0), (-1, -1), 2), ("TOPPADDING", (0, 0), (-1, -1), 1), ("BOTTOMPADDING", (0, 0), (-1, -1), 1)]))
    return table


def _pdf_table_style(*, header_rows: int = 0) -> TableStyle:
    commands = [
        ("GRID", (0, 0), (-1, -1), 0.55, colors.black),
        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
        ("LEFTPADDING", (0, 0), (-1, -1), 3),
        ("RIGHTPADDING", (0, 0), (-1, -1), 3),
        ("TOPPADDING", (0, 0), (-1, -1), 2),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 2),
    ]
    if header_rows:
        commands.extend([("BACKGROUND", (0, 0), (-1, header_rows - 1), colors.HexColor("#eeeeee")), ("ALIGN", (0, 0), (-1, header_rows - 1), "CENTER")])
    return TableStyle(commands)


def _add_section_title(document: Document, title: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(4)
    paragraph.paragraph_format.space_after = Pt(2)
    _keep_with_next(paragraph)
    _set_run_font(paragraph.add_run(title), 9.5, True)
