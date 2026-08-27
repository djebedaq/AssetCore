"""Transfer issue/return snapshots, rendering and document construction."""

from __future__ import annotations

import hashlib
import io
from xml.sax.saxutils import escape

from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Mm, Pt
from reportlab.lib.pagesizes import A4
from reportlab.lib.units import mm
from reportlab.platypus import (
    Paragraph,
    SimpleDocTemplate,
    Spacer,
    Table,
)
from sqlalchemy.orm import Session

from ..models import (
    DocumentType,
    GeneratedDocument,
    ProtocolDocument,
    TransferBatch,
    TransferProtocol,
)
from ..template_engine import convert_docx_to_pdf, render_docx
from .common import (
    DOCX_MEDIA_TYPE,
    PDF_MEDIA_TYPE,
    REPAIR_REFERENCE,
    TEXT,
    _language,
    safe_filename,
)
from .registration import (
    _generated_documents,
    _next_generated_number,
    _register_official_version,
)
from .rendering import (
    _add_centered,
    _pdf_header,
    _pdf_styles,
    _pdf_table_style,
    _prepare_document,
    _set_cell,
    _set_repeat_table_header,
    _set_run_font,
)
from .templates import (
    _preparer_values,
    _signature_status,
    _template_version,
)

CHECKLIST_CODES = [
    "pump",
    "supply_hose",
    "hp_hose",
    "gun",
    "nozzle",
    "tips",
    "cable",
    "plug",
    "chassis",
    "body",
]


CHECKLIST = {
    "bg": [
        "Помпа",
        "Шланг захранващ",
        "Шланг изходящ ВН",
        "Пистолет",
        "Дюза метла / ротационна",
        "Накрайници",
        "Кабел - метри",
        "Куплунг / Еврокак",
        "Ходова част",
        "Корпус",
    ],
    "en": [
        "Pump",
        "Supply hose",
        "High-pressure outlet hose",
        "Gun",
        "Broom / rotary nozzle",
        "Nozzles",
        "Cable - metres",
        "Coupling / Euro coupling",
        "Running gear",
        "Body",
    ],
    "ru": [
        "Насос",
        "Подающий шланг",
        "Выходной шланг ВД",
        "Пистолет",
        "Щелевая / ротационная форсунка",
        "Наконечники",
        "Кабель - метры",
        "Муфта / Евросоединение",
        "Ходовая часть",
        "Корпус",
    ],
}


def _machine_model(transfer: TransferProtocol) -> str:
    machine = transfer.machine
    brand = (machine.brand or machine.manufacturer or "").strip()
    pressure = f"{machine.pressure_bar}bar" if machine.pressure_bar else ""
    family = " - ".join(value for value in (brand, pressure) if value)
    return "; ".join(value for value in (family, machine.model or "") if value)


def _usage_text(transfer: TransferProtocol) -> str:
    return "; ".join(
        value
        for value in (
            transfer.company_unit,
            transfer.department,
            transfer.vessel,
            transfer.dock,
            transfer.pier,
            transfer.work_area,
            transfer.location_text,
        )
        if value
    )


def _equipment_text(transfer: TransferProtocol, language: str) -> str:
    labels = {
        "bg": ("Оборудване", "Шлангове", "Дюзи", "Пистолети", "Принадлежности"),
        "en": ("Equipment", "Hoses", "Nozzles", "Guns", "Accessories"),
        "ru": ("Оборудование", "Шланги", "Сопла", "Пистолеты", "Принадлежности"),
    }[language]
    values = (
        transfer.equipment,
        transfer.hoses,
        transfer.nozzles,
        transfer.guns,
        transfer.accessories,
    )
    return "; ".join(
        f"{label}: {value}" for label, value in zip(labels, values, strict=True) if value
    )


def _return_findings(transfer: TransferProtocol, language: str) -> str:
    labels = {
        "bg": {
            "missing": "Липсващо оборудване",
            "damage": "Повреди",
            "contamination": "Замърсяване",
            "cleaning": "Необходимо почистване",
            "inspection": "Необходим преглед",
            "repair": "Необходим ремонт",
            "yes": "да",
            "no": "не",
        },
        "en": {
            "missing": "Missing equipment", "damage": "Damage",
            "contamination": "Contamination", "cleaning": "Cleaning required",
            "inspection": "Inspection required", "repair": "Repair required",
            "yes": "yes", "no": "no",
        },
        "ru": {
            "missing": "Недостающее оборудование", "damage": "Повреждения",
            "contamination": "Загрязнение", "cleaning": "Требуется очистка",
            "inspection": "Требуется осмотр", "repair": "Требуется ремонт",
            "yes": "да", "no": "нет",
        },
    }[language]
    values = [
        (labels["missing"], transfer.return_missing_equipment),
        (labels["damage"], transfer.return_damage),
        (labels["contamination"], transfer.return_contamination),
        (labels["cleaning"], labels["yes"] if transfer.return_cleaning_required else labels["no"]),
        (labels["inspection"], labels["yes"] if transfer.return_inspection_required else labels["no"]),
        (labels["repair"], labels["yes"] if transfer.return_repair_required else labels["no"]),
    ]
    return "; ".join(f"{label}: {value}" for label, value in values if value)


def _protocol_remarks(transfer: TransferProtocol, operation: str, language: str) -> str:
    if operation == "return":
        return "; ".join(
            value
            for value in (
                transfer.return_result_text,
                _return_findings(transfer, language),
                transfer.return_notes,
            )
            if value
        )
    return "; ".join(
        value
        for value in (_equipment_text(transfer, language), transfer.remarks)
        if value
    )


def _checklist_rows(transfer: TransferProtocol, operation: str, language: str) -> list[tuple[str, str]]:
    stored = transfer.return_checklist if operation == "return" else transfer.issue_checklist
    condition_labels = {
        "bg": {"GOOD": "Добро", "SATISFACTORY": "Задоволително", "REPAIR": "За ремонт", "FAULTY": "Неизправно", "MISSING": "Липсва", "NA": "Не се прилага"},
        "en": {"GOOD": "Good", "SATISFACTORY": "Satisfactory", "REPAIR": "For repair", "FAULTY": "Faulty", "MISSING": "Missing", "NA": "N/A"},
        "ru": {"GOOD": "Хорошее", "SATISFACTORY": "Удовлетворительное", "REPAIR": "В ремонт", "FAULTY": "Неисправно", "MISSING": "Отсутствует", "NA": "Не применяется"},
    }[language]
    if not stored:
        return [(component, "") for component in CHECKLIST[language]]
    label_by_code = dict(zip(CHECKLIST_CODES, CHECKLIST[language], strict=True))
    rows = []
    for item in stored:
        code = str(item.get("code") or "")
        label = label_by_code.get(code) or str(item.get("label") or code)
        value = condition_labels.get(str(item.get("condition") or ""), str(item.get("condition") or ""))
        if item.get("length_m") is not None:
            value += f" · {item['length_m']} m"
        if item.get("note"):
            value += f" · {item['note']}"
        rows.append((label, value))
    return rows


def _identity_rows(transfer: TransferProtocol, operation: str, language: str) -> list[tuple[str, str]]:
    t = TEXT[language]
    date_value = (
        (transfer.returned_at or transfer.return_requested_at)
        if operation == "return"
        else transfer.issued_at
    ) or transfer.created_at
    return [
        (t["date"], date_value.strftime("%d.%m.%Y")),
        (t["equipment"], transfer.machine.category or "HPWJ"),
        (t["model"], _machine_model(transfer)),
        (t["inventory"], f"№ {transfer.machine.inventory_number}"),
        (t["serial"], transfer.machine.serial_number or ""),
    ]


def _protocol_snapshot(
    transfer: TransferProtocol, batch_reference: str, operation: str, language: str
) -> dict:
    return {
        "operation": operation,
        "language": language,
        "protocol_number": transfer.protocol_number,
        "batch_reference": batch_reference,
        "machine_id": transfer.machine_id,
        "machine_number": transfer.machine.inventory_number,
        "machine_name": transfer.machine.name,
        "brand": transfer.machine.brand,
        "model": transfer.machine.model,
        "pressure_bar": transfer.machine.pressure_bar,
        "serial_number": transfer.machine.serial_number,
        "company_unit": transfer.company_unit,
        "department": transfer.department,
        "vessel": transfer.vessel,
        "dock": transfer.dock,
        "pier": transfer.pier,
        "work_area": transfer.work_area,
        "location_text": transfer.location_text,
        "equipment": transfer.equipment,
        "hoses": transfer.hoses,
        "nozzles": transfer.nozzles,
        "guns": transfer.guns,
        "accessories": transfer.accessories,
        "condition": (
            transfer.return_condition_text
            if operation == "return"
            else transfer.condition_text
        ),
        "result": transfer.return_result_text if operation == "return" else None,
        "missing_equipment": transfer.return_missing_equipment if operation == "return" else None,
        "damage": transfer.return_damage if operation == "return" else None,
        "contamination": transfer.return_contamination if operation == "return" else None,
        "cleaning_required": transfer.return_cleaning_required if operation == "return" else None,
        "inspection_required": transfer.return_inspection_required if operation == "return" else None,
        "repair_required": transfer.return_repair_required if operation == "return" else None,
        "remarks": transfer.return_notes if operation == "return" else transfer.remarks,
        "handed_over_by": (
            transfer.returned_by_name if operation == "return" else transfer.handed_over_by
        ),
        "handed_over_job_title": (
            transfer.returned_by_job_title
            if operation == "return"
            else transfer.handed_over_job_title
        ),
        "handed_over_organization": (
            transfer.returned_by_company
            if operation == "return"
            else transfer.handed_over_department
        ),
        "accepted_by": (
            transfer.return_accepted_by if operation == "return" else transfer.accepted_by
        ),
        "accepted_by_job_title": (
            transfer.return_accepted_job_title
            if operation == "return"
            else transfer.accepted_by_job_title
        ),
        "accepted_by_organization": (
            transfer.return_accepted_department
            if operation == "return"
            else transfer.accepted_by_company
        ),
        "document_date": (
            (
                (transfer.returned_at or transfer.return_requested_at)
                if operation == "return"
                else transfer.issued_at
            )
            or transfer.created_at
        ).isoformat(),
    }


def _build_protocol_docx(
    transfer: TransferProtocol,
    batch_reference: str,
    *,
    operation: str,
    language: str,
) -> bytes:
    language = _language(language)
    t = TEXT[language]
    document = _prepare_document(REPAIR_REFERENCE)
    number = transfer.protocol_number + ("-R" if operation == "return" else "")
    _add_centered(document, f'{t["protocol"]} № {number}', 10.5, True)
    _add_centered(
        document,
        t["return_title"] if operation == "return" else t["issue_title"],
        10.5,
        True,
    )

    identity = document.add_table(rows=0, cols=2)
    identity.style = "Table Grid"
    identity.alignment = WD_TABLE_ALIGNMENT.CENTER
    identity.autofit = False
    for label, value in _identity_rows(transfer, operation, language):
        cells = identity.add_row().cells
        cells[0].width = Mm(46)
        cells[1].width = Mm(146)
        _set_cell(cells[0], f"{label}:", bold=True, size=8.5)
        _set_cell(cells[1], value, size=8.5)

    document.add_paragraph().paragraph_format.space_after = Pt(0)
    checklist = document.add_table(rows=1, cols=3)
    checklist.style = "Table Grid"
    checklist.alignment = WD_TABLE_ALIGNMENT.CENTER
    checklist.autofit = False
    widths = (Mm(10), Mm(66), Mm(116))
    for index, value in enumerate((t["number"], t["element"], t["condition"])):
        checklist.rows[0].cells[index].width = widths[index]
        _set_cell(checklist.rows[0].cells[index], value, bold=True, size=8.5)
        checklist.rows[0].cells[index].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_repeat_table_header(checklist.rows[0])
    for index, label in enumerate(CHECKLIST[language], start=1):
        cells = checklist.add_row().cells
        for cell_index, width in enumerate(widths):
            cells[cell_index].width = width
        _set_cell(cells[0], f"{index}.", size=8.5)
        _set_cell(cells[1], label, size=8.5)
        _set_cell(cells[2], "", size=8.5)

    overall_condition = (
        transfer.return_condition_text if operation == "return" else transfer.condition_text
    ) or ""
    condition = document.add_paragraph()
    condition.paragraph_format.space_before = Pt(2)
    condition.paragraph_format.space_after = Pt(1)
    run = condition.add_run(f'{t["overall_condition"]}: ')
    _set_run_font(run, 8.5, True)
    _set_run_font(condition.add_run(overall_condition), 8.5)

    usage = document.add_paragraph()
    usage.paragraph_format.space_before = Pt(1)
    usage.paragraph_format.space_after = Pt(1)
    usage_label = t["usage_return"] if operation == "return" else t["usage_issue"]
    _set_run_font(usage.add_run(f"{usage_label}: "), 8.5, True)
    _set_run_font(usage.add_run(_usage_text(transfer)), 8.5)

    remarks_value = _protocol_remarks(transfer, operation, language)
    remarks = document.add_paragraph()
    remarks.paragraph_format.space_before = Pt(1)
    remarks.paragraph_format.space_after = Pt(2)
    _set_run_font(remarks.add_run(f'{t["remarks"]}: '), 8.5, True)
    _set_run_font(remarks.add_run(remarks_value or ""), 8.5)

    signatures = document.add_table(rows=2, cols=3)
    signatures.style = "Table Grid"
    signatures.alignment = WD_TABLE_ALIGNMENT.CENTER
    signatures.autofit = False
    handed_label = t["returned"] if operation == "return" else t["handed"]
    handed_name = (
        transfer.returned_by_name if operation == "return" else transfer.handed_over_by
    ) or ""
    accepted_name = (
        transfer.return_accepted_by if operation == "return" else transfer.accepted_by
    ) or ""
    for row_index, (label, name) in enumerate(
        ((handed_label, handed_name), (t["accepted"], accepted_name))
    ):
        cells = signatures.rows[row_index].cells
        cells[0].width = Mm(38)
        cells[1].width = Mm(118)
        cells[2].width = Mm(36)
        _set_cell(cells[0], label, bold=True, size=8)
        _set_cell(cells[1], f"{name}\n({t['name']})", size=8)
        cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        _set_cell(cells[2], f"\n({t['signature']})", size=8)
        cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    footer = document.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer.paragraph_format.space_before = Pt(2)
    _set_run_font(
        footer.add_run(f'{t["batch"]}: {batch_reference} · {number}'), 7.5
    )
    output = io.BytesIO()
    document.save(output)
    return output.getvalue()


def build_protocol_docx(
    transfer: TransferProtocol, batch_reference: str, language: str = "bg"
) -> bytes:
    return _build_protocol_docx(
        transfer, batch_reference, operation="issue", language=language
    )


def build_return_protocol_docx(
    transfer: TransferProtocol, batch_reference: str, language: str = "bg"
) -> bytes:
    return _build_protocol_docx(
        transfer, batch_reference, operation="return", language=language
    )


def _build_protocol_pdf(
    transfer: TransferProtocol,
    batch_reference: str,
    *,
    operation: str,
    language: str,
) -> bytes:
    language = _language(language)
    t = TEXT[language]
    _, _, body, label, title, small = _pdf_styles()
    output = io.BytesIO()
    number = transfer.protocol_number + ("-R" if operation == "return" else "")
    pdf = SimpleDocTemplate(output, pagesize=A4, leftMargin=10 * mm, rightMargin=8 * mm, topMargin=7 * mm, bottomMargin=8 * mm, title=number, author="AssetCore")
    story = [
        _pdf_header(),
        Spacer(1, 2 * mm),
        Paragraph(f'{escape(t["protocol"])} № {escape(number)}', title),
        Paragraph(escape(t["return_title"] if operation == "return" else t["issue_title"]), title),
        Spacer(1, 2 * mm),
    ]
    identity = Table([[Paragraph(escape(label_text) + ":", label), Paragraph(escape(value), body)] for label_text, value in _identity_rows(transfer, operation, language)], colWidths=[46 * mm, 146 * mm])
    identity.setStyle(_pdf_table_style())
    story.extend([identity, Spacer(1, 3 * mm)])
    checklist_data = [[Paragraph(escape(t["number"]), label), Paragraph(escape(t["element"]), label), Paragraph(escape(t["condition"]), label)]]
    checklist_data.extend([[Paragraph(f"{index}.", body), Paragraph(escape(component), body), Paragraph(escape(value), body)] for index, (component, value) in enumerate(_checklist_rows(transfer, operation, language), start=1)])
    checklist = Table(checklist_data, colWidths=[10 * mm, 66 * mm, 116 * mm], repeatRows=1)
    checklist.setStyle(_pdf_table_style(header_rows=1))
    condition_value = (transfer.return_condition_text if operation == "return" else transfer.condition_text) or ""
    usage_label = t["usage_return"] if operation == "return" else t["usage_issue"]
    remarks_value = _protocol_remarks(transfer, operation, language)
    story.extend([checklist, Spacer(1, 2 * mm), Paragraph(f'<b>{escape(t["overall_condition"])}:</b> {escape(condition_value)}', body), Paragraph(f"<b>{escape(usage_label)}:</b> {escape(_usage_text(transfer))}", body), Paragraph(f'<b>{escape(t["remarks"])}:</b> {escape(remarks_value or "")}', body), Spacer(1, 3 * mm)])
    handed_label = t["returned"] if operation == "return" else t["handed"]
    handed_name = (transfer.returned_by_name if operation == "return" else transfer.handed_over_by) or ""
    accepted_name = (transfer.return_accepted_by if operation == "return" else transfer.accepted_by) or ""
    signature_data = [[Paragraph(escape(handed_label), label), Paragraph(f'{escape(handed_name)}<br/><font size="7">({escape(t["name"])})</font>', body), Paragraph(f'<br/><font size="7">({escape(t["signature"])})</font>', small)], [Paragraph(escape(t["accepted"]), label), Paragraph(f'{escape(accepted_name)}<br/><font size="7">({escape(t["name"])})</font>', body), Paragraph(f'<br/><font size="7">({escape(t["signature"])})</font>', small)]]
    signatures = Table(signature_data, colWidths=[38 * mm, 118 * mm, 36 * mm], rowHeights=[18 * mm, 18 * mm])
    signatures.setStyle(_pdf_table_style())
    story.extend([signatures, Spacer(1, 1.5 * mm), Paragraph(f'{escape(t["batch"])}: {escape(batch_reference)} · {escape(number)}', small)])
    pdf.build(story)
    return output.getvalue()


def build_protocol_pdf(
    transfer: TransferProtocol, batch_reference: str, language: str = "bg"
) -> bytes:
    return _build_protocol_pdf(
        transfer, batch_reference, operation="issue", language=language
    )


def build_return_protocol_pdf(
    transfer: TransferProtocol, batch_reference: str, language: str = "bg"
) -> bytes:
    return _build_protocol_pdf(
        transfer, batch_reference, operation="return", language=language
    )


def _protocol_template_values(
    db: Session,
    transfer: TransferProtocol,
    batch_reference: str,
    operation: str,
    language: str,
    created_by_id: int,
) -> dict[str, object]:
    language = _language(language)
    t = TEXT[language]
    date_value = (
        transfer.returned_at if operation == "return" else transfer.issued_at
    ) or transfer.created_at
    left_name = transfer.returned_by_name if operation == "return" else transfer.handed_over_by
    right_name = transfer.return_accepted_by if operation == "return" else transfer.accepted_by
    left_job = (
        transfer.returned_by_job_title
        if operation == "return"
        else transfer.handed_over_job_title
    )
    right_job = (
        transfer.return_accepted_job_title
        if operation == "return"
        else transfer.accepted_by_job_title
    )
    if operation == "return":
        left_role = t["returned"]
        right_role = " · ".join(value for value in (t["accepted"], right_job) if value)
    else:
        left_role = " · ".join(value for value in (t["handed"], left_job) if value)
        right_role = t["accepted"]
    values: dict[str, object] = {
        "DOCUMENT_NUMBER": f"{transfer.protocol_number}-R" if operation == "return" else transfer.protocol_number,
        "CREATION_DATE": date_value.strftime("%d.%m.%Y"),
        "EQUIPMENT_TYPE": transfer.machine.category or "HPWJ",
        "MACHINE_NAME": transfer.machine.name,
        "MACHINE_NUMBER": transfer.machine.inventory_number,
        "BRAND": transfer.machine.brand,
        "MODEL": transfer.machine.model or "",
        "MODEL_DISPLAY": _machine_model(transfer),
        "SERIAL_NUMBER": transfer.machine.serial_number or "",
        "PRESSURE_BAR": transfer.machine.pressure_bar,
        "BATCH_REFERENCE": batch_reference,
        "CONDITION_LABEL": t["overall_condition"],
        "CONDITION_TEXT": transfer.return_condition_text if operation == "return" else transfer.condition_text,
        "USAGE_TEXT": _usage_text(transfer),
        "REMARKS": _protocol_remarks(transfer, operation, language),
        "LEFT_SIGNER_NAME": left_name or "",
        "LEFT_SIGNER_JOB_TITLE": left_job or "",
        "LEFT_SIGNER_ROLE": left_role or "",
        "RIGHT_SIGNER_NAME": right_name or "",
        "RIGHT_SIGNER_JOB_TITLE": right_job or "",
        "RIGHT_SIGNER_ROLE": right_role or "",
        "LEFT_SIGNATURE": "[[ASSETCORE_LEFT_SIGNATURE]]",
        "RIGHT_SIGNATURE": "[[ASSETCORE_RIGHT_SIGNATURE]]",
        "SIGNATURE_STATUS": _signature_status(language),
    }
    checklist_rows = _checklist_rows(transfer, operation, language)
    for index in range(1, 11):
        values[f"CHECK_{index}"] = (
            checklist_rows[index - 1][1] if index <= len(checklist_rows) else ""
        )
    values.update(_preparer_values(db, created_by_id))
    return values


def make_protocol_documents(
    db: Session,
    transfer: TransferProtocol,
    batch: TransferBatch,
    created_by_id: int,
    language: str = "bg",
) -> list[ProtocolDocument]:
    template = _template_version(db, DocumentType.TRANSFER_ISSUE.value, language)
    snapshot = _protocol_snapshot(
        transfer, batch.batch_reference, "issue", _language(language)
    )
    stem = safe_filename(transfer.protocol_number)
    docx = render_docx(
        template,
        _protocol_template_values(
            db, transfer, batch.batch_reference, "issue", language, created_by_id
        ),
    )
    pdf = convert_docx_to_pdf(docx) or build_protocol_pdf(
        transfer, batch.batch_reference, language
    )
    _register_official_version(
        db,
        number=transfer.protocol_number,
        document_type=DocumentType.TRANSFER_ISSUE.value,
        language=language,
        docx=docx,
        pdf=pdf,
        snapshot=snapshot,
        created_by_id=created_by_id,
        machine_id=transfer.machine_id,
        transfer_id=transfer.id,
        batch_id=batch.id,
        template_version_id=template.id,
    )
    generated = [
        ("docx", DOCX_MEDIA_TYPE, docx),
        ("pdf", PDF_MEDIA_TYPE, pdf),
    ]
    return [
        ProtocolDocument(
            transfer_id=transfer.id,
            machine_id=transfer.machine_id,
            batch_id=batch.id,
            format=format_name,
            filename=f"{stem}.{format_name}",
            media_type=media_type,
            content=content,
            sha256=hashlib.sha256(content).hexdigest(),
            document_number=transfer.protocol_number,
            language=_language(language),
            template_version_id=template.id if template else None,
            snapshot=snapshot,
            created_by_id=created_by_id,
            created_at=transfer.created_at,
        )
        for format_name, media_type, content in generated
    ]


def make_return_documents(
    db: Session,
    transfer: TransferProtocol,
    batch: TransferBatch | None,
    created_by_id: int,
    language: str = "bg",
) -> list[GeneratedDocument]:
    template = _template_version(db, DocumentType.TRANSFER_RETURN.value, language)
    batch_reference = batch.batch_reference if batch else "-"
    base = f"{transfer.protocol_number}-R"
    number = _next_generated_number(db, base)
    values = _protocol_template_values(
        db, transfer, batch_reference, "return", language, created_by_id
    )
    values["DOCUMENT_NUMBER"] = number
    docx = render_docx(template, values)
    pdf = convert_docx_to_pdf(docx) or build_return_protocol_pdf(
        transfer, batch_reference, language
    )
    snapshot = _protocol_snapshot(transfer, batch_reference, "return", _language(language))
    _register_official_version(
        db,
        number=number,
        document_type=DocumentType.TRANSFER_RETURN.value,
        language=language,
        docx=docx,
        pdf=pdf,
        snapshot=snapshot,
        created_by_id=created_by_id,
        machine_id=transfer.machine_id,
        transfer_id=transfer.id,
        batch_id=batch.id if batch else None,
        template_version_id=template.id,
    )
    return _generated_documents(
        number=number,
        document_type=DocumentType.TRANSFER_RETURN.value,
        language=language,
        template_version=template,
        docx=docx,
        pdf=pdf,
        snapshot=snapshot,
        created_by_id=created_by_id,
        machine_id=transfer.machine_id,
        transfer_id=transfer.id,
        batch_id=batch.id if batch else None,
    )
