"""Canonical parts-request snapshots and protocol construction."""

from __future__ import annotations

import io
from xml.sax.saxutils import escape

from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Mm, Pt
from reportlab.lib.pagesizes import A4
from reportlab.lib.units import mm
from reportlab.platypus import (
    Paragraph,
    SimpleDocTemplate,
    Spacer,
    Table,
)
from sqlalchemy.orm import Session

from ..models import (
    DocumentType,
    GeneratedDocument,
    PartRequest,
    PartRequestLine,
)
from ..template_engine import convert_docx_to_pdf, render_docx
from .common import (
    PARTS_REFERENCE,
    TEXT,
    _language,
)
from .registration import (
    _generated_documents,
    _register_official_version,
)
from .rendering import (
    _add_centered,
    _pdf_header,
    _pdf_styles,
    _pdf_table_style,
    _prepare_document,
    _set_cell,
    _set_repeat_table_header,
    _set_run_font,
)
from .templates import (
    _preparer_values,
    _signature_status,
    _template_version,
)


def _request_snapshot(request: PartRequest) -> dict:
    return {
        "request_id": request.id,
        "request_reference": request.request_reference,
        "machine_id": request.machine_id,
        "machine_number": request.machine.inventory_number if request.machine else None,
        "repair_id": request.repair_id,
        "repair_reference": request.repair.repair_reference if request.repair else None,
        "priority": request.priority,
        "status": request.status,
        "language": request.language,
        "reason": request.reason,
        "department": request.department,
        "supplier": request.supplier,
        "delivery_note": request.delivery_note,
        "ordered_at": request.ordered_at.isoformat() if request.ordered_at else None,
        "delivered_at": request.delivered_at.isoformat() if request.delivered_at else None,
        "requested_by_id": request.requested_by_id,
        "submitted_at": request.submitted_at.isoformat() if request.submitted_at else None,
        "decided_at": request.decided_at.isoformat() if request.decided_at else None,
        "decision_note": request.decision_note,
        "lines": [
            {
                "catalog_part_id": line.catalog_part_id,
                "position": line.position,
                "part_number": line.part_number,
                "description": line.description,
                "quantity": line.quantity,
                "unit": line.unit,
                "source_document": line.source_document,
                "source_page": line.source_page,
                "delivered_quantity": line.delivered_quantity,
                "is_unknown_part": line.is_unknown_part,
                "assembly": line.assembly,
                "note": line.note,
                "linked_catalog_part_id": line.linked_catalog_part_id,
                "linked_part_number": line.linked_catalog_part.part_number if line.linked_catalog_part else None,
                "linked_at": line.linked_at.isoformat() if line.linked_at else None,
            }
            for line in request.lines
        ],
    }


def _part_request_line_description(line: PartRequestLine, language: str) -> str:
    if not line.is_unknown_part:
        return line.description
    label = {
        "bg": "Част без потвърден part number",
        "en": "Part without a confirmed part number",
        "ru": "Деталь без подтверждённого part number",
    }[_language(language)]
    assembly = {"bg": "Възел", "en": "Assembly", "ru": "Узел"}[_language(language)]
    linked = ""
    if line.linked_catalog_part:
        linked_label = {"bg": "Свързана с", "en": "Linked to", "ru": "Связана с"}[_language(language)]
        linked = f"; {linked_label}: {line.linked_catalog_part.part_number}"
    return f"[{label}] {assembly}: {line.assembly or '-'}; {line.description}{linked}"


def build_part_request_docx(request: PartRequest, language: str = "bg") -> bytes:
    language = _language(language)
    t = TEXT[language]
    document = _prepare_document(PARTS_REFERENCE)
    for line in t["part_request_title"].splitlines():
        _add_centered(document, line, 10.5, True)
    if request.machine:
        machine = request.machine
        for label_text, value in ((t["machine"], f"{machine.name}; {machine.brand} {machine.model or ''}".strip()), (t["inventory"], machine.inventory_number), (t["serial"], machine.serial_number or "")):
            paragraph = document.add_paragraph()
            paragraph.paragraph_format.space_after = Pt(0)
            _set_run_font(paragraph.add_run(f"{label_text}: "), 8.5, True)
            _set_run_font(paragraph.add_run(value), 8.5)
    document.add_paragraph()
    table = document.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = (t["position"], t["part_number"], t["description"], t["quantity"])
    widths = (Mm(14), Mm(31), Mm(127), Mm(20))
    for index, value in enumerate(headers):
        table.rows[0].cells[index].width = widths[index]
        _set_cell(table.rows[0].cells[index], value, bold=True, size=8)
        table.rows[0].cells[index].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_repeat_table_header(table.rows[0])
    for line in request.lines:
        cells = table.add_row().cells
        values = (line.position or "", line.part_number or "", _part_request_line_description(line, language), f"{line.quantity:g} {line.unit or ''}".strip())
        for index, value in enumerate(values):
            cells[index].width = widths[index]
            _set_cell(cells[index], value, size=8)
    if request.reason:
        paragraph = document.add_paragraph()
        _set_run_font(paragraph.add_run(f'{t["remarks"]}: '), 8.5, True)
        _set_run_font(paragraph.add_run(request.reason), 8.5)
    provenance = [f"{line.source_document}, p. {line.source_page}" for line in request.lines if line.source_document and line.source_page]
    if provenance:
        paragraph = document.add_paragraph()
        _set_run_font(paragraph.add_run(f'{t["source"]}: '), 8, True)
        _set_run_font(paragraph.add_run("; ".join(provenance)), 8)
    footer = document.add_table(rows=1, cols=3)
    footer.style = "Table Grid"
    request_ref = request.request_reference or f"PR-{request.id:06d}"
    requester = request.requested_by.full_name if request.requested_by else ""
    decision = request.decision_note or request.status
    for cell, value in zip(footer.rows[0].cells, (f'{t["request_number"]}: {request_ref}', f'{t["date"]}: {request.created_at:%d.%m.%Y}', f'{t["requester"]}: {requester}\n{t["decision"]}: {decision}'), strict=True):
        _set_cell(cell, value, size=8)
    output = io.BytesIO()
    document.save(output)
    return output.getvalue()


def build_part_request_pdf(request: PartRequest, language: str = "bg") -> bytes:
    language = _language(language)
    t = TEXT[language]
    _, _, body, label, title, small = _pdf_styles()
    output = io.BytesIO()
    request_ref = request.request_reference or f"PR-{request.id:06d}"
    pdf = SimpleDocTemplate(output, pagesize=A4, leftMargin=10 * mm, rightMargin=8 * mm, topMargin=7 * mm, bottomMargin=8 * mm, title=request_ref, author="AssetCore")
    story = [_pdf_header(), Spacer(1, 3 * mm)]
    story.extend(Paragraph(escape(line), title) for line in t["part_request_title"].splitlines())
    if request.machine:
        machine = request.machine
        for label_text, value in ((t["machine"], f"{machine.name}; {machine.brand} {machine.model or ''}".strip()), (t["inventory"], machine.inventory_number), (t["serial"], machine.serial_number or "")):
            story.append(Paragraph(f"<b>{escape(label_text)}:</b> {escape(value)}", body))
    story.append(Spacer(1, 3 * mm))
    data = [[Paragraph(escape(value), label) for value in (t["position"], t["part_number"], t["description"], t["quantity"])]]
    data.extend([[Paragraph(escape(line.position or ""), small), Paragraph(escape(line.part_number or ""), small), Paragraph(escape(_part_request_line_description(line, language)), small), Paragraph(escape(f"{line.quantity:g} {line.unit or ''}".strip()), small)] for line in request.lines])
    table = Table(data, colWidths=[14 * mm, 31 * mm, 127 * mm, 20 * mm], repeatRows=1)
    table.setStyle(_pdf_table_style(header_rows=1))
    story.append(table)
    if request.reason:
        story.extend([Spacer(1, 2 * mm), Paragraph(f'<b>{escape(t["remarks"])}:</b> {escape(request.reason)}', body)])
    provenance = [f"{line.source_document}, p. {line.source_page}" for line in request.lines if line.source_document and line.source_page]
    if provenance:
        story.append(Paragraph(f'<b>{escape(t["source"])}:</b> {escape("; ".join(provenance))}', small))
    requester = request.requested_by.full_name if request.requested_by else ""
    decision = request.decision_note or request.status
    footer = Table([[Paragraph(f'{escape(t["request_number"])}: {escape(request_ref)}', small), Paragraph(f'{escape(t["date"])}: {request.created_at:%d.%m.%Y}', small), Paragraph(f'{escape(t["requester"])}: {escape(requester)}<br/>{escape(t["decision"])}: {escape(decision)}', small)]], colWidths=[64 * mm, 52 * mm, 76 * mm])
    footer.setStyle(_pdf_table_style())
    story.extend([Spacer(1, 4 * mm), footer])
    pdf.build(story)
    return output.getvalue()


def make_part_request_documents(
    db: Session, request: PartRequest, created_by_id: int, language: str = "bg"
) -> list[GeneratedDocument]:
    template = _template_version(db, DocumentType.PART_REQUEST.value, language)
    number = request.request_reference or f"PR-{request.id:06d}"
    machine = request.machine
    values: dict[str, object] = {
        "DOCUMENT_NUMBER": number,
        "CREATION_DATE": request.created_at.strftime("%d.%m.%Y"),
        "MACHINE_NAME": machine.name if machine else "",
        "MACHINE_NUMBER": machine.inventory_number if machine else "",
        "BRAND": machine.brand if machine else "",
        "MODEL": machine.model or "" if machine else "",
        "SERIAL_NUMBER": machine.serial_number or "" if machine else "",
        "PRESSURE_BAR": machine.pressure_bar if machine else "",
        "BATCH_REFERENCE": "",
        "REMARKS": request.reason or "",
        "DECISION": request.decision_note or request.status,
        "LEFT_SIGNER_NAME": request.requested_by.full_name if request.requested_by else "",
        "LEFT_SIGNER_JOB_TITLE": request.requested_by.job_title if request.requested_by else "",
        "RIGHT_SIGNER_NAME": request.decided_by.full_name if request.decided_by else "",
        "RIGHT_SIGNER_JOB_TITLE": request.decided_by.job_title if request.decided_by else "",
        "LEFT_SIGNATURE": "",
        "RIGHT_SIGNATURE": "",
        "SIGNATURE_STATUS": _signature_status(language),
    }
    values.update(_preparer_values(db, created_by_id))
    line_rows = [["Поз.", "PART №", "Описание", "Количество", "Източник"]] + [
        [line.position or "", line.part_number or "", _part_request_line_description(line, language), f"{line.quantity:g} {line.unit or ''}".strip(), f"{line.source_document or ''} / {line.source_page or ''}".strip(" /")]
        for line in request.lines
    ]
    docx = render_docx(template, values, {"REQUEST_LINES": line_rows})
    pdf = convert_docx_to_pdf(docx) or build_part_request_pdf(request, language)
    snapshot = _request_snapshot(request)
    _register_official_version(
        db,
        number=number,
        document_type=DocumentType.PART_REQUEST.value,
        language=language,
        docx=docx,
        pdf=pdf,
        snapshot=snapshot,
        created_by_id=created_by_id,
        machine_id=request.machine_id,
        template_version_id=template.id,
    )
    return _generated_documents(
        number=number,
        document_type=DocumentType.PART_REQUEST.value,
        language=language,
        template_version=template,
        docx=docx,
        pdf=pdf,
        snapshot=snapshot,
        created_by_id=created_by_id,
        machine_id=request.machine_id,
        part_request_id=request.id,
    )
