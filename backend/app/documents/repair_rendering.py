"""Repair DOCX and fallback PDF layouts, including legacy compatibility builders."""

from __future__ import annotations

import io
from xml.sax.saxutils import escape

from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Mm, Pt
from reportlab.lib.pagesizes import A4
from reportlab.lib.units import mm
from reportlab.platypus import (
    Image as ReportLabImage,
)
from reportlab.platypus import (
    KeepTogether,
    PageBreak,
    Paragraph,
    SimpleDocTemplate,
    Spacer,
    Table,
)

from ..models import (
    Repair,
)
from .common import (
    REPAIR_REFERENCE,
    TEXT,
    _language,
)
from .rendering import (
    _add_centered,
    _add_section_title,
    _pdf_header,
    _pdf_styles,
    _pdf_table_style,
    _prepare_document,
    _set_cell,
    _set_repeat_table_header,
    _set_run_font,
)
from .templates import (
    _signature_status,
)


def build_repair_protocol_docx(repair: Repair, language: str = "bg") -> bytes:
    language = _language(language)
    t = TEXT[language]
    document = _prepare_document(REPAIR_REFERENCE)
    _add_centered(document, f'{t["repair_protocol"]} № {repair.repair_reference or repair.id}', 11, True)
    machine_table = document.add_table(rows=0, cols=2)
    machine_table.style = "Table Grid"
    values = [
        (t["machine"], f"{repair.machine.name}; {repair.machine.brand} {repair.machine.model or ''}".strip()),
        (t["inventory"], repair.machine.inventory_number),
        (t["serial"], repair.machine.serial_number or ""),
        (t["date"], repair.opened_at.strftime("%d.%m.%Y")),
    ]
    for label_text, value in values:
        cells = machine_table.add_row().cells
        _set_cell(cells[0], label_text, bold=True)
        _set_cell(cells[1], value)
    for title_key, value in (
        ("reported_problem", repair.reported_problem),
        ("symptoms", repair.symptoms),
        ("condition_before", repair.condition_before),
        ("diagnosis", repair.diagnosis),
        ("required_work", repair.required_work),
        ("removed_parts", repair.removed_parts_text),
        ("cleaning", repair.diagnostic_cleaning),
        ("work", repair.work_performed),
    ):
        _add_section_title(document, t[title_key])
        paragraph = document.add_paragraph(value or "")
        paragraph.paragraph_format.space_after = Pt(1)
    if repair.events:
        _add_section_title(document, t["events"])
        event_table = document.add_table(rows=1, cols=3)
        event_table.style = "Table Grid"
        for cell, value in zip(event_table.rows[0].cells, (t["date"], t["condition"], t["description"]), strict=True):
            _set_cell(cell, value, bold=True, size=8)
        _set_repeat_table_header(event_table.rows[0])
        for event in sorted(repair.events, key=lambda item: (item.created_at, item.id)):
            cells = event_table.add_row().cells
            _set_cell(cells[0], event.created_at.strftime("%d.%m.%Y %H:%M"), size=8)
            _set_cell(cells[1], event.event_type, size=8)
            _set_cell(cells[2], event.description or "", size=8)
    if repair.parts_used:
        _add_section_title(document, t["parts_used"])
        parts = document.add_table(rows=1, cols=4)
        parts.style = "Table Grid"
        for cell, value in zip(parts.rows[0].cells, (t["position"], t["part_number"], t["description"], t["quantity"]), strict=True):
            _set_cell(cell, value, bold=True, size=8)
        _set_repeat_table_header(parts.rows[0])
        for index, part in enumerate(repair.parts_used, start=1):
            cells = parts.add_row().cells
            for cell, value in zip(cells, (str(index), part.part_number or "", part.description, f"{part.quantity:g} {part.unit or ''}".strip()), strict=True):
                _set_cell(cell, value, size=8)
    if repair.participants:
        _add_section_title(document, t.get("participants", "Допълнителни участници"))
        participants = document.add_table(rows=1, cols=4)
        participants.style = "Table Grid"
        for cell, value in zip(participants.rows[0].cells, (t.get("full_name", "Три имена"), t.get("job_title", "Длъжност"), t.get("contribution", "Участие"), t.get("duration", "Време")), strict=True):
            _set_cell(cell, value, bold=True, size=8)
        for participant in repair.participants:
            cells = participants.add_row().cells
            for cell, value in zip(cells, (participant.full_name_snapshot, participant.job_title_snapshot or "", participant.contribution or "", _repair_duration(participant.minutes_worked, language)), strict=True):
                _set_cell(cell, value, size=8)
    _add_section_title(document, t["test"])
    test_text = "; ".join(
        value
        for value in (
            f"cleaning={repair.cleaning_completed_at.strftime('%d.%m.%Y %H:%M')}" if repair.cleaning_completed_at else None,
            f"test={'PASS' if repair.test_passed else 'FAIL'}" if repair.test_passed is not None else None,
            f'{t["test_method"]}: {repair.test_method}' if repair.test_method else None,
            f'{t["test_pressure"]}: {repair.test_pressure_bar} bar' if repair.test_pressure_bar is not None else None,
            f'{t["leaks"]}: {t["yes"] if repair.leaks_detected else t["no"]}' if repair.leaks_detected is not None else None,
            f'{t["electrical_test"]}: {repair.electrical_test_result}' if repair.electrical_test_result else None,
            f'{t["functional_test"]}: {repair.functional_test_result}' if repair.functional_test_result else None,
            repair.test_details,
        )
        if value
    )
    document.add_paragraph(test_text)
    _add_section_title(document, t["condition_after"])
    document.add_paragraph("; ".join(value for value in (repair.condition_after, repair.result) if value))
    if repair.attachments:
        image_items = [item for item in repair.attachments if item.media_type in {"image/jpeg", "image/png"}]
        if image_items:
            document.add_page_break()
            for item in image_items:
                paragraph = document.add_paragraph()
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                try:
                    paragraph.add_run().add_picture(io.BytesIO(item.content), width=Mm(170))
                    caption = document.add_paragraph(item.caption or item.filename)
                    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
                except (ValueError, TypeError):
                    document.add_paragraph(item.filename)
    signatures = document.add_table(rows=1, cols=2)
    signatures.style = "Table Grid"
    responsible_identity = " · ".join(value for value in ((repair.responsible_user.full_name if repair.responsible_user else ""), (repair.responsible_user.job_title if repair.responsible_user else "")) if value)
    _set_cell(signatures.cell(0, 0), f'{t["responsible"]}: {responsible_identity}', bold=True)
    _set_cell(signatures.cell(0, 1), t["signature"])
    status = document.add_paragraph(_signature_status(language, finalized_internal=True))
    _set_run_font(status.runs[0], 8, True)
    output = io.BytesIO()
    document.save(output)
    return output.getvalue()


def _build_repair_protocol_pdf_legacy(repair: Repair, language: str = "bg") -> bytes:
    language = _language(language)
    t = TEXT[language]
    _, _, body, label, title, small = _pdf_styles()
    output = io.BytesIO()
    number = repair.repair_reference or f"REP-{repair.id:06d}"
    pdf = SimpleDocTemplate(output, pagesize=A4, leftMargin=10 * mm, rightMargin=8 * mm, topMargin=7 * mm, bottomMargin=8 * mm, title=number, author="AssetCore")
    story = [_pdf_header(), Spacer(1, 2 * mm), Paragraph(f'{escape(t["repair_protocol"])} № {escape(number)}', title), Spacer(1, 2 * mm)]
    identity_values = [(t["machine"], f"{repair.machine.name}; {repair.machine.brand} {repair.machine.model or ''}".strip()), (t["inventory"], repair.machine.inventory_number), (t["serial"], repair.machine.serial_number or ""), (t["date"], repair.opened_at.strftime("%d.%m.%Y"))]
    identity = Table([[Paragraph(escape(key), label), Paragraph(escape(value), body)] for key, value in identity_values], colWidths=[45 * mm, 147 * mm])
    identity.setStyle(_pdf_table_style())
    story.append(identity)
    for key, value in (("reported_problem", repair.reported_problem), ("symptoms", repair.symptoms), ("condition_before", repair.condition_before), ("diagnosis", repair.diagnosis), ("required_work", repair.required_work), ("removed_parts", repair.removed_parts_text), ("cleaning", repair.diagnostic_cleaning), ("work", repair.work_performed)):
        story.extend([Spacer(1, 2 * mm), Paragraph(escape(t[key]), label), Paragraph(escape(value or ""), body)])
    if repair.events:
        data = [[Paragraph(escape(t["date"]), label), Paragraph(escape(t["condition"]), label), Paragraph(escape(t["description"]), label)]]
        data.extend([[Paragraph(event.created_at.strftime("%d.%m.%Y %H:%M"), small), Paragraph(escape(event.event_type), small), Paragraph(escape(event.description or ""), small)] for event in sorted(repair.events, key=lambda item: (item.created_at, item.id))])
        table = Table(data, colWidths=[34 * mm, 40 * mm, 118 * mm], repeatRows=1)
        table.setStyle(_pdf_table_style(header_rows=1))
        story.extend([Spacer(1, 2 * mm), Paragraph(escape(t["events"]), label), table])
    if repair.parts_used:
        data = [[Paragraph(escape(value), label) for value in (t["position"], t["part_number"], t["description"], t["quantity"])]]
        data.extend([[Paragraph(str(index), small), Paragraph(escape(part.part_number or ""), small), Paragraph(escape(part.description), small), Paragraph(escape(f"{part.quantity:g} {part.unit or ''}".strip()), small)] for index, part in enumerate(repair.parts_used, start=1)])
        table = Table(data, colWidths=[14 * mm, 38 * mm, 110 * mm, 30 * mm], repeatRows=1)
        table.setStyle(_pdf_table_style(header_rows=1))
        story.extend([Spacer(1, 2 * mm), Paragraph(escape(t["parts_used"]), label), table])
    if repair.participants:
        data = [[Paragraph(escape(value), label) for value in (t.get("full_name", "Три имена"), t.get("job_title", "Длъжност"), t.get("contribution", "Участие"), t.get("duration", "Време"))]]
        data.extend([
            [Paragraph(escape(participant.full_name_snapshot), small), Paragraph(escape(participant.job_title_snapshot or ""), small), Paragraph(escape(participant.contribution or ""), small), Paragraph(escape(_repair_duration(participant.minutes_worked, language)), small)]
            for participant in repair.participants
        ])
        table = Table(data, colWidths=[55 * mm, 42 * mm, 70 * mm, 25 * mm], repeatRows=1)
        table.setStyle(_pdf_table_style(header_rows=1))
        story.extend([Spacer(1, 2 * mm), Paragraph(escape(t.get("participants", "Допълнителни участници")), label), table])
    test_text = _repair_test_summary(repair, language)
    story.extend([Spacer(1, 2 * mm), Paragraph(escape(t["test"]), label), Paragraph(escape(test_text), body), Spacer(1, 2 * mm), Paragraph(escape(t["condition_after"]), label), Paragraph(escape("; ".join(value for value in (repair.condition_after, repair.result) if value)), body)])
    if repair.attachments:
        for attachment in [item for item in repair.attachments if item.media_type in {"image/jpeg", "image/png"}]:
            story.extend([PageBreak(), ReportLabImage(io.BytesIO(attachment.content), width=170 * mm, height=220 * mm, kind="proportional"), Paragraph(escape(attachment.caption or attachment.filename), small)])
    responsible_identity = " · ".join(value for value in ((repair.responsible_user.full_name if repair.responsible_user else ""), (repair.responsible_user.job_title if repair.responsible_user else "")) if value)
    signature_data = [[Paragraph(f'<b>{escape(t["responsible"])}:</b> {escape(responsible_identity)}', body), Paragraph(escape(t["signature"]), body)]]
    signatures = Table(signature_data, colWidths=[156 * mm, 36 * mm], rowHeights=[14 * mm])
    signatures.setStyle(_pdf_table_style())
    story.extend([
        Spacer(1, 3 * mm),
        KeepTogether(signatures),
        Spacer(1, 2 * mm),
        Paragraph(escape(_signature_status(language, finalized_internal=True)), label),
    ])
    pdf.build(story)
    return output.getvalue()


def build_repair_protocol_pdf(
    repair: Repair,
    language: str = "bg",
    source_reference: str = "",
) -> bytes:
    """Three-part PDF fallback mirroring the controlled v6 DOCX package."""
    language = _language(language)
    _, _, body, label, title, small = _pdf_styles()
    labels = {
        "bg": {
            "accept": "ПРОТОКОЛ ЗА ПРИЕМАНЕ НА ОБОРУДВАНЕ ЗА РЕМОНТ",
            "done": "ПРОТОКОЛ ЗА ИЗВЪРШЕН РЕМОНТ",
            "number": "Протокол №", "date": "Дата", "equipment": "Оборудване",
            "ownership": "Собственост", "condition": "Състояние при приемане",
            "problem": "Описание на проблема", "cleaning": "Почистване при диагностиката",
            "required": "Описание на необходимия ремонт", "removed": "А) Демонтаж и подготовка",
            "diagnosis": "Б) Констатирано след разглобяване, почистване и дефектация",
            "needed": "В) Нужни части за ремонт", "diagnosis_time": "Реално време за диагностика",
            "participants": "Извършили ремонта / диагностиката", "handed": "Предал оборудването",
            "accepted": "Приел оборудването", "work": "Извършени ремонтни дейности",
            "repair_time": "Реално време за ремонт", "testing_time": "Реално време за тестове",
            "parts": "Вложени резервни части", "hours": "РЕАЛНО ОТРАБОТЕНО ВРЕМЕ",
            "start": "Начало", "end": "Край", "total": "Общо реално време",
            "participant_total": "Общо време на участниците",
            "tests": "Тестове и резултат", "after": "Състояние след ремонта",
            "result": "Краен резултат", "attachments": "Приложения",
            "approved": "Приел ремонта", "reference": "Ремонтна референция",
            "source": "Свързано връщане",
        },
        "en": {
            "accept": "EQUIPMENT ACCEPTANCE FOR REPAIR PROTOCOL", "done": "COMPLETED REPAIR PROTOCOL",
            "number": "Protocol No.", "date": "Date", "equipment": "Equipment", "ownership": "Ownership",
            "condition": "Condition on acceptance", "problem": "Reported problem",
            "cleaning": "Cleaning during diagnosis", "required": "Required repair", "removed": "A) Disassembly and preparation",
            "diagnosis": "B) Findings after inspection", "needed": "C) Parts required", "diagnosis_time": "Actual diagnosis time",
            "participants": "Repair / diagnosis participants", "handed": "Handed over by", "accepted": "Accepted by",
            "work": "Repair work performed", "repair_time": "Actual repair time", "testing_time": "Actual testing time",
            "parts": "Spare parts used", "hours": "ACTUAL LABOUR TIME", "start": "Start", "end": "End", "total": "Total actual time",
            "participant_total": "Total participant labour time",
            "tests": "Tests and result", "after": "Condition after repair", "result": "Final result", "attachments": "Attachments",
            "approved": "Repair accepted by", "reference": "Repair reference", "source": "Linked return",
        },
        "ru": {
            "accept": "ПРОТОКОЛ ПРИЕМА ОБОРУДОВАНИЯ В РЕМОНТ", "done": "ПРОТОКОЛ ВЫПОЛНЕННОГО РЕМОНТА",
            "number": "Протокол №", "date": "Дата", "equipment": "Оборудование", "ownership": "Собственность",
            "condition": "Состояние при приеме", "problem": "Описание проблемы",
            "cleaning": "Очистка при диагностике", "required": "Необходимый ремонт", "removed": "А) Демонтаж и подготовка",
            "diagnosis": "Б) Результаты дефектации", "needed": "В) Необходимые детали", "diagnosis_time": "Фактическое время диагностики",
            "participants": "Участники ремонта / диагностики", "handed": "Передал", "accepted": "Принял",
            "work": "Выполненные работы", "repair_time": "Фактическое время ремонта", "testing_time": "Фактическое время испытаний",
            "parts": "Использованные детали", "hours": "ФАКТИЧЕСКОЕ РАБОЧЕЕ ВРЕМЯ", "start": "Начало", "end": "Окончание", "total": "Общее время",
            "participant_total": "Общее время участников",
            "tests": "Испытания и результат", "after": "Состояние после ремонта", "result": "Итоговый результат", "attachments": "Приложения",
            "approved": "Ремонт принял", "reference": "Референция ремонта", "source": "Связанный возврат",
        },
    }[language]
    output = io.BytesIO()
    number = repair.repair_reference or f"REP-{repair.id:06d}"
    pdf = SimpleDocTemplate(
        output, pagesize=A4, leftMargin=10 * mm, rightMargin=8 * mm,
        topMargin=7 * mm, bottomMargin=8 * mm, title=number, author="AssetCore",
    )

    def field_table(rows: list[tuple[str, str]]) -> Table:
        table = Table(
            [[Paragraph(escape(key), label), Paragraph(escape(value or ""), body)] for key, value in rows],
            colWidths=[48 * mm, 144 * mm],
        )
        table.setStyle(_pdf_table_style())
        return table

    def box(heading: str, value: str | None) -> Table:
        table = Table(
            [[Paragraph(escape(heading), label)], [Paragraph(escape(value or ""), body)]],
            colWidths=[192 * mm],
        )
        table.setStyle(_pdf_table_style(header_rows=1))
        return table

    participant_text = "; ".join(
        " — ".join(
            value for value in (
                item.full_name_snapshot, item.job_title_snapshot, item.contribution,
                _repair_duration(item.minutes_worked, language),
            ) if value
        )
        for item in repair.participants
    )
    machine_text = (
        f"{repair.machine.name} · №{repair.machine.inventory_number} · "
        f"{repair.machine.brand} {repair.machine.model or ''} · S/N {repair.machine.serial_number or ''}"
    )
    total_minutes = sum(
        value or 0 for value in (
            repair.diagnosis_minutes, repair.repair_minutes, repair.testing_minutes
        )
    )
    participant_total_minutes = sum(
        item.minutes_worked or 0 for item in repair.participants
    )
    story: list[object] = [
        _pdf_header(), Spacer(1, 2 * mm), Paragraph(escape(labels["accept"]), title),
        field_table([
            (labels["number"], number), (labels["date"], repair.opened_at.strftime("%d.%m.%Y")),
            (labels["reference"], number), (labels["source"], source_reference),
            (labels["equipment"], machine_text), (labels["ownership"], repair.machine.ownership or ""),
        ]), Spacer(1, 2 * mm), box(labels["problem"], repair.reported_problem),
        Spacer(1, 2 * mm), box(labels["condition"], repair.condition_before),
        Spacer(1, 2 * mm), box(labels["required"], repair.required_work),
        Spacer(1, 2 * mm), box(labels["removed"], repair.removed_parts_text),
        Spacer(1, 2 * mm), box(labels["cleaning"], repair.diagnostic_cleaning),
        Spacer(1, 2 * mm), box(labels["diagnosis"], repair.diagnosis),
        Spacer(1, 2 * mm), box(labels["needed"], repair.required_parts_text),
        Spacer(1, 2 * mm), field_table([
            (labels["diagnosis_time"], _repair_duration(repair.diagnosis_minutes, language)),
            (labels["participants"], participant_text),
            (labels["handed"], repair.reported_by_name or ""),
            (labels["accepted"], repair.accepted_by.full_name if repair.accepted_by else ""),
        ]), PageBreak(), _pdf_header(), Spacer(1, 2 * mm),
        Paragraph(escape(labels["done"]), title),
        field_table([
            (labels["number"], number),
            (labels["date"], (repair.closed_at or repair.opened_at).strftime("%d.%m.%Y")),
            (labels["equipment"], machine_text), (labels["ownership"], repair.machine.ownership or ""),
        ]), Spacer(1, 2 * mm), box(labels["work"], repair.work_performed),
        Spacer(1, 2 * mm), field_table([
            (labels["repair_time"], _repair_duration(repair.repair_minutes, language)),
            (labels["testing_time"], _repair_duration(repair.testing_minutes, language)),
            (labels["participants"], participant_text),
        ]), Spacer(1, 2 * mm), Paragraph(escape(labels["parts"]), label),
    ]
    part_rows = [[Paragraph(escape(value), label) for value in ("№", "Part No.", "Описание", "Количество")]]
    part_rows.extend([
        [Paragraph(str(index), small), Paragraph(escape(part.part_number or ""), small),
         Paragraph(escape(part.description), small), Paragraph(escape(f"{part.quantity:g} {part.unit or ''}".strip()), small)]
        for index, part in enumerate(repair.parts_used, start=1)
    ])
    if len(part_rows) == 1:
        part_rows.append([Paragraph("—", small), Paragraph("", small), Paragraph("", small), Paragraph("", small)])
    parts_table = Table(part_rows, colWidths=[12 * mm, 40 * mm, 108 * mm, 32 * mm], repeatRows=1)
    parts_table.setStyle(_pdf_table_style(header_rows=1))
    story.extend([
        parts_table, PageBreak(), _pdf_header(), Spacer(1, 2 * mm),
        Paragraph(escape(labels["hours"]), title),
        field_table([
            (labels["start"], repair.started_at.strftime("%d.%m.%Y %H:%M") if repair.started_at else ""),
            (labels["end"], repair.closed_at.strftime("%d.%m.%Y %H:%M") if repair.closed_at else ""),
            (labels["total"], _repair_duration(total_minutes, language)),
            (labels["participant_total"], _repair_duration(participant_total_minutes, language)),
        ]), Spacer(1, 2 * mm), box(labels["tests"], _repair_test_summary(repair, language)),
        Spacer(1, 2 * mm), box(labels["after"], repair.condition_after),
        Spacer(1, 2 * mm), box(labels["result"], repair.result),
        Spacer(1, 2 * mm), box(labels["attachments"], "; ".join(item.filename for item in repair.attachments)),
        Spacer(1, 2 * mm), field_table([
            (labels["participants"], participant_text),
            (labels["approved"], " — ".join(value for value in (
                repair.approved_by.full_name if repair.approved_by else "",
                repair.approved_by.job_title if repair.approved_by else "",
            ) if value)),
            (labels["date"], (repair.closed_at or repair.opened_at).strftime("%d.%m.%Y")),
        ]),
    ])
    pdf.build(story)
    return output.getvalue()


def _repair_test_summary(repair: Repair, language: str = "bg") -> str:
    language = _language(language)
    t = TEXT[language]
    values = (
        f'{t["test_method"]}: {repair.test_method}' if repair.test_method else None,
        f'{t["test_pressure"]}: {repair.test_pressure_bar} bar' if repair.test_pressure_bar is not None else None,
        f'{t["leaks"]}: {t["yes"] if repair.leaks_detected else t["no"]}' if repair.leaks_detected is not None else None,
        f'{t["electrical_test"]}: {repair.electrical_test_result}' if repair.electrical_test_result else None,
        f'{t["functional_test"]}: {repair.functional_test_result}' if repair.functional_test_result else None,
        repair.test_details,
        f'{ {"bg": "Краен резултат", "en": "Final result", "ru": "Итоговый результат"}[language]}: {repair.result}' if repair.result else None,
    )
    passed = None
    if repair.test_passed is not None:
        passed = f'{t["test"]}: {t["yes"] if repair.test_passed else t["no"]}'
    return "; ".join(value for value in (passed, *values) if value)


def _repair_duration(minutes: int | None, language: str = "bg") -> str:
    if minutes is None:
        return ""
    hours, remainder = divmod(minutes, 60)
    labels = {
        "bg": ("ч", "мин"),
        "en": ("h", "min"),
        "ru": ("ч", "мин"),
    }[_language(language)]
    parts = []
    if hours:
        parts.append(f"{hours} {labels[0]}")
    if remainder or not parts:
        parts.append(f"{remainder} {labels[1]}")
    return " ".join(parts)
