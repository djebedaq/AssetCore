from __future__ import annotations

import base64
import hashlib
import io
from pathlib import Path

from cryptography.fernet import Fernet
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Mm
from PIL import Image
from pypdf import PdfReader, PdfWriter
from reportlab.lib.pagesizes import A4
from reportlab.lib.units import mm
from reportlab.lib.utils import ImageReader
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.pdfgen import canvas
from sqlalchemy import select
from sqlalchemy.orm import Session

from .models import DocumentParticipant, DocumentSignature, OfficialDocumentVersion
from .settings import settings
from .template_engine import convert_docx_to_pdf

LEFT_SIGNATURE_MARKER = "[[ASSETCORE_LEFT_SIGNATURE]]"
RIGHT_SIGNATURE_MARKER = "[[ASSETCORE_RIGHT_SIGNATURE]]"

# Exact slot-to-field mapping for the original two-row protocol layout.
_SIGNATURE_MARKER_BY_SLOT = {
    "HANDOVER": LEFT_SIGNATURE_MARKER,
    "RETURNED_BY": LEFT_SIGNATURE_MARKER,
    "ACCEPTANCE": RIGHT_SIGNATURE_MARKER,
    "ACCEPTED_RETURN": RIGHT_SIGNATURE_MARKER,
}


def _sha(value: bytes) -> str:
    return hashlib.sha256(value).hexdigest()


def _fernet() -> Fernet:
    material = (settings.signature_encryption_key or settings.secret_key).encode("utf-8")
    return Fernet(base64.urlsafe_b64encode(hashlib.sha256(material).digest()))


def _all_paragraphs(document: Document):
    yield from document.paragraphs
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                yield from cell.paragraphs
    for section in document.sections:
        yield from section.header.paragraphs
        yield from section.footer.paragraphs


def _replace_signed_status(document: Document, language: str) -> None:
    replacements = {
        "bg": ("НЕПЪЛНО ПОДПИСАН", "ПОДПИСАН"),
        "en": ("NOT FULLY SIGNED", "SIGNED"),
        "ru": ("ПОДПИСАН НЕ ПОЛНОСТЬЮ", "ПОДПИСАН"),
    }
    old, new = replacements.get(language, replacements["bg"])
    for paragraph in _all_paragraphs(document):
        if old not in paragraph.text:
            continue
        value = paragraph.text.replace(old, new)
        if paragraph.runs:
            paragraph.runs[0].text = value
            for run in paragraph.runs[1:]:
                run.text = ""


def _marker_exists(document: Document) -> bool:
    return any(
        marker in paragraph.text
        for paragraph in _all_paragraphs(document)
        for marker in (LEFT_SIGNATURE_MARKER, RIGHT_SIGNATURE_MARKER)
    )


def _embed_signature_image(document: Document, marker: str, image: bytes) -> bool:
    for paragraph in _all_paragraphs(document):
        if marker not in paragraph.text:
            continue
        for run in paragraph.runs:
            run.text = ""
        with Image.open(io.BytesIO(image)) as source_image:
            width, height = source_image.size
        max_width_mm, max_height_mm = 29.0, 13.5
        ratio = min(
            max_width_mm / max(width, 1),
            max_height_mm / max(height, 1),
        )
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.space_before = 0
        paragraph.paragraph_format.space_after = 0
        paragraph.add_run().add_picture(
            io.BytesIO(image),
            width=Mm(max(width * ratio, 1)),
            height=Mm(max(height * ratio, 1)),
        )
        return True
    return False


def signature_rows(
    db: Session,
    participants: list[DocumentParticipant],
) -> list[tuple[DocumentParticipant, DocumentSignature, bytes]]:
    rows: list[tuple[DocumentParticipant, DocumentSignature, bytes]] = []
    for participant in participants:
        signature = db.scalar(
            select(DocumentSignature).where(
                DocumentSignature.participant_id == participant.id,
                DocumentSignature.confirmed_at.is_not(None),
            )
        )
        if signature:
            rows.append((participant, signature, _fernet().decrypt(signature.image_encrypted)))
    return rows


def finalize_signed_files(
    db: Session,
    version: OfficialDocumentVersion,
    participants: list[DocumentParticipant],
) -> None:
    finalize_signed_files_from_rows(version, signature_rows(db, participants))


def finalize_signed_files_from_rows(
    version: OfficialDocumentVersion,
    rows: list[tuple[DocumentParticipant, DocumentSignature, bytes]],
) -> None:
    if not rows:
        return
    final_docx = version.docx_content
    signatures_embedded = False
    if final_docx:
        document = Document(io.BytesIO(final_docx))
        _replace_signed_status(document, version.language)
        has_protocol_markers = _marker_exists(document)
        if has_protocol_markers:
            embedded = 0
            expected = 0
            for participant, _, image in rows:
                marker = _SIGNATURE_MARKER_BY_SLOT.get(participant.slot_code)
                if marker is None:
                    continue
                expected += 1
                if _embed_signature_image(document, marker, image):
                    embedded += 1
            if expected != 2 or embedded != expected:
                raise RuntimeError(
                    "Подписите не могат да бъдат поставени надеждно в оригиналните полета на протокола."
                )
            signatures_embedded = True
        else:
            _append_signature_annex_docx(document, rows, version.language)
        output = io.BytesIO()
        document.save(output)
        final_docx = output.getvalue()
        version.docx_content = final_docx
        version.docx_sha256 = _sha(final_docx)

    converted = convert_docx_to_pdf(final_docx) if final_docx else None
    if converted:
        version.pdf_content = converted
        version.pdf_sha256 = _sha(converted)
    elif version.pdf_content:
        if signatures_embedded:
            version.pdf_content = _stamp_protocol_signatures_pdf(version.pdf_content, rows)
        else:
            version.pdf_content = _append_signature_annex_pdf(version, rows)
        version.pdf_sha256 = _sha(version.pdf_content)


def _append_signature_annex_docx(
    document: Document,
    rows: list[tuple[DocumentParticipant, DocumentSignature, bytes]],
    language: str,
) -> None:
    document.add_page_break()
    document.add_heading(
        {
            "bg": "Потвърдени подписи",
            "en": "Confirmed signatures",
            "ru": "Подтверждённые подписи",
        }.get(language, "Потвърдени подписи"),
        level=1,
    )
    for participant, signature, image in rows:
        snapshot = participant.identity_snapshot
        table = document.add_table(rows=1, cols=2)
        table.style = "Table Grid"
        table.cell(0, 0).text = "\n".join(
            filter(
                None,
                [
                    str(snapshot.get("display_name") or ""),
                    str(snapshot.get("job_title") or ""),
                    str(snapshot.get("company") or ""),
                    participant.operation_role,
                    f"SHA-256: {signature.signature_sha256}",
                    f"UTC: {signature.confirmed_at.isoformat()}Z",
                ],
            )
        )
        with Image.open(io.BytesIO(image)) as source_image:
            width, height = source_image.size
        max_width, max_height = 55.0, 18.0
        ratio = min(max_width / max(width, 1), max_height / max(height, 1))
        table.cell(0, 1).paragraphs[0].add_run().add_picture(
            io.BytesIO(image), width=Mm(width * ratio), height=Mm(height * ratio)
        )


def _stamp_protocol_signatures_pdf(
    original_pdf: bytes,
    rows: list[tuple[DocumentParticipant, DocumentSignature, bytes]],
) -> bytes:
    reader = PdfReader(io.BytesIO(original_pdf))
    if not reader.pages:
        return original_pdf
    first_page = reader.pages[0]
    page_width = float(first_page.mediabox.width)
    page_height = float(first_page.mediabox.height)
    overlay_stream = io.BytesIO()
    overlay = canvas.Canvas(overlay_stream, pagesize=(page_width, page_height), pageCompression=1)

    # Coordinates match transfer_* v3 templates: signature images are in the right
    # 35 mm column, first and second 22 mm signature rows respectively.
    positions = {
        LEFT_SIGNATURE_MARKER: (165 * mm, 111 * mm),
        RIGHT_SIGNATURE_MARKER: (165 * mm, 88.5 * mm),
    }
    for participant, _, image in rows:
        marker = _SIGNATURE_MARKER_BY_SLOT.get(participant.slot_code)
        if marker not in positions:
            continue
        x, y = positions[marker]
        overlay.drawImage(
            ImageReader(io.BytesIO(image)),
            x,
            y,
            width=29 * mm,
            height=13.5 * mm,
            preserveAspectRatio=True,
            anchor="c",
            mask="auto",
        )
    overlay.save()
    overlay_stream.seek(0)
    overlay_page = PdfReader(overlay_stream).pages[0]
    first_page.merge_page(overlay_page)
    writer = PdfWriter()
    for page in reader.pages:
        writer.add_page(page)
    output = io.BytesIO()
    writer.write(output)
    return output.getvalue()


def _append_signature_annex_pdf(
    version: OfficialDocumentVersion,
    rows: list[tuple[DocumentParticipant, DocumentSignature, bytes]],
) -> bytes:
    annex = _signature_annex_pdf(version, rows)
    writer = PdfWriter()
    for page in PdfReader(io.BytesIO(version.pdf_content)).pages:
        writer.add_page(page)
    for page in PdfReader(io.BytesIO(annex)).pages:
        writer.add_page(page)
    output = io.BytesIO()
    writer.write(output)
    return output.getvalue()


def _signature_annex_pdf(
    version: OfficialDocumentVersion,
    rows: list[tuple[DocumentParticipant, DocumentSignature, bytes]],
) -> bytes:
    output = io.BytesIO()
    pdf = canvas.Canvas(output, pagesize=A4, pageCompression=1)
    font_name = "Helvetica"
    for candidate in (
        "C:/Windows/Fonts/arial.ttf",
        "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",
    ):
        if Path(candidate).is_file():
            font_name = "AssetCoreUnicode"
            if font_name not in pdfmetrics.getRegisteredFontNames():
                pdfmetrics.registerFont(TTFont(font_name, candidate))
            break
    _, height = A4
    pdf.setFont(font_name, 15)
    pdf.drawString(
        42,
        height - 48,
        {
            "bg": "Потвърдени подписи",
            "en": "Confirmed signatures",
            "ru": "Подтверждённые подписи",
        }.get(version.language, "Confirmed signatures"),
    )
    y = height - 84
    for participant, signature, image in rows:
        if y < 150:
            pdf.showPage()
            y = height - 55
        snapshot = participant.identity_snapshot
        pdf.setFont(font_name, 10)
        pdf.drawString(42, y, str(snapshot.get("display_name") or ""))
        y -= 15
        pdf.setFont(font_name, 8)
        pdf.drawString(
            42,
            y,
            " · ".join(
                filter(
                    None,
                    [
                        str(snapshot.get("job_title") or ""),
                        str(snapshot.get("company") or ""),
                        participant.operation_role,
                    ],
                )
            ),
        )
        y -= 55
        pdf.drawImage(
            ImageReader(io.BytesIO(image)),
            42,
            y,
            width=170,
            height=48,
            preserveAspectRatio=True,
            anchor="sw",
            mask="auto",
        )
        pdf.drawString(230, y + 28, f"SHA-256: {signature.signature_sha256[:32]}…")
        pdf.drawString(230, y + 13, f"UTC: {signature.confirmed_at.isoformat()}Z")
        y -= 35
    pdf.save()
    return output.getvalue()
