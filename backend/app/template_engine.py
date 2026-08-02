from __future__ import annotations

import hashlib
import io
import os
import re
import shutil
import signal
import subprocess
import tempfile
from pathlib import Path
from threading import RLock

from docx import Document
from docx.document import Document as DocumentObject
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph

from .models import DocumentTemplateVersion, DocumentType

TOKEN_RE = re.compile(r"\{\{([A-Za-z0-9_:.-]+)\}\}")
DOCX_MEDIA_TYPE = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
_LIBREOFFICE_LOCK = RLock()


class TemplateValidationError(ValueError):
    pass


def source_bytes(version: DocumentTemplateVersion) -> bytes:
    if version.source_content is not None:
        return version.source_content
    if not version.source_path:
        raise TemplateValidationError("Липсва изходен файл на шаблона.")
    resources = Path(__file__).resolve().parents[1] / "resources"
    path = (resources / version.source_path).resolve()
    if resources.resolve() not in path.parents or not path.is_file():
        raise TemplateValidationError("Изходният файл на шаблона не е намерен.")
    return path.read_bytes()


def _paragraphs(document: DocumentObject):
    yield from document.paragraphs
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                yield from cell.paragraphs
    for section in document.sections:
        yield from section.header.paragraphs
        yield from section.footer.paragraphs


def _load_docx(content: bytes) -> DocumentObject:
    if not content.startswith(b"PK"):
        raise TemplateValidationError("Шаблонът не е DOCX файл.")
    try:
        return Document(io.BytesIO(content))
    except Exception as exc:
        raise TemplateValidationError("DOCX шаблонът е повреден или нечетим.") from exc


def validate_template(version: DocumentTemplateVersion) -> dict:
    errors: list[str] = []
    warnings: list[str] = []
    try:
        content = source_bytes(version)
    except TemplateValidationError as exc:
        return {"valid": False, "errors": [str(exc)], "warnings": [], "tokens": []}
    actual_hash = hashlib.sha256(content).hexdigest()
    if version.source_sha256 and actual_hash != version.source_sha256:
        errors.append("SHA-256 на изходния файл не съвпада с регистъра.")
    filename = (version.source_filename or version.source_path or "").casefold()
    media_type = (version.source_media_type or "").casefold()
    if not filename.endswith(".docx") and DOCX_MEDIA_TYPE not in media_type:
        errors.append("Официалният шаблон трябва да бъде DOCX файл.")
    try:
        document = _load_docx(content)
        text = "\n".join(paragraph.text for paragraph in _paragraphs(document))
        tokens = sorted(set(TOKEN_RE.findall(text)))
    except TemplateValidationError as exc:
        errors.append(str(exc))
        tokens = []
    language_token = f"TEMPLATE_LANGUAGE:{version.language}"
    required_tokens = [
        (language_token, "Шаблонът не декларира правилния език."),
        ("DOCUMENT_NUMBER", "Липсва поле за номер на документа."),
        ("SIGNATURE_STATUS", "Липсва поле за статуса на подписите."),
        ("PREPARER_NAME", "Липсва поле за съставителя."),
        ("LEFT_SIGNATURE", "Липсва лява подписна позиция."),
    ]
    document_type = version.template.document_type if version.template else None
    is_internal_repair = (
        document_type == DocumentType.REPAIR_PROTOCOL.value
        or version.source_filename.startswith("repair_protocol-")
    )
    if not is_internal_repair:
        required_tokens.append(
            ("RIGHT_SIGNATURE", "Липсва дясна подписна позиция.")
        )
    for token, message in required_tokens:
        if token not in tokens:
            errors.append(message)
    required = [str(value).upper() for value in (version.required_fields or [])]
    for token in required:
        if token not in tokens:
            errors.append(f"Липсва задължително поле {token}.")
    if not version.layout_contract:
        errors.append("Липсва договор за оформление.")
    if version.layout_contract.get("reference_only"):
        warnings.append("Източникът е означен само като референция и не може да се публикува.")
        errors.append("Шаблонът е референтен, а не машинно използваем източник.")
    return {
        "valid": not errors,
        "errors": errors,
        "warnings": warnings,
        "tokens": tokens,
        "source_sha256": actual_hash,
        "source_size": len(content),
        "language": version.language,
    }


def _replace_paragraph(paragraph: Paragraph, values: dict[str, object]) -> None:
    original = paragraph.text
    if "{{" not in original:
        return

    def replacement(match: re.Match[str]) -> str:
        token = match.group(1)
        if token.startswith("TEMPLATE_LANGUAGE:"):
            return ""
        if token.startswith("TABLE:"):
            return match.group(0)
        if token not in values:
            raise TemplateValidationError(f"Няма стойност за полето {token}.")
        value = values[token]
        return "" if value is None else str(value)

    updated = TOKEN_RE.sub(replacement, original)
    if updated == original:
        return
    if paragraph.runs:
        paragraph.runs[0].text = updated
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(updated)


def _insert_table_after(paragraph: Paragraph, rows: list[list[object]]) -> None:
    table = OxmlElement("w:tbl")
    props = OxmlElement("w:tblPr")
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = OxmlElement(f"w:{edge}")
        element.set("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val", "single")
        element.set("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}sz", "4")
        borders.append(element)
    props.append(borders)
    table.append(props)
    for values in rows:
        row = OxmlElement("w:tr")
        for value in values:
            cell = OxmlElement("w:tc")
            cell_paragraph = OxmlElement("w:p")
            run = OxmlElement("w:r")
            text = OxmlElement("w:t")
            text.text = "" if value is None else str(value)
            run.append(text)
            cell_paragraph.append(run)
            cell.append(cell_paragraph)
            row.append(cell)
        table.append(row)
    paragraph._p.addnext(table)
    paragraph._element.getparent().remove(paragraph._element)


def render_docx(
    version: DocumentTemplateVersion,
    values: dict[str, object],
    tables: dict[str, list[list[object]]] | None = None,
) -> bytes:
    validation = validate_template(version)
    if not validation["valid"]:
        raise TemplateValidationError("; ".join(validation["errors"]))
    document = _load_docx(source_bytes(version))
    table_values = tables or {}
    for paragraph in list(document.paragraphs):
        match = re.fullmatch(r"\s*\{\{TABLE:([A-Z0-9_.-]+)\}\}\s*", paragraph.text)
        if match:
            code = match.group(1)
            if code not in table_values:
                raise TemplateValidationError(f"Няма таблични данни за {code}.")
            _insert_table_after(paragraph, table_values[code])
    for paragraph in _paragraphs(document):
        _replace_paragraph(paragraph, values)
    output = io.BytesIO()
    document.save(output)
    rendered = output.getvalue()
    remaining = TOKEN_RE.findall("\n".join(p.text for p in _paragraphs(_load_docx(rendered))))
    if remaining:
        raise TemplateValidationError(
            f"След генериране останаха непопълнени полета: {', '.join(sorted(set(remaining)))}."
        )
    return rendered


def convert_docx_to_pdf(docx: bytes) -> bytes | None:
    """Convert from the filled DOCX with an isolated, bounded LO lifecycle."""
    executable = shutil.which("soffice") or shutil.which("libreoffice")
    if executable is None:
        return None
    with _LIBREOFFICE_LOCK, tempfile.TemporaryDirectory(
        prefix="assetcore-docx-"
    ) as temp_name:
        temp = Path(temp_name)
        source = temp / "document.docx"
        profile = temp / "lo-profile"
        output = temp / "output"
        profile.mkdir()
        output.mkdir()
        source.write_bytes(docx)
        environment = os.environ.copy()
        environment.update(
            {
                "HOME": str(profile),
                "XDG_CACHE_HOME": str(profile / "cache"),
                "XDG_CONFIG_HOME": str(profile / "config"),
                "XDG_DATA_HOME": str(profile / "data"),
            }
        )
        command = [
            executable,
            "--headless",
            "--nologo",
            "--nodefault",
            "--nolockcheck",
            "--nofirststartwizard",
            f"-env:UserInstallation={profile.resolve().as_uri()}",
            "--convert-to",
            "pdf",
            "--outdir",
            str(output),
            str(source),
        ]
        process_options: dict[str, object] = {
            "stdout": subprocess.PIPE,
            "stderr": subprocess.PIPE,
            "env": environment,
        }
        if os.name == "nt":
            process_options["creationflags"] = subprocess.CREATE_NEW_PROCESS_GROUP
        else:
            process_options["start_new_session"] = True
        process = subprocess.Popen(command, **process_options)
        try:
            process.communicate(timeout=45)
        except subprocess.TimeoutExpired:
            _terminate_process_tree(process)
            return None
        finally:
            if process.poll() is None:
                _terminate_process_tree(process)
        target = output / "document.pdf"
        if process.returncode != 0 or not target.is_file() or target.stat().st_size == 0:
            return None
        return target.read_bytes()


def _terminate_process_tree(process: subprocess.Popen) -> None:
    """Terminate the exact conversion process and children after a timeout."""
    if process.poll() is not None:
        return
    try:
        if os.name == "nt":
            subprocess.run(
                ["taskkill", "/PID", str(process.pid), "/T", "/F"],
                capture_output=True,
                timeout=10,
                check=False,
            )
        else:
            os.killpg(process.pid, signal.SIGKILL)
        process.communicate(timeout=5)
    except (OSError, subprocess.SubprocessError):
        process.kill()
        try:
            process.communicate(timeout=5)
        except subprocess.SubprocessError:
            pass
