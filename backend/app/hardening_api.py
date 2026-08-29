from __future__ import annotations

import base64
import hashlib
import io
import json
import re
import secrets
from datetime import timedelta
from pathlib import Path

from cryptography.fernet import Fernet
from docx import Document
from docx.shared import Mm
from fastapi import APIRouter, Depends, HTTPException, Request, Response, status
from PIL import Image
from pypdf import PdfReader, PdfWriter
from reportlab.lib.pagesizes import A4
from reportlab.lib.utils import ImageReader
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.pdfgen import canvas
from sqlalchemy import func, select
from sqlalchemy.exc import IntegrityError
from sqlalchemy.orm import Session

from .audit import add_audit_log
from .database import get_db
from .governance import emergency_routes, license_routes, owner_routes
from .governance.audit_context import _correlation_id as _correlation_id
from .governance.emergency_routes import emergency_access_status as emergency_access_status
from .governance.emergency_routes import end_emergency_access as end_emergency_access
from .governance.emergency_routes import start_emergency_access as start_emergency_access
from .governance.emergency_service import _emergency_status as _emergency_status
from .governance.license_routes import install_license as install_license
from .governance.license_routes import license_status as license_status
from .governance.license_service import _parse_license_date as _parse_license_date
from .governance.owner_routes import owner_audit_history as owner_audit_history
from .governance.owner_routes import owner_status as owner_status
from .governance.owner_routes import transfer_owner as transfer_owner
from .governance.owner_service import _ownership as _ownership
from .governance.profile_checks import _profile_complete as _profile_complete
from .governance.profile_checks import _require_complete_profile as _require_complete_profile
from .hardening_schemas import (
    ExternalSignerCreate,
    ExternalSignerOut,
    ExternalSignerUpdate,
    OfficialDocumentCreate,
    OfficialDocumentOut,
    OfficialDocumentVersionOut,
    ParticipantsAssign,
    ProfileUpdate,
    ReasonRequest,
    SignatureSessionCreate,
    SignatureSessionOut,
    SignatureSlotOut,
    SignatureSlotUpdate,
    SignatureSubmit,
    SupersedeDocumentRequest,
)
from .models import (
    Department,
    DocumentParticipant,
    DocumentSignature,
    DocumentType,
    ExternalSigner,
    Machine,
    OfficialDocument,
    OfficialDocumentStatus,
    OfficialDocumentVersion,
    ProfileStatus,
    SignatureSession,
    SignatureSlot,
    User,
    UserRole,
    utcnow,
)
from .official_documents import build_official_document_registry
from .official_documents.integrity import require_current_version, set_current_version
from .official_documents.schemas import OfficialDocumentRegistryOut
from .permissions import Permission, require_permission
from .security import get_current_active_user
from .settings import settings
from .template_engine import convert_docx_to_pdf
from .transfer_service import (
    TransferServiceError,
    finalize_signed_transfer_workflow,
)
from .user_api import serialize_user

router = APIRouter(prefix="/api", tags=["production-hardening"])

_SIGNABLE_DOCUMENT_STATUSES = frozenset(
    {
        OfficialDocumentStatus.READY_FOR_SIGNATURE.value,
        OfficialDocumentStatus.PARTIALLY_SIGNED.value,
    }
)


def _canonical(value: object) -> bytes:
    return json.dumps(value, ensure_ascii=False, sort_keys=True, separators=(",", ":"), default=str).encode("utf-8")


def _sha(value: bytes) -> str:
    return hashlib.sha256(value).hexdigest()


def _signing_hash(
    *,
    document_number: str,
    document_type: str,
    version: int,
    snapshot_sha256: str,
    docx_sha256: str | None,
    pdf_sha256: str | None,
) -> str:
    return _sha(
        _canonical(
            {
                "document_number": document_number,
                "document_type": document_type,
                "version": version,
                "snapshot_sha256": snapshot_sha256,
                "docx_sha256": docx_sha256,
                "pdf_sha256": pdf_sha256,
            }
        )
    )


def _signature_consent(language: str) -> str:
    return {
        "bg": "Потвърждавам, че положеният подпис е мой и че приемам съдържанието на документа.",
        "en": "I confirm that the signature is mine and that I accept the contents of the document.",
        "ru": "Подтверждаю, что подпись принадлежит мне и что я принимаю содержание документа.",
    }.get(language, "Потвърждавам, че положеният подпис е мой и че приемам съдържанието на документа.")


def _identity_snapshot(user: User, operation_role: str) -> dict:
    _require_complete_profile(user)
    department = user.profile_department.name_bg if user.profile_department else None
    return {
        "participant_kind": "INTERNAL",
        "user_id": user.id,
        "first_name": user.first_name,
        "middle_name": user.middle_name,
        "last_name": user.last_name,
        "display_name": user.full_name,
        "job_title": user.job_title,
        "department_id": user.department_id,
        "department": department,
        "operation_role": operation_role,
        "captured_at": utcnow().isoformat(timespec="seconds") + "Z",
    }


def _external_snapshot(signer: ExternalSigner, operation_role: str) -> dict:
    return {
        "participant_kind": "EXTERNAL",
        "external_signer_id": signer.id,
        "first_name": signer.first_name,
        "middle_name": signer.middle_name,
        "last_name": signer.last_name,
        "display_name": " ".join(filter(None, [signer.first_name, signer.middle_name, signer.last_name])),
        "job_title": signer.job_title,
        "company": signer.company,
        "is_foreign_person": signer.is_foreign_person,
        "name_exception_reason": signer.name_exception_reason,
        "participant_role": signer.participant_role,
        "operation_role": operation_role,
        "captured_at": utcnow().isoformat(timespec="seconds") + "Z",
    }


def _decode_file(value: str | None, label: str) -> bytes | None:
    if value is None:
        return None
    try:
        data = base64.b64decode(value, validate=True)
    except ValueError as exc:
        raise HTTPException(
            status_code=status.HTTP_422_UNPROCESSABLE_ENTITY,
            detail={"code": "invalid_base64", "message": f"Файлът {label} не е валиден Base64."},
        ) from exc
    if not data or len(data) > 15 * 1024 * 1024:
        raise HTTPException(
            status_code=status.HTTP_422_UNPROCESSABLE_ENTITY,
            detail={"code": "invalid_document_file", "message": f"Файлът {label} е празен или над 15 MB."},
        )
    return data


def _fernet() -> Fernet:
    material = (settings.signature_encryption_key or settings.secret_key).encode("utf-8")
    return Fernet(base64.urlsafe_b64encode(hashlib.sha256(material).digest()))


@router.put("/users/me/profile")
def complete_my_profile(
    data: ProfileUpdate,
    request: Request,
    user: User = Depends(get_current_active_user),
    db: Session = Depends(get_db),
) -> dict:
    return _update_profile(user, data, user, request, db, may_approve_exception=False)


@router.put("/users/{user_id}/profile")
def update_profile_by_id(
    user_id: int,
    data: ProfileUpdate,
    request: Request,
    actor: User = Depends(require_permission(Permission.USERS_EDIT)),
    db: Session = Depends(get_db),
) -> dict:
    target = db.get(User, user_id)
    if target is None:
        raise HTTPException(404, detail={"code": "user_not_found", "message": "Потребителят не е намерен."})
    if target.is_system_owner and actor.id != target.id:
        raise HTTPException(403, detail={"code": "owner_profile_protected", "message": "Профилът на собственика се променя само от самия собственик."})
    may_manage = (
        actor.id == target.id
        or (
            actor.is_system_owner
            and actor.role == UserRole.ADMINISTRATOR.value
            and target.role
            in {UserRole.DIRECTOR.value, UserRole.MECHANIC.value, UserRole.OBSERVER.value}
        )
        or (
            actor.role == UserRole.DIRECTOR.value
            and target.role in {UserRole.MECHANIC.value, UserRole.OBSERVER.value}
        )
    )
    if not may_manage:
        raise HTTPException(
            403,
            detail={
                "code": "user_scope_denied",
                "message": "Нямате право да променяте профила на този потребител.",
            },
        )
    return _update_profile(
        target,
        data,
        actor,
        request,
        db,
        may_approve_exception=actor.role == UserRole.ADMINISTRATOR.value,
    )


def _update_profile(
    target: User,
    data: ProfileUpdate,
    actor: User,
    request: Request,
    db: Session,
    *,
    may_approve_exception: bool,
) -> dict:
    if data.department_id is not None and db.get(Department, data.department_id) is None:
        raise HTTPException(404, detail={"code": "department_not_found", "message": "Отделът не е намерен."})
    previous = {
        "profile_status": target.profile_status,
        "display_name": target.full_name,
        "job_title": target.job_title,
        "department_id": target.department_id,
    }
    if data.legal_name_exception and not data.middle_name:
        already_approved = bool(
            target.legal_name_exception
            and target.legal_name_exception_approved_by_id
            and target.legal_name_exception_approved_at
        )
        if not may_approve_exception and not already_approved:
            add_audit_log(
                db,
                actor,
                "user_profile",
                target.id,
                "Отказано самоодобряване на изключение за име",
                {"result": "rejected", "reason": "legal_name_exception_requires_admin"},
                _correlation_id(request),
            )
            db.commit()
            raise HTTPException(
                403,
                detail={
                    "code": "legal_name_exception_requires_admin",
                    "message_key": "errors.legalNameExceptionRequiresAdmin",
                    "message": "Изключение за липсващо бащино име се одобрява от упълномощен администратор.",
                },
            )
    target.first_name = data.first_name
    target.middle_name = data.middle_name
    target.last_name = data.last_name
    target.job_title = data.job_title
    target.department_id = data.department_id
    target.legal_name_exception = data.legal_name_exception
    target.legal_name_exception_reason = data.legal_name_exception_reason
    if data.legal_name_exception and not data.middle_name and may_approve_exception:
        target.legal_name_exception_approved_by_id = actor.id
        target.legal_name_exception_approved_at = utcnow()
    elif not data.legal_name_exception:
        target.legal_name_exception_approved_by_id = None
        target.legal_name_exception_approved_at = None
    if data.preferred_language is not None:
        target.preferred_language = data.preferred_language.value
    target.full_name = " ".join(filter(None, [data.first_name, data.middle_name, data.last_name]))
    target.profile_status = ProfileStatus.COMPLETE.value
    target.updated_at = utcnow()
    add_audit_log(
        db, actor, "user_profile", target.id, "Потвърден потребителски профил",
        {"previous": previous, "new": {"profile_status": target.profile_status, "display_name": target.full_name, "job_title": target.job_title, "department_id": target.department_id, "legal_name_exception": target.legal_name_exception, "legal_name_exception_approved_by_id": target.legal_name_exception_approved_by_id}},
        _correlation_id(request),
    )
    db.commit()
    db.refresh(target)
    return serialize_user(target)


router.include_router(owner_routes.router)


router.include_router(emergency_routes.router)


router.include_router(owner_routes.transfer_router)


router.include_router(license_routes.router)


@router.get("/external-signers", response_model=list[ExternalSignerOut])
def list_external_signers(
    include_inactive: bool = False,
    _: User = Depends(require_permission(Permission.DOCUMENTS_VIEW)),
    db: Session = Depends(get_db),
) -> list[ExternalSigner]:
    query = select(ExternalSigner)
    if not include_inactive:
        query = query.where(ExternalSigner.is_active.is_(True))
    return list(db.scalars(query.order_by(ExternalSigner.last_name, ExternalSigner.first_name)))


@router.get("/document-participants/internal-candidates")
def internal_document_participant_candidates(
    _: User = Depends(require_permission(Permission.DOCUMENTS_GENERATE)),
    db: Session = Depends(get_db),
) -> list[dict]:
    users = db.scalars(
        select(User).where(
            User.is_active.is_(True),
            User.profile_status == ProfileStatus.COMPLETE.value,
        ).order_by(User.full_name, User.id)
    )
    return [
        {
            "id": user.id,
            "display_name": user.full_name,
            "job_title": user.job_title,
            "role": user.role,
        }
        for user in users
    ]


@router.post("/external-signers", response_model=ExternalSignerOut, status_code=201)
def create_external_signer(
    data: ExternalSignerCreate,
    request: Request,
    actor: User = Depends(require_permission(Permission.DOCUMENTS_GENERATE)),
    db: Session = Depends(get_db),
) -> ExternalSigner:
    _require_complete_profile(actor)
    item = ExternalSigner(**data.model_dump(), created_by_id=actor.id)
    db.add(item)
    db.flush()
    add_audit_log(db, actor, "external_signer", item.id, "Създаден външен подписващ", {"participant_role": item.participant_role, "company": item.company}, _correlation_id(request))
    db.commit()
    db.refresh(item)
    return item


@router.patch("/external-signers/{signer_id}", response_model=ExternalSignerOut)
def update_external_signer(
    signer_id: int,
    data: ExternalSignerUpdate,
    request: Request,
    actor: User = Depends(require_permission(Permission.DOCUMENTS_GENERATE)),
    db: Session = Depends(get_db),
) -> ExternalSigner:
    _require_complete_profile(actor)
    item = db.get(ExternalSigner, signer_id)
    if item is None:
        raise HTTPException(
            404,
            detail={"code": "external_signer_not_found", "message": "Външният подписващ не е намерен."},
        )
    previous = {"is_active": item.is_active, "job_title": item.job_title, "company": item.company}
    for field, value in data.model_dump(exclude_unset=True, exclude_none=True).items():
        setattr(item, field, value)
    if not item.middle_name and not (
        item.is_foreign_person
        and item.name_exception_reason
        and len(item.name_exception_reason.strip()) >= 10
    ):
        raise HTTPException(
            422,
            detail={
                "code": "external_signer_full_name_required",
                "message": "Външният подписващ трябва да има три имена; изключение е допустимо само за изрично отбелязано чуждестранно лице.",
            },
        )
    item.updated_at = utcnow()
    add_audit_log(
        db,
        actor,
        "external_signer",
        item.id,
        "Променен външен подписващ",
        {
            "previous": previous,
            "new": {"is_active": item.is_active, "job_title": item.job_title, "company": item.company},
        },
        _correlation_id(request),
    )
    db.commit()
    db.refresh(item)
    return item


@router.get("/signature-slots", response_model=list[SignatureSlotOut])
def list_signature_slots(
    document_type: str | None = None,
    _: User = Depends(require_permission(Permission.DOCUMENTS_VIEW)),
    db: Session = Depends(get_db),
) -> list[SignatureSlot]:
    query = select(SignatureSlot)
    if document_type:
        query = query.where(SignatureSlot.document_type == document_type)
    return list(db.scalars(query.order_by(SignatureSlot.document_type, SignatureSlot.sequence, SignatureSlot.id)))


@router.patch("/signature-slots/{slot_id}", response_model=SignatureSlotOut)
def update_signature_slot(
    slot_id: int,
    data: SignatureSlotUpdate,
    request: Request,
    actor: User = Depends(require_permission(Permission.SETTINGS_MANAGE)),
    db: Session = Depends(get_db),
) -> SignatureSlot:
    item = db.get(SignatureSlot, slot_id)
    if item is None:
        raise HTTPException(
            404,
            detail={"code": "signature_slot_not_found", "message": "Подписната позиция не е намерена."},
        )
    previous = {
        "required": item.required,
        "sequence": item.sequence,
        "signing_mode": item.signing_mode,
        "is_active": item.is_active,
    }
    for field, value in data.model_dump(exclude_unset=True, exclude_none=True).items():
        setattr(item, field, value)
    add_audit_log(
        db,
        actor,
        "signature_slot",
        item.id,
        "Променена конфигурация на подписна позиция",
        {
            "document_type": item.document_type,
            "code": item.code,
            "previous": previous,
            "new": {
                "required": item.required,
                "sequence": item.sequence,
                "signing_mode": item.signing_mode,
                "is_active": item.is_active,
            },
        },
        _correlation_id(request),
    )
    db.commit()
    db.refresh(item)
    return item


def _build_participants(db: Session, version: OfficialDocumentVersion, document_type: str, inputs: list) -> None:
    # A document may be saved as a draft before signers are known. Required
    # slots are enforced when the draft is opened for signing.
    if not inputs:
        return
    configured = {slot.code: slot for slot in db.scalars(select(SignatureSlot).where(SignatureSlot.document_type == document_type, SignatureSlot.is_active.is_(True)))}
    supplied = {item.slot_code for item in inputs}
    missing = sorted(code for code, slot in configured.items() if slot.required and code not in supplied)
    if missing:
        raise HTTPException(422, detail={"code": "required_signature_slots_missing", "message": f"Липсват задължителни подписни позиции: {', '.join(missing)}."})
    for item in inputs:
        slot = configured.get(item.slot_code)
        if configured and slot is None:
            raise HTTPException(422, detail={"code": "unknown_signature_slot", "message": f"Непозната подписна позиция: {item.slot_code}."})
        if item.user_id is not None:
            user = db.get(User, item.user_id)
            if user is None or not user.is_active:
                raise HTTPException(404, detail={"code": "participant_not_found", "message": "Вътрешният участник не е намерен или не е активен."})
            snapshot = _identity_snapshot(user, item.operation_role)
            kind = "INTERNAL"
            external_id = None
            user_id = user.id
        else:
            signer = db.get(ExternalSigner, item.external_signer_id)
            if signer is None or not signer.is_active:
                raise HTTPException(404, detail={"code": "participant_not_found", "message": "Външният участник не е намерен или не е активен."})
            snapshot = _external_snapshot(signer, item.operation_role)
            kind = "EXTERNAL"
            external_id = signer.id
            user_id = None
        if slot and slot.allowed_participant_kind not in {"ANY", kind}:
            raise HTTPException(422, detail={"code": "participant_kind_not_allowed", "message": f"Участникът не е допустим за позиция {item.slot_code}."})
        db.add(DocumentParticipant(document_version_id=version.id, slot_code=item.slot_code, participant_kind=kind, user_id=user_id, external_signer_id=external_id, operation_role=item.operation_role, identity_snapshot=snapshot, identity_snapshot_sha256=_sha(_canonical(snapshot))))


def _version_out(db: Session, item: OfficialDocumentVersion) -> dict:
    return {"id": item.id, "version": item.version, "status": item.status, "language": item.language, "snapshot_sha256": item.snapshot_sha256, "signing_sha256": item.signing_sha256, "docx_sha256": item.docx_sha256, "pdf_sha256": item.pdf_sha256, "correction_reason": item.correction_reason, "created_at": item.created_at, "finalized_at": item.finalized_at}


def _document_out(db: Session, item: OfficialDocument) -> dict:
    version = db.get(OfficialDocumentVersion, item.current_version_id)
    if version is None:
        raise HTTPException(
            409,
            detail={
                "code": "official_document_version_missing",
                "message": "Официалният документ няма достъпна канонична версия.",
                "document_id": item.id,
            },
        )
    participants = list(db.scalars(select(DocumentParticipant).where(DocumentParticipant.document_version_id == version.id)))
    confirmed = db.scalar(select(func.count(DocumentSignature.id)).where(DocumentSignature.document_version_id == version.id, DocumentSignature.confirmed_at.is_not(None))) or 0
    required_codes = {slot.code for slot in db.scalars(select(SignatureSlot).where(SignatureSlot.document_type == item.document_type, SignatureSlot.required.is_(True), SignatureSlot.is_active.is_(True)))}
    required = len([p for p in participants if not required_codes or p.slot_code in required_codes])
    signatures = {
        signature.participant_id: signature
        for signature in db.scalars(
            select(DocumentSignature).where(DocumentSignature.document_version_id == version.id)
        )
    }
    participant_rows = [
        {
            "id": participant.id,
            "slot_code": participant.slot_code,
            "participant_kind": participant.participant_kind,
            "operation_role": participant.operation_role,
            "identity_snapshot": participant.identity_snapshot,
            "signed": bool(signatures.get(participant.id) and signatures[participant.id].confirmed_at),
            "signature_id": signatures[participant.id].id if participant.id in signatures and signatures[participant.id].confirmed_at else None,
        }
        for participant in participants
    ]
    return {"id": item.id, "document_number": item.document_number, "document_type": item.document_type, "machine_id": item.machine_id, "transfer_id": item.transfer_id, "batch_id": item.batch_id, "created_at": item.created_at, "current_version": _version_out(db, version), "signed_count": confirmed, "required_count": required, "participants": participant_rows}


@router.post("/official-documents", response_model=OfficialDocumentOut, status_code=201)
def create_official_document(
    data: OfficialDocumentCreate,
    request: Request,
    actor: User = Depends(require_permission(Permission.DOCUMENTS_GENERATE)),
    db: Session = Depends(get_db),
) -> dict:
    _require_complete_profile(actor)
    if data.document_type in {
        DocumentType.TRANSFER_ISSUE.value,
        DocumentType.TRANSFER_RETURN.value,
    }:
        raise HTTPException(
            409,
            detail={
                "code": "transfer_document_requires_workflow",
                "message": (
                    "Протоколите за издаване и връщане се създават само "
                    "чрез съответната операция с директно подписване."
                ),
            },
        )
    if db.scalar(select(OfficialDocument.id).where(OfficialDocument.document_number == data.document_number)):
        raise HTTPException(409, detail={"code": "document_number_exists", "message": "Номерът на документа вече съществува."})
    docx, pdf = _decode_file(data.docx_base64, "DOCX"), _decode_file(data.pdf_base64, "PDF")
    snapshot = dict(data.snapshot)
    snapshot["prepared_by"] = _identity_snapshot(actor, "PREPARER")
    snapshot["document_number"] = data.document_number
    snapshot["document_type"] = data.document_type
    document = OfficialDocument(document_number=data.document_number, document_type=data.document_type, machine_id=data.machine_id, transfer_id=data.transfer_id, batch_id=data.batch_id, created_by_id=actor.id)
    db.add(document)
    db.flush()
    snapshot_sha256 = _sha(_canonical(snapshot))
    docx_sha256 = _sha(docx) if docx else None
    pdf_sha256 = _sha(pdf) if pdf else None
    version = OfficialDocumentVersion(document_id=document.id, version=1, status=OfficialDocumentStatus.READY_FOR_SIGNATURE.value if data.participants else OfficialDocumentStatus.DRAFT.value, language=data.language.value, snapshot=snapshot, snapshot_sha256=snapshot_sha256, signing_sha256=_signing_hash(document_number=data.document_number, document_type=data.document_type, version=1, snapshot_sha256=snapshot_sha256, docx_sha256=docx_sha256, pdf_sha256=pdf_sha256), docx_content=docx, docx_sha256=docx_sha256, pdf_content=pdf, pdf_sha256=pdf_sha256, prepared_by_id=actor.id)
    db.add(version)
    db.flush()
    _build_participants(db, version, data.document_type, data.participants)
    set_current_version(db, document, version)
    add_audit_log(db, actor, "official_document", document.id, "Създадена неизменяема версия на официален документ", {"document_number": document.document_number, "version": 1, "snapshot_sha256": version.snapshot_sha256, "docx_sha256": version.docx_sha256, "pdf_sha256": version.pdf_sha256, "participant_slots": [p.slot_code for p in data.participants]}, _correlation_id(request))
    db.commit()
    return _document_out(db, document)


@router.get(
    "/official-documents/registry", response_model=OfficialDocumentRegistryOut
)
def official_document_registry(
    _: User = Depends(require_permission(Permission.DOCUMENTS_VIEW)),
    db: Session = Depends(get_db),
) -> dict:
    return build_official_document_registry(db)


@router.get("/official-documents/{document_id}", response_model=OfficialDocumentOut)
def get_official_document(document_id: int, _: User = Depends(require_permission(Permission.DOCUMENTS_VIEW)), db: Session = Depends(get_db)) -> dict:
    item = db.get(OfficialDocument, document_id)
    if item is None:
        raise HTTPException(404, detail={"code": "document_not_found", "message": "Документът не е намерен."})
    return _document_out(db, item)


@router.get("/official-documents", response_model=list[OfficialDocumentOut])
def list_official_documents(
    _: User = Depends(require_permission(Permission.DOCUMENTS_VIEW)),
    db: Session = Depends(get_db),
) -> list[dict]:
    return [
        _document_out(db, item)
        for item in db.scalars(
            select(OfficialDocument)
            .join(
                OfficialDocumentVersion,
                OfficialDocumentVersion.id == OfficialDocument.current_version_id,
            )
            .order_by(OfficialDocument.created_at.desc(), OfficialDocument.id.desc())
        )
    ]


@router.post("/official-documents/{document_id}/prepare-for-signatures", response_model=OfficialDocumentOut)
@router.post("/official-documents/{document_id}/participants", response_model=OfficialDocumentOut)
def assign_document_participants(
    document_id: int,
    data: ParticipantsAssign,
    request: Request,
    actor: User = Depends(require_permission(Permission.DOCUMENTS_GENERATE)),
    db: Session = Depends(get_db),
) -> dict:
    _require_complete_profile(actor)
    document = db.scalar(select(OfficialDocument).where(OfficialDocument.id == document_id).with_for_update())
    if document is None:
        raise HTTPException(404, detail={"code": "document_not_found", "message": "Документът не е намерен."})
    version = db.get(OfficialDocumentVersion, document.current_version_id)
    if version.status != OfficialDocumentStatus.DRAFT.value:
        raise HTTPException(409, detail={"code": "participants_immutable", "message": "Участниците са заключени след отваряне за подписване."})
    existing = db.scalar(select(func.count(DocumentParticipant.id)).where(DocumentParticipant.document_version_id == version.id)) or 0
    if existing:
        raise HTTPException(409, detail={"code": "participants_already_assigned", "message": "Участниците вече са определени."})
    _build_participants(db, version, document.document_type, data.participants)
    version.status = OfficialDocumentStatus.READY_FOR_SIGNATURE.value
    add_audit_log(db, actor, "official_document", document.id, "Определени участници за подписване", {"document_version": version.version, "participant_slots": [item.slot_code for item in data.participants]}, _correlation_id(request))
    db.commit()
    return _document_out(db, document)


@router.get("/official-documents/{document_id}/versions", response_model=list[OfficialDocumentVersionOut])
def list_document_versions(document_id: int, _: User = Depends(require_permission(Permission.DOCUMENTS_VIEW)), db: Session = Depends(get_db)) -> list[dict]:
    if db.get(OfficialDocument, document_id) is None:
        raise HTTPException(404, detail={"code": "document_not_found", "message": "Документът не е намерен."})
    return [_version_out(db, item) for item in db.scalars(select(OfficialDocumentVersion).where(OfficialDocumentVersion.document_id == document_id).order_by(OfficialDocumentVersion.version))]


@router.get("/official-documents/{document_id}/versions/{version_number}/download/{file_format}")
def download_document_version(document_id: int, version_number: int, file_format: str, _: User = Depends(require_permission(Permission.DOCUMENTS_VIEW)), db: Session = Depends(get_db)) -> Response:
    document = db.get(OfficialDocument, document_id)
    version = db.scalar(select(OfficialDocumentVersion).where(OfficialDocumentVersion.document_id == document_id, OfficialDocumentVersion.version == version_number))
    if document is None or version is None:
        raise HTTPException(404, detail={"code": "document_version_not_found", "message": "Версията на документа не е намерена."})
    if (
        document.document_type in {"TRANSFER_ISSUE", "TRANSFER_RETURN"}
        and version.status not in {
            OfficialDocumentStatus.SIGNED.value,
            OfficialDocumentStatus.SUPERSEDED.value,
        }
    ):
        raise HTTPException(
            409,
            detail={
                "code": "document_awaiting_signatures",
                "message": "Окончателният документ ще бъде достъпен след всички задължителни подписи.",
            },
        )
    if file_format == "docx":
        content, media = version.docx_content, "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    elif file_format == "pdf":
        content, media = version.pdf_content, "application/pdf"
    else:
        raise HTTPException(404, detail={"code": "format_not_found", "message": "Форматът не е намерен."})
    if not content:
        raise HTTPException(404, detail={"code": "file_not_generated", "message": "Файлът не е генериран за тази версия."})
    safe_number = re.sub(r"[^A-Za-z0-9._-]", "_", document.document_number)
    return Response(content=content, media_type=media, headers={"Content-Disposition": f'attachment; filename="{safe_number}-v{version.version}.{file_format}"', "X-Content-Type-Options": "nosniff"})


@router.get("/official-documents/{document_id}/preview/{file_format}")
def preview_current_document(
    document_id: int,
    file_format: str,
    actor: User = Depends(require_permission(Permission.DOCUMENTS_VIEW)),
    db: Session = Depends(get_db),
) -> Response:
    document = db.get(OfficialDocument, document_id)
    if document is None:
        raise HTTPException(
            404,
            detail={"code": "document_not_found", "message": "Документът не е намерен."},
        )
    version = db.get(OfficialDocumentVersion, document.current_version_id)
    return download_document_version(document_id, version.version, file_format, actor, db)


@router.get("/official-documents/{document_id}/verify-hash")
def verify_official_document_hashes(
    document_id: int,
    _: User = Depends(require_permission(Permission.DOCUMENTS_VIEW)),
    db: Session = Depends(get_db),
) -> dict:
    document = db.get(OfficialDocument, document_id)
    if document is None:
        raise HTTPException(
            404,
            detail={"code": "document_not_found", "message": "Документът не е намерен."},
        )
    version = db.get(OfficialDocumentVersion, document.current_version_id)
    checks = {
        "snapshot": _sha(_canonical(version.snapshot)) == version.snapshot_sha256,
        "docx": version.docx_content is None or _sha(version.docx_content) == version.docx_sha256,
        "pdf": version.pdf_content is None or _sha(version.pdf_content) == version.pdf_sha256,
    }
    return {
        "document_id": document.id,
        "document_number": document.document_number,
        "version": version.version,
        "valid": all(checks.values()),
        "checks": checks,
    }


@router.post("/official-documents/{document_id}/finalize", response_model=OfficialDocumentOut)
def finalize_official_document(
    document_id: int,
    _: User = Depends(require_permission(Permission.DOCUMENTS_GENERATE)),
    db: Session = Depends(get_db),
) -> dict:
    document = db.get(OfficialDocument, document_id)
    if document is None:
        raise HTTPException(
            404,
            detail={"code": "document_not_found", "message": "Документът не е намерен."},
        )
    version = db.get(OfficialDocumentVersion, document.current_version_id)
    if version.status != OfficialDocumentStatus.SIGNED.value:
        raise HTTPException(
            409,
            detail={
                "code": "signatures_incomplete",
                "message_key": "errors.signaturesIncomplete",
                "message": "Документът се финализира само след потвърждаване на всички задължителни подписи.",
            },
        )
    return _document_out(db, document)


@router.get("/official-documents/{document_id}/signature-status", response_model=OfficialDocumentOut)
def official_document_signature_status(
    document_id: int,
    _: User = Depends(require_permission(Permission.DOCUMENTS_VIEW)),
    db: Session = Depends(get_db),
) -> dict:
    document = db.get(OfficialDocument, document_id)
    if document is None:
        raise HTTPException(
            404,
            detail={"code": "document_not_found", "message": "Документът не е намерен."},
        )
    return _document_out(db, document)


@router.post("/official-documents/{document_id}/supersede", response_model=OfficialDocumentOut, status_code=201)
def supersede_document(document_id: int, data: SupersedeDocumentRequest, request: Request, actor: User = Depends(require_permission(Permission.DOCUMENTS_GENERATE)), db: Session = Depends(get_db)) -> dict:
    _require_complete_profile(actor)
    document = db.scalar(select(OfficialDocument).where(OfficialDocument.id == document_id).with_for_update())
    if document is None:
        raise HTTPException(404, detail={"code": "document_not_found", "message": "Документът не е намерен."})
    previous = require_current_version(db, document)
    if previous.status in {OfficialDocumentStatus.SUPERSEDED.value, OfficialDocumentStatus.CANCELLED.value}:
        raise HTTPException(409, detail={"code": "document_not_current", "message": "Тази версия не може да бъде коригирана."})
    docx, pdf = _decode_file(data.docx_base64, "DOCX"), _decode_file(data.pdf_base64, "PDF")
    snapshot = dict(data.snapshot)
    snapshot["prepared_by"] = _identity_snapshot(actor, "PREPARER")
    snapshot["document_number"] = document.document_number
    snapshot["document_type"] = document.document_type
    snapshot["correction_reason"] = data.reason.strip()
    next_version = previous.version + 1
    snapshot_sha256 = _sha(_canonical(snapshot))
    docx_sha256 = _sha(docx) if docx else None
    pdf_sha256 = _sha(pdf) if pdf else None
    version = OfficialDocumentVersion(document_id=document.id, version=next_version, status=OfficialDocumentStatus.READY_FOR_SIGNATURE.value if data.participants else OfficialDocumentStatus.DRAFT.value, language=previous.language, template_version_id=previous.template_version_id, snapshot=snapshot, snapshot_sha256=snapshot_sha256, signing_sha256=_signing_hash(document_number=document.document_number, document_type=document.document_type, version=next_version, snapshot_sha256=snapshot_sha256, docx_sha256=docx_sha256, pdf_sha256=pdf_sha256), docx_content=docx, docx_sha256=docx_sha256, pdf_content=pdf, pdf_sha256=pdf_sha256, correction_reason=data.reason.strip(), supersedes_version_id=previous.id, prepared_by_id=actor.id)
    db.add(version)
    db.flush()
    _build_participants(db, version, document.document_type, data.participants)
    previous.status = OfficialDocumentStatus.SUPERSEDED.value
    previous.finalized_at = previous.finalized_at or utcnow()
    set_current_version(db, document, version)
    add_audit_log(db, actor, "official_document", document.id, "Създадена коригираща версия на официален документ", {"previous_version": previous.version, "new_version": version.version, "reason": data.reason.strip(), "snapshot_sha256": version.snapshot_sha256}, _correlation_id(request))
    db.commit()
    return _document_out(db, document)


@router.post("/official-documents/{document_id}/cancel", response_model=OfficialDocumentOut)
def cancel_document(
    document_id: int,
    data: ReasonRequest,
    request: Request,
    actor: User = Depends(require_permission(Permission.DOCUMENTS_GENERATE)),
    db: Session = Depends(get_db),
) -> dict:
    _require_complete_profile(actor)
    document = db.scalar(
        select(OfficialDocument).where(OfficialDocument.id == document_id).with_for_update()
    )
    if document is None:
        raise HTTPException(
            404, detail={"code": "document_not_found", "message": "Документът не е намерен."}
        )
    version = db.get(OfficialDocumentVersion, document.current_version_id)
    if version.status in {
        OfficialDocumentStatus.SIGNED.value,
        OfficialDocumentStatus.SUPERSEDED.value,
        OfficialDocumentStatus.CANCELLED.value,
    }:
        raise HTTPException(
            409,
            detail={
                "code": "document_immutable",
                "message": "Подписана, заменена или вече анулирана версия не може да бъде анулирана.",
            },
        )
    participant_ids = list(
        db.scalars(
            select(DocumentParticipant.id).where(
                DocumentParticipant.document_version_id == version.id
            )
        )
    )
    if participant_ids:
        for signing_session in db.scalars(
            select(SignatureSession).where(
                SignatureSession.participant_id.in_(participant_ids),
                SignatureSession.consumed_at.is_(None),
                SignatureSession.rejected_at.is_(None),
            )
        ):
            signing_session.rejected_at = utcnow()
    version.status = OfficialDocumentStatus.CANCELLED.value
    version.finalized_at = utcnow()
    version.correction_reason = data.reason.strip()
    add_audit_log(
        db,
        actor,
        "official_document",
        document.id,
        "Анулирана неподписана версия на официален документ",
        {"version": version.version, "reason": data.reason.strip()},
        _correlation_id(request),
    )
    db.commit()
    return _document_out(db, document)


@router.post("/signatures/sessions", response_model=SignatureSessionOut, status_code=201)
def create_signature_session(data: SignatureSessionCreate, request: Request, actor: User = Depends(require_permission(Permission.DOCUMENTS_GENERATE)), db: Session = Depends(get_db)) -> dict:
    _require_complete_profile(actor)
    participant = db.get(DocumentParticipant, data.participant_id)
    if participant is None:
        raise HTTPException(404, detail={"code": "participant_not_found", "message": "Участникът не е намерен."})
    version = db.get(OfficialDocumentVersion, participant.document_version_id)
    if version is None:
        raise HTTPException(409, detail={"code": "document_not_signable", "message": "Версията не е отворена за подписване."})
    document = db.scalar(
        select(OfficialDocument)
        .where(OfficialDocument.id == version.document_id)
        .with_for_update()
        .execution_options(populate_existing=True)
    )
    version = db.scalar(
        select(OfficialDocumentVersion)
        .where(OfficialDocumentVersion.id == version.id)
        .with_for_update()
        .execution_options(populate_existing=True)
    )
    _require_current_signable_version(document, version)
    if db.scalar(select(DocumentSignature.id).where(DocumentSignature.participant_id == participant.id)):
        raise HTTPException(409, detail={"code": "already_signed", "message": "Този участник вече е положил подпис."})
    _ensure_sequence(db, participant, version)
    raw_token = secrets.token_urlsafe(32)
    expires_at = utcnow() + timedelta(minutes=data.expires_minutes)
    session = SignatureSession(participant_id=participant.id, token_hash=_sha(raw_token.encode()), expires_at=expires_at, created_by_id=actor.id)
    db.add(session)
    db.flush()
    add_audit_log(db, actor, "signature_session", session.id, "Създадена еднократна сесия за подпис", {"participant_id": participant.id, "document_version_id": version.id, "expires_at": expires_at.isoformat()}, _correlation_id(request))
    db.commit()
    return {"signing_token": raw_token, "signing_endpoint": f"/api/signing/{raw_token}", "expires_at": expires_at}


@router.post("/signatures/sessions/{session_id}/cancel")
def cancel_signature_session(
    session_id: int,
    data: ReasonRequest,
    request: Request,
    actor: User = Depends(require_permission(Permission.DOCUMENTS_GENERATE)),
    db: Session = Depends(get_db),
) -> dict:
    item = db.scalar(
        select(SignatureSession).where(SignatureSession.id == session_id).with_for_update()
    )
    if item is None:
        raise HTTPException(
            404,
            detail={"code": "signing_session_not_found", "message": "Сесията за подпис не е намерена."},
        )
    if item.consumed_at or item.rejected_at:
        raise HTTPException(
            409,
            detail={"code": "signing_session_closed", "message": "Сесията вече е приключена."},
        )
    signature = db.scalar(
        select(DocumentSignature).where(DocumentSignature.participant_id == item.participant_id)
    )
    if signature is not None and signature.confirmed_at is None:
        db.delete(signature)
    item.rejected_at = utcnow()
    add_audit_log(
        db,
        actor,
        "signature_session",
        item.id,
        "Отменена еднократна сесия за подпис",
        {"participant_id": item.participant_id, "reason": data.reason.strip()},
        _correlation_id(request),
    )
    db.commit()
    return {"message": "Сесията за подпис е отменена."}


def _ensure_sequence(db: Session, participant: DocumentParticipant, version: OfficialDocumentVersion) -> None:
    document = db.get(OfficialDocument, version.document_id)
    current_slot = db.scalar(select(SignatureSlot).where(SignatureSlot.document_type == document.document_type, SignatureSlot.code == participant.slot_code))
    if current_slot is None or current_slot.signing_mode != "SEQUENTIAL":
        return
    previous_codes = {slot.code for slot in db.scalars(select(SignatureSlot).where(SignatureSlot.document_type == document.document_type, SignatureSlot.required.is_(True), SignatureSlot.sequence < current_slot.sequence))}
    if not previous_codes:
        return
    signed_codes = {row[0] for row in db.execute(select(DocumentParticipant.slot_code).join(DocumentSignature, DocumentSignature.participant_id == DocumentParticipant.id).where(DocumentParticipant.document_version_id == version.id, DocumentSignature.confirmed_at.is_not(None)))}
    if not previous_codes.issubset(signed_codes):
        raise HTTPException(409, detail={"code": "signature_sequence_blocked", "message": "Предходните задължителни подписи още не са потвърдени."})


def _require_current_signable_version(
    document: OfficialDocument | None,
    version: OfficialDocumentVersion | None,
    *,
    signing_session: bool = False,
) -> None:
    if (
        document is not None
        and version is not None
        and document.current_version_id == version.id
        and version.status in _SIGNABLE_DOCUMENT_STATUSES
    ):
        return
    if signing_session:
        raise HTTPException(
            410,
            detail={
                "code": "signing_session_closed",
                "message": "Сесията за подпис е изтекла или вече е приключена.",
            },
        )
    raise HTTPException(
        409,
        detail={
            "code": "document_not_signable",
            "message": "Версията не е отворена за подписване.",
        },
    )


def _require_open_signing_session(session: SignatureSession | None) -> SignatureSession:
    if session is None:
        raise HTTPException(404, detail={"code": "signing_session_not_found", "message": "Сесията за подпис не е намерена."})
    if session.expires_at < utcnow() or session.consumed_at or session.rejected_at:
        raise HTTPException(410, detail={"code": "signing_session_closed", "message": "Сесията за подпис е изтекла или вече е приключена."})
    return session


def _signing_context(token: str, db: Session, lock: bool = False):
    token_hash = _sha(token.encode())
    session = _require_open_signing_session(
        db.scalar(select(SignatureSession).where(SignatureSession.token_hash == token_hash))
    )
    participant = db.get(DocumentParticipant, session.participant_id)
    if lock:
        # Superseding locks the canonical document before updating its version.
        # Use the same order here to avoid document/version deadlocks in PostgreSQL.
        version_reference = db.get(
            OfficialDocumentVersion, participant.document_version_id
        )
        document = db.scalar(
            select(OfficialDocument)
            .where(OfficialDocument.id == version_reference.document_id)
            .with_for_update()
            .execution_options(populate_existing=True)
        )
        version = db.scalar(
            select(OfficialDocumentVersion)
            .where(OfficialDocumentVersion.id == participant.document_version_id)
            .with_for_update()
            .execution_options(populate_existing=True)
        )
        session = _require_open_signing_session(
            db.scalar(
                select(SignatureSession)
                .where(SignatureSession.id == session.id)
                .with_for_update()
                .execution_options(populate_existing=True)
            )
        )
    else:
        version = db.get(OfficialDocumentVersion, participant.document_version_id)
        document = (
            db.get(OfficialDocument, version.document_id) if version is not None else None
        )
    _require_current_signable_version(document, version, signing_session=True)
    return session, participant, version, document


@router.get("/signing/{token}")
def signing_summary(token: str, db: Session = Depends(get_db)) -> dict:
    _, participant, version, document = _signing_context(token, db)
    machine = db.get(Machine, document.machine_id) if document.machine_id else None
    operation_labels = {
        "TRANSFER_ISSUE": {"bg": "Издаване / предаване на машина", "en": "Machine issue / handover", "ru": "Выдача / передача машины"},
        "TRANSFER_RETURN": {"bg": "Връщане / приемане на машина", "en": "Machine return / acceptance", "ru": "Возврат / приём машины"},
        "REPAIR_PROTOCOL": {"bg": "Вътрешен ремонтен протокол", "en": "Internal repair protocol", "ru": "Внутренний ремонтный протокол"},
    }
    operation_description = operation_labels.get(document.document_type, {}).get(
        version.language, document.document_type
    )
    batch_signing = (version.snapshot or {}).get("batch_signing") or {}
    batch_machines = [
        {
            "id": item.get("machine_id"),
            "number": item.get("machine_number"),
            "brand": item.get("brand"),
            "model": item.get("model"),
            "protocol_number": item.get("protocol_number"),
        }
        for item in batch_signing.get("machines", [])
        if isinstance(item, dict)
    ]
    if batch_machines:
        batch_labels = {
            DocumentType.TRANSFER_ISSUE.value: {
                "bg": "Издаване / предаване на избраните машини",
                "en": "Issue / handover of the selected machines",
                "ru": "Выдача / передача выбранных машин",
            },
            DocumentType.TRANSFER_RETURN.value: {
                "bg": "Връщане / приемане на избраните машини",
                "en": "Return / acceptance of the selected machines",
                "ru": "Возврат / приём выбранных машин",
            },
        }
        operation_description = batch_labels.get(document.document_type, {}).get(
            version.language, operation_description
        )
    return {
        "document_number": document.document_number,
        "document_type": document.document_type,
        "document_version": version.version,
        "document_status": version.status,
        "document_sha256": version.signing_sha256 or version.snapshot_sha256,
        "machine": {
            "id": machine.id,
            "number": machine.inventory_number,
            "name": machine.name,
            "brand": machine.brand,
        } if machine else None,
        "machines": batch_machines,
        "batch_reference": batch_signing.get("batch_reference"),
        "batch_manifest_sha256": (version.snapshot or {}).get("batch_manifest_sha256"),
        "operation_description": operation_description,
        "operation_datetime": version.created_at,
        "participant": participant.identity_snapshot,
        "operation_role": participant.operation_role,
        "consent_notice": _signature_consent(version.language),
        "requires_confirmation": True,
    }


@router.post("/signing/{token}", status_code=201)
def submit_signature(token: str, data: SignatureSubmit, db: Session = Depends(get_db)) -> dict:
    session, participant, version, document = _signing_context(token, db, lock=True)
    _ensure_sequence(db, participant, version)
    if db.scalar(select(DocumentSignature.id).where(DocumentSignature.participant_id == participant.id)):
        raise HTTPException(409, detail={"code": "already_signed", "message": "Подписът вече е подаден."})
    try:
        image = base64.b64decode(data.image_base64.split(",", 1)[-1], validate=True)
    except ValueError as exc:
        raise HTTPException(422, detail={"code": "invalid_signature_image", "message": "Изображението на подписа е невалидно."}) from exc
    if len(image) < 40 or len(image) > 2 * 1024 * 1024 or not image.startswith(b"\x89PNG\r\n\x1a\n"):
        raise HTTPException(422, detail={"code": "invalid_signature_image", "message": "Подписът трябва да е валидно PNG изображение до 2 MB."})
    try:
        with Image.open(io.BytesIO(image)) as source_image:
            source_image.verify()
            if source_image.format != "PNG":
                raise ValueError("not PNG")
    except (OSError, ValueError) as exc:
        raise HTTPException(
            422,
            detail={
                "code": "invalid_signature_image",
                "message": "Изображението на подписа е повредено или не е валиден PNG файл.",
            },
        ) from exc
    expected_consent = _signature_consent(version.language)
    if data.consent_text.strip() != expected_consent:
        raise HTTPException(
            422,
            detail={
                "code": "signature_consent_mismatch",
                "message": "Текстът за съгласие не съответства на тази версия на документа.",
            },
        )
    image_sha256 = _sha(image)
    if db.scalar(
        select(DocumentSignature.id).where(
            DocumentSignature.image_sha256 == image_sha256
        )
    ):
        raise HTTPException(
            409,
            detail={
                "code": "signature_reuse_forbidden",
                "message": "Това изображение на подпис вече е използвано и не може да се приложи към друг документ.",
            },
        )
    strokes = [[point.model_dump() for point in stroke] for stroke in data.strokes]
    strokes_bytes = _canonical(strokes)
    signed_at = utcnow()
    document_hash = version.signing_sha256 or version.snapshot_sha256
    signature_binding = {"document_sha256": document_hash, "document_version": version.version, "participant_snapshot_sha256": participant.identity_snapshot_sha256, "strokes_sha256": _sha(strokes_bytes), "image_sha256": image_sha256, "canvas_width": data.canvas_width, "canvas_height": data.canvas_height, "signed_at": signed_at.isoformat(timespec="microseconds") + "Z", "signature_kind": "MANUAL_GRAPHIC"}
    signature = DocumentSignature(participant_id=participant.id, document_version_id=version.id, consent_text=data.consent_text, strokes_encrypted=_fernet().encrypt(strokes_bytes), image_encrypted=_fernet().encrypt(image), canvas_width=data.canvas_width, canvas_height=data.canvas_height, stroke_count=len(strokes), point_count=sum(len(stroke) for stroke in strokes), document_sha256=document_hash, image_sha256=image_sha256, signature_sha256=_sha(_canonical(signature_binding)), signed_at=signed_at)
    db.add(signature)
    try:
        db.flush()
    except IntegrityError as exc:
        db.rollback()
        raise HTTPException(
            409,
            detail={
                "code": "signature_reuse_forbidden",
                "message": "Подписът вече е използван или позицията вече е подписана.",
            },
        ) from exc
    creator = db.get(User, session.created_by_id)
    add_audit_log(db, creator, "document_signature", signature.id, "Подаден ръчен графичен подпис за потвърждение", {"document_id": document.id, "document_version": version.version, "participant_id": participant.id, "document_sha256": document_hash, "signature_sha256": signature.signature_sha256, "stroke_count": signature.stroke_count, "point_count": signature.point_count})
    db.commit()
    return {"message": "Подписът е записан. Прегледайте обобщението и го потвърдете.", "signature_sha256": signature.signature_sha256, "requires_confirmation": True}


@router.post("/signing/{token}/confirm")
def confirm_signature(token: str, db: Session = Depends(get_db)) -> dict:
    session, participant, version, document = _signing_context(token, db, lock=True)
    signature = db.scalar(select(DocumentSignature).where(DocumentSignature.participant_id == participant.id).with_for_update())
    if signature is None:
        raise HTTPException(409, detail={"code": "signature_not_submitted", "message": "Първо положете подпис."})
    now = utcnow()
    signature.confirmed_at = now
    session.consumed_at = now
    # Test sessions deliberately disable autoflush. Flush explicitly so the
    # status query always observes the signature confirmed in this transaction.
    db.flush()
    _refresh_document_status(db, document, version)
    creator = db.get(User, session.created_by_id)
    try:
        if version.status == OfficialDocumentStatus.SIGNED.value:
            finalize_signed_transfer_workflow(db, document, version, creator)
    except TransferServiceError as exc:
        db.rollback()
        raise HTTPException(exc.status_code, detail=exc.as_detail()) from exc
    add_audit_log(db, creator, "document_signature", signature.id, "Потвърден ръчен графичен подпис", {"document_id": document.id, "document_version": version.version, "participant_id": participant.id, "document_sha256": signature.document_sha256, "signature_sha256": signature.signature_sha256, "new_document_status": version.status})
    db.commit()
    return {"message": "Подписът е потвърден.", "document_status": version.status, "signature_sha256": signature.signature_sha256}


@router.post("/signing/{token}/reject")
def reject_signature(token: str, db: Session = Depends(get_db)) -> dict:
    session, participant, version, document = _signing_context(token, db, lock=True)
    signature = db.scalar(select(DocumentSignature).where(DocumentSignature.participant_id == participant.id))
    if signature is not None and signature.confirmed_at is None:
        db.delete(signature)
    session.rejected_at = utcnow()
    creator = db.get(User, session.created_by_id)
    add_audit_log(db, creator, "signature_session", session.id, "Отказан ръчен графичен подпис", {"document_id": document.id, "document_version": version.version, "participant_id": participant.id})
    db.commit()
    return {"message": "Подписването е отказано."}


def _refresh_document_status(db: Session, document: OfficialDocument, version: OfficialDocumentVersion) -> None:
    _require_current_signable_version(document, version, signing_session=True)
    participants = list(db.scalars(select(DocumentParticipant).where(DocumentParticipant.document_version_id == version.id)))
    required_codes = {slot.code for slot in db.scalars(select(SignatureSlot).where(SignatureSlot.document_type == document.document_type, SignatureSlot.required.is_(True), SignatureSlot.is_active.is_(True)))}
    required = [participant for participant in participants if not required_codes or participant.slot_code in required_codes]
    confirmed_ids = set(db.scalars(select(DocumentSignature.participant_id).where(DocumentSignature.document_version_id == version.id, DocumentSignature.confirmed_at.is_not(None))))
    if required and all(item.id in confirmed_ids for item in required):
        version.status = OfficialDocumentStatus.SIGNED.value
        version.finalized_at = utcnow()
        _finalize_signed_files(db, version, required)
    elif confirmed_ids:
        version.status = OfficialDocumentStatus.PARTIALLY_SIGNED.value
    else:
        version.status = OfficialDocumentStatus.READY_FOR_SIGNATURE.value


def _replace_signed_status(document: Document, language: str) -> None:
    replacements = {
        "bg": ("НЕПЪЛНО ПОДПИСАН", "ПОДПИСАН"),
        "en": ("NOT FULLY SIGNED", "SIGNED"),
        "ru": ("ПОДПИСАН НЕ ПОЛНОСТЬЮ", "ПОДПИСАН"),
    }
    old, new = replacements.get(language, replacements["bg"])
    paragraphs = list(document.paragraphs)
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                paragraphs.extend(cell.paragraphs)
    for paragraph in paragraphs:
        if old not in paragraph.text:
            continue
        value = paragraph.text.replace(old, new)
        if paragraph.runs:
            paragraph.runs[0].text = value
            for run in paragraph.runs[1:]:
                run.text = ""


def _finalize_signed_files(
    db: Session,
    version: OfficialDocumentVersion,
    participants: list[DocumentParticipant],
) -> None:
    rows: list[tuple[DocumentParticipant, DocumentSignature, bytes]] = []
    for participant in participants:
        signature = db.scalar(
            select(DocumentSignature).where(
                DocumentSignature.participant_id == participant.id,
                DocumentSignature.confirmed_at.is_not(None),
            )
        )
        if signature:
            rows.append((participant, signature, _fernet().decrypt(signature.image_encrypted)))
    if not rows:
        return
    final_docx = version.docx_content
    if final_docx:
        document = Document(io.BytesIO(final_docx))
        _replace_signed_status(document, version.language)
        document.add_page_break()
        document.add_heading(
            {"bg": "Потвърдени подписи", "en": "Confirmed signatures", "ru": "Подтверждённые подписи"}.get(version.language, "Потвърдени подписи"),
            level=1,
        )
        for participant, signature, image in rows:
            snapshot = participant.identity_snapshot
            table = document.add_table(rows=1, cols=2)
            table.style = "Table Grid"
            table.cell(0, 0).text = "\n".join(
                filter(
                    None,
                    [
                        str(snapshot.get("display_name") or ""),
                        str(snapshot.get("job_title") or ""),
                        str(snapshot.get("company") or ""),
                        participant.operation_role,
                        f"SHA-256: {signature.signature_sha256}",
                        f"UTC: {signature.confirmed_at.isoformat()}Z",
                    ],
                )
            )
            with Image.open(io.BytesIO(image)) as source_image:
                width, height = source_image.size
            max_width, max_height = 55.0, 18.0
            ratio = min(max_width / max(width, 1), max_height / max(height, 1))
            table.cell(0, 1).paragraphs[0].add_run().add_picture(
                io.BytesIO(image), width=Mm(width * ratio), height=Mm(height * ratio)
            )
        output = io.BytesIO()
        document.save(output)
        final_docx = output.getvalue()
        version.docx_content = final_docx
        version.docx_sha256 = _sha(final_docx)
    converted = convert_docx_to_pdf(final_docx) if final_docx else None
    if converted:
        version.pdf_content = converted
        version.pdf_sha256 = _sha(converted)
    elif version.pdf_content:
        annex = _signature_annex_pdf(version, rows)
        writer = PdfWriter()
        for page in PdfReader(io.BytesIO(version.pdf_content)).pages:
            writer.add_page(page)
        for page in PdfReader(io.BytesIO(annex)).pages:
            writer.add_page(page)
        output = io.BytesIO()
        writer.write(output)
        version.pdf_content = output.getvalue()
        version.pdf_sha256 = _sha(version.pdf_content)


def _signature_annex_pdf(
    version: OfficialDocumentVersion,
    rows: list[tuple[DocumentParticipant, DocumentSignature, bytes]],
) -> bytes:
    output = io.BytesIO()
    pdf = canvas.Canvas(output, pagesize=A4, pageCompression=1)
    font_name = "Helvetica"
    for candidate in (
        "C:/Windows/Fonts/arial.ttf",
        "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",
    ):
        if Path(candidate).is_file():
            font_name = "AssetCoreUnicode"
            if font_name not in pdfmetrics.getRegisteredFontNames():
                pdfmetrics.registerFont(TTFont(font_name, candidate))
            break
    width, height = A4
    pdf.setFont(font_name, 15)
    pdf.drawString(42, height - 48, {"bg": "Потвърдени подписи", "en": "Confirmed signatures", "ru": "Подтверждённые подписи"}.get(version.language, "Confirmed signatures"))
    y = height - 84
    for participant, signature, image in rows:
        if y < 150:
            pdf.showPage()
            y = height - 55
        snapshot = participant.identity_snapshot
        pdf.setFont(font_name, 10)
        pdf.drawString(42, y, str(snapshot.get("display_name") or ""))
        y -= 15
        pdf.setFont(font_name, 8)
        pdf.drawString(42, y, " · ".join(filter(None, [str(snapshot.get("job_title") or ""), str(snapshot.get("company") or ""), participant.operation_role])))
        y -= 55
        pdf.drawImage(ImageReader(io.BytesIO(image)), 42, y, width=170, height=48, preserveAspectRatio=True, anchor="sw", mask="auto")
        pdf.drawString(230, y + 28, f"SHA-256: {signature.signature_sha256[:32]}…")
        pdf.drawString(230, y + 13, f"UTC: {signature.confirmed_at.isoformat()}Z")
        y -= 35
    pdf.save()
    return output.getvalue()


@router.get("/signatures/{signature_id}/image")
def signature_image(signature_id: int, _: User = Depends(require_permission(Permission.DOCUMENTS_VIEW)), db: Session = Depends(get_db)) -> Response:
    item = db.get(DocumentSignature, signature_id)
    if item is None or item.confirmed_at is None:
        raise HTTPException(404, detail={"code": "signature_not_found", "message": "Подписът не е намерен."})
    return Response(content=_fernet().decrypt(item.image_encrypted), media_type="image/png", headers={"Cache-Control": "private, no-store", "Content-Disposition": "inline", "X-Content-Type-Options": "nosniff"})
