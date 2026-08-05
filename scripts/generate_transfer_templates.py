from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_ROW_HEIGHT_RULE, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Mm, Pt, RGBColor
from PIL import Image, ImageEnhance, ImageOps

ROOT = Path(__file__).resolve().parents[1]
ASSETS = ROOT / "backend" / "resources" / "assets"
TEMPLATES = ROOT / "backend" / "resources" / "templates"

LANG = {
    "bg": {
        "protocol": "ПРОТОКОЛ № {{DOCUMENT_NUMBER}}",
        "issue_title": "ПРОТОКОЛ ПРЕДАВАНЕ НА МИЕЩА ТЕХНИКА",
        "return_title": "ПРОТОКОЛ ПРИЕМАНЕ НА МИЕЩА ТЕХНИКА СЛЕД ИЗПОЛЗВАНЕ",
        "date": "ДАТА:",
        "equipment": "ОБОРУДВАНЕ:",
        "model": "МОДЕЛ:",
        "factory": "ЗАВОДСКИ НОМЕР:",
        "serial": "СЕРИЕН НОМЕР:",
        "no": "№",
        "element": "Елемент",
        "condition": "Състояние",
        "usage_issue": "Оборудването ще се използва за:",
        "usage_return": "Оборудването е било използвано за:",
        "remarks": "Забележки:",
        "issue_left": "Предал\nоборудването\nот страна на\nДИРП",
        "issue_right": "Приел\nоборудването",
        "return_left": "Върнал\nоборудването",
        "return_right": "Приел\nоборудването\nот страна на\nДИРП",
        "three_names": "(Три имена)",
        "external_names": "(Три имена)",
        "signature": "(Подпис)",
        "prepared": "Съставил:",
        "batch": "Batch manifest:",
        "status": "Статус:",
        "components": [
            "Помпа", "Шланг захранващ", "Шланг изходящ ВН", "Пистолет",
            "Дюза метла / ротационна", "Накрайници", "Кабел",
            "Куплунг / Еврощек", "Ходова част", "Корпус",
        ],
    },
    "en": {
        "protocol": "PROTOCOL No. {{DOCUMENT_NUMBER}}",
        "issue_title": "HIGH-PRESSURE WASHING EQUIPMENT ISSUE PROTOCOL",
        "return_title": "HIGH-PRESSURE WASHING EQUIPMENT RETURN PROTOCOL",
        "date": "DATE:", "equipment": "EQUIPMENT:", "model": "MODEL:",
        "factory": "FACTORY / INVENTORY No.:", "serial": "SERIAL No.:",
        "no": "No.", "element": "Component", "condition": "Condition",
        "usage_issue": "The equipment will be used for:",
        "usage_return": "The equipment was used for:",
        "remarks": "Remarks:",
        "issue_left": "Handed over\nby DIRP", "issue_right": "Accepted by",
        "return_left": "Returned by", "return_right": "Accepted by\nDIRP",
        "three_names": "(Full name)", "external_names": "(Full name)",
        "signature": "(Signature)", "prepared": "Prepared by:",
        "batch": "Batch manifest:", "status": "Status:",
        "components": [
            "Pump", "Supply hose", "High-pressure outlet hose", "Gun",
            "Fan / rotary nozzle", "Tips", "Cable", "Coupling / Euro plug",
            "Running gear", "Body",
        ],
    },
    "ru": {
        "protocol": "ПРОТОКОЛ № {{DOCUMENT_NUMBER}}",
        "issue_title": "ПРОТОКОЛ ВЫДАЧИ МОЕЧНОЙ ТЕХНИКИ",
        "return_title": "ПРОТОКОЛ ПРИЕМА МОЕЧНОЙ ТЕХНИКИ ПОСЛЕ ИСПОЛЬЗОВАНИЯ",
        "date": "ДАТА:", "equipment": "ОБОРУДОВАНИЕ:", "model": "МОДЕЛЬ:",
        "factory": "ЗАВОДСКОЙ / ИНВЕНТАРНЫЙ №:", "serial": "СЕРИЙНЫЙ №:",
        "no": "№", "element": "Элемент", "condition": "Состояние",
        "usage_issue": "Оборудование будет использовано для:",
        "usage_return": "Оборудование использовалось для:",
        "remarks": "Примечания:",
        "issue_left": "Передал\nоборудование\nсо стороны ДИРП", "issue_right": "Принял\nоборудование",
        "return_left": "Вернул\nоборудование", "return_right": "Принял\nоборудование\nсо стороны ДИРП",
        "three_names": "(ФИО)", "external_names": "(ФИО)",
        "signature": "(Подпись)", "prepared": "Составил:",
        "batch": "Batch manifest:", "status": "Статус:",
        "components": [
            "Насос", "Подающий шланг", "Выходной шланг ВД", "Пистолет",
            "Веерное / ротационное сопло", "Наконечники", "Кабель",
            "Муфта / евровилка", "Ходовая часть", "Корпус",
        ],
    },
}


def _bw(source: Path, target: Path) -> Path:
    image = Image.open(source).convert("RGBA")
    alpha = image.getchannel("A")
    gray = ImageOps.grayscale(image.convert("RGB"))
    gray = ImageEnhance.Contrast(gray).enhance(1.45)
    rgba = Image.merge("RGBA", (gray, gray, gray, alpha))
    rgba.save(target)
    return target


def _set_table_widths(table, widths_mm: tuple[float, ...]) -> None:
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width_mm in widths_mm:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(int(width_mm * 56.6929)))
        grid.append(col)
    for col, width_mm in zip(table.columns, widths_mm, strict=True):
        col.width = Mm(width_mm)
        for cell in col.cells:
            _set_cell_width(cell, width_mm)


def _set_cell_width(cell, width_mm: float) -> None:
    cell.width = Mm(width_mm)
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(int(width_mm * 56.6929)))
    tc_w.set(qn("w:type"), "dxa")


def _cell_margins(cell, top=40, start=60, bottom=40, end=60) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def _set_repeat_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def _set_exact_height(row, mm_value: float) -> None:
    row.height = Mm(mm_value)
    row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY


def _set_min_height(row, mm_value: float) -> None:
    row.height = Mm(mm_value)
    row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST


def _run_font(run, size=8.0, bold=False, color=None, hidden=False) -> None:
    run.font.name = "Arial"
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Arial")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)
    if hidden:
        vanish = OxmlElement("w:vanish")
        run._element.get_or_add_rPr().append(vanish)


def _paragraph(cell, text="", size=8.0, bold=False, align=WD_ALIGN_PARAGRAPH.LEFT, space=0):
    p = cell.paragraphs[0] if len(cell.paragraphs) == 1 and not cell.paragraphs[0].text else cell.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(space)
    p.paragraph_format.line_spacing = 1
    r = p.add_run(text)
    _run_font(r, size, bold)
    return p


def _clear_cell(cell):
    for p in list(cell.paragraphs):
        if p._element.getparent() is not None:
            p._element.getparent().remove(p._element)
    cell.add_paragraph()


def _set_cell_text(cell, text, *, size=8.0, bold=False, align=WD_ALIGN_PARAGRAPH.LEFT):
    _clear_cell(cell)
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1
    lines = text.split("\n")
    for idx, line in enumerate(lines):
        if idx:
            p.add_run().add_break()
        _run_font(p.add_run(line), size, bold)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    _cell_margins(cell)


def _add_image(cell, image_path: Path, width_mm: float, height_mm: float | None = None):
    _clear_cell(cell)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    if height_mm:
        run.add_picture(str(image_path), width=Mm(width_mm), height=Mm(height_mm))
    else:
        run.add_picture(str(image_path), width=Mm(width_mm))
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    _cell_margins(cell, 20, 20, 20, 20)


def _no_cell_split(cell):
    tc_pr = cell._tc.get_or_add_tcPr()
    cant_split = OxmlElement("w:cantSplit")
    tc_pr.append(cant_split)


def build_template(language: str, operation: str, output: Path, krz: Path, odessos: Path, rina: Path) -> None:
    t = LANG[language]
    doc = Document()
    section = doc.sections[0]
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.top_margin = Mm(9)
    section.bottom_margin = Mm(8)
    section.left_margin = Mm(11)
    section.right_margin = Mm(11)
    section.header_distance = Mm(2)
    section.footer_distance = Mm(3)

    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Arial")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    normal.font.size = Pt(8)

    header = doc.add_table(rows=1, cols=3)
    header.style = "Table Grid"
    header.alignment = WD_TABLE_ALIGNMENT.CENTER
    _set_table_widths(header, (27, 124, 37))
    _set_exact_height(header.rows[0], 25)
    _add_image(header.cell(0, 0), krz, 18.5, 20.5)
    _add_image(header.cell(0, 1), odessos, 113, 22)
    _add_image(header.cell(0, 2), rina, 31, 23)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(1)
    _run_font(p.add_run(t["protocol"]), 9.2, False)
    hidden = p.add_run(f" {{{{TEMPLATE_LANGUAGE:{language}}}}}")
    _run_font(hidden, 1, False, color=(255, 255, 255), hidden=True)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(0)
    title.paragraph_format.space_after = Pt(3)
    _run_font(title.add_run(t["return_title"] if operation == "return" else t["issue_title"]), 10.0, True)

    identity = doc.add_table(rows=5, cols=2)
    identity.style = "Table Grid"
    identity.alignment = WD_TABLE_ALIGNMENT.CENTER
    _set_table_widths(identity, (48, 140))
    labels = (t["date"], t["equipment"], t["model"], t["factory"], t["serial"])
    values = ("{{CREATION_DATE}}", "{{EQUIPMENT_TYPE}}", "{{MODEL_DISPLAY}}", "№ {{MACHINE_NUMBER}}", "{{SERIAL_NUMBER}}")
    for row, label, value in zip(identity.rows, labels, values, strict=True):
        _set_exact_height(row, 5.4)
        _set_cell_width(row.cells[0], 48)
        _set_cell_width(row.cells[1], 140)
        _set_cell_text(row.cells[0], label, size=8.0, bold=True)
        _set_cell_text(row.cells[1], value, size=8.2)

    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_before = Pt(0)
    spacer.paragraph_format.space_after = Pt(1)
    spacer.paragraph_format.line_spacing = 0.5
    _run_font(spacer.add_run(" "), 2)

    checklist = doc.add_table(rows=11, cols=3)
    checklist.style = "Table Grid"
    checklist.alignment = WD_TABLE_ALIGNMENT.CENTER
    widths = (10, 62, 116)
    _set_table_widths(checklist, widths)
    for cell, width in zip(checklist.rows[0].cells, widths, strict=True):
        _set_cell_width(cell, width)
    _set_exact_height(checklist.rows[0], 5.8)
    _set_cell_text(checklist.rows[0].cells[0], t["no"], size=8.0, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    _set_cell_text(checklist.rows[0].cells[1], t["element"], size=8.0, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    _set_cell_text(checklist.rows[0].cells[2], t["condition"], size=8.0, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    _set_repeat_header(checklist.rows[0])
    for idx, component in enumerate(t["components"], start=1):
        row = checklist.rows[idx]
        _set_min_height(row, 5.2)
        for cell, width in zip(row.cells, widths, strict=True):
            _set_cell_width(cell, width)
            _no_cell_split(cell)
        _set_cell_text(row.cells[0], f"{idx}.", size=7.8, align=WD_ALIGN_PARAGRAPH.CENTER)
        _set_cell_text(row.cells[1], component, size=7.8)
        _set_cell_text(row.cells[2], f"{{{{CHECK_{idx}}}}}", size=7.8)

    condition = doc.add_table(rows=1, cols=1)
    condition.style = "Table Grid"
    condition.alignment = WD_TABLE_ALIGNMENT.CENTER
    _set_table_widths(condition, (188,))
    _set_min_height(condition.rows[0], 6.5)
    _set_cell_width(condition.cell(0, 0), 188)
    _set_cell_text(condition.cell(0, 0), "{{CONDITION_LABEL}}: {{CONDITION_TEXT}}", size=7.8)

    usage = doc.add_table(rows=1, cols=1)
    usage.style = "Table Grid"
    usage.alignment = WD_TABLE_ALIGNMENT.CENTER
    _set_table_widths(usage, (188,))
    _set_min_height(usage.rows[0], 10.5)
    _set_cell_width(usage.cell(0, 0), 188)
    _set_cell_text(usage.cell(0, 0), (t["usage_return"] if operation == "return" else t["usage_issue"]) + " {{USAGE_TEXT}}", size=7.8)

    remarks = doc.add_table(rows=1, cols=1)
    remarks.style = "Table Grid"
    remarks.alignment = WD_TABLE_ALIGNMENT.CENTER
    _set_table_widths(remarks, (188,))
    _set_min_height(remarks.rows[0], 9.0)
    _set_cell_width(remarks.cell(0, 0), 188)
    _set_cell_text(remarks.cell(0, 0), t["remarks"] + " {{REMARKS}}", size=7.6)

    signatures = doc.add_table(rows=2, cols=3)
    signatures.style = "Table Grid"
    signatures.alignment = WD_TABLE_ALIGNMENT.CENTER
    widths = (36, 117, 35)
    _set_table_widths(signatures, widths)
    left_label = t["return_left"] if operation == "return" else t["issue_left"]
    right_label = t["return_right"] if operation == "return" else t["issue_right"]
    rows = [
        (left_label, "{{LEFT_SIGNER_NAME}}", "{{LEFT_SIGNER_ROLE}}", "{{LEFT_SIGNATURE}}"),
        (right_label, "{{RIGHT_SIGNER_NAME}}", "{{RIGHT_SIGNER_ROLE}}", "{{RIGHT_SIGNATURE}}"),
    ]
    for row, (label, name, role, marker) in zip(signatures.rows, rows, strict=True):
        _set_exact_height(row, 22.0)
        for cell, width in zip(row.cells, widths, strict=True):
            _set_cell_width(cell, width)
            _no_cell_split(cell)
        _set_cell_text(row.cells[0], label, size=7.8, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
        _clear_cell(row.cells[1])
        p_name = row.cells[1].paragraphs[0]
        p_name.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_name.paragraph_format.space_before = Pt(0)
        p_name.paragraph_format.space_after = Pt(0)
        _run_font(p_name.add_run(name), 8.4, False)
        p_role = row.cells[1].add_paragraph()
        p_role.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_role.paragraph_format.space_before = Pt(1)
        p_role.paragraph_format.space_after = Pt(0)
        _run_font(p_role.add_run(role), 7.2, False)
        p_hint = row.cells[1].add_paragraph()
        p_hint.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_hint.paragraph_format.space_before = Pt(1)
        p_hint.paragraph_format.space_after = Pt(0)
        _run_font(p_hint.add_run(t["three_names"]), 6.8, False)
        row.cells[1].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        _cell_margins(row.cells[1], 40, 50, 40, 50)

        _clear_cell(row.cells[2])
        p_marker = row.cells[2].paragraphs[0]
        p_marker.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_marker.paragraph_format.space_before = Pt(0)
        p_marker.paragraph_format.space_after = Pt(0)
        marker_run = p_marker.add_run(marker)
        _run_font(marker_run, 1, False, color=(255, 255, 255), hidden=True)
        p_sig = row.cells[2].add_paragraph()
        p_sig.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_sig.paragraph_format.space_before = Pt(0)
        p_sig.paragraph_format.space_after = Pt(0)
        _run_font(p_sig.add_run(t["signature"]), 6.8, False)
        row.cells[2].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        _cell_margins(row.cells[2], 20, 20, 20, 20)

    footer = doc.add_paragraph()
    footer.paragraph_format.space_before = Pt(1)
    footer.paragraph_format.space_after = Pt(0)
    footer.alignment = WD_ALIGN_PARAGRAPH.LEFT
    _run_font(footer.add_run(f'{t["prepared"]} {{{{PREPARER_NAME}}}} · {{{{PREPARER_JOB_TITLE}}}}'), 6.5)
    footer.add_run("\n")
    _run_font(footer.add_run(f'{t["batch"]} {{{{BATCH_REFERENCE}}}} · {t["status"]} {{{{SIGNATURE_STATUS}}}}'), 6.3)

    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output)


def main() -> None:
    krz = _bw(ASSETS / "krz_logo.png", ASSETS / "krz_logo_bw.png")
    odessos = _bw(ASSETS / "odessos_logo.png", ASSETS / "odessos_logo_bw.png")
    rina = _bw(ASSETS / "rina_aqap.jpeg", ASSETS / "rina_aqap_bw.png")
    for language in LANG:
        for operation in ("issue", "return"):
            name = f"transfer_{operation}-{language}-v3.docx"
            build_template(language, operation, TEMPLATES / name, krz, odessos, rina)
            print(TEMPLATES / name)


if __name__ == "__main__":
    main()
