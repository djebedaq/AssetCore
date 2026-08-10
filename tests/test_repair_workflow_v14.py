from __future__ import annotations

import base64
import hashlib
import io
import zipfile
from concurrent.futures import ThreadPoolExecutor
from pathlib import Path

from app.models import (
    AuditLog,
    GeneratedDocument,
    Location,
    Machine,
    PartCatalog,
    Repair,
    RepairAttachment,
    RepairEvent,
    RepairPart,
    RepairParticipant,
    User,
)
from docx import Document
from pypdf import PdfReader
from sqlalchemy import select

ROOT = Path(__file__).resolve().parents[1]


def _create_repair(client, headers, machine_id: int) -> int:
    response = client.post(
        "/api/repair-cases",
        headers=headers,
        json={
            "machine_id": machine_id,
            "reported_problem": "Контролен проблем в изолираната тестова база",
            "condition_before": "Приета за контролна диагностика",
        },
    )
    assert response.status_code == 201, response.text
    return response.json()["id"]


def _advance_to_final_stage(client, headers, repair_id: int) -> None:
    transitions = [
        {
            "status": "DIAGNOSIS",
        },
        {
            "status": "REPAIRING",
            "diagnosis": "Контролна диагностика в изолираната тестова база",
            "required_work": "Проверена ремонтна операция",
            "required_parts_text": "Не са необходими допълнителни части",
            "diagnostic_cleaning": "Извършено контролно почистване при диагностиката",
            "diagnosis_minutes": 35,
        },
        {
            "work_performed": "Изпълнена и проверена ремонтна операция",
            "removed_parts_text": "Няма демонтирани части",
            "repair_minutes": 70,
            "advance_to_final": True,
        },
    ]
    for payload in transitions:
        response = client.patch(
            f"/api/repair-cases/{repair_id}", headers=headers, json=payload
        )
        assert response.status_code == 200, response.text


def _media_hashes(content: bytes) -> set[str]:
    with zipfile.ZipFile(io.BytesIO(content)) as package:
        return {
            hashlib.sha256(package.read(name)).hexdigest()
            for name in package.namelist()
            if name.startswith("word/media/")
        }


def test_invalid_transition_rolls_back_fields_and_returns_concrete_bulgarian_error(
    client, auth_headers, machine_ids, session_factory
):
    created = client.post(
        "/api/repair-cases",
        headers=auth_headers,
        json={
            "machine_id": machine_ids["9"],
            "reported_problem": "Първоначален контролен проблем",
        },
    )
    assert created.status_code == 201, created.text
    repair_id = created.json()["id"]

    rejected = client.patch(
        f"/api/repair-cases/{repair_id}",
        headers=auth_headers,
        json={
            "status": "DIAGNOSIS",
            "reported_problem": "Тази промяна трябва да бъде върната",
            "inspection_complete": True,
        },
    )
    assert rejected.status_code == 409, rejected.text
    detail = rejected.json()["detail"]
    assert detail["code"] == "repair_stage_requirements_missing"
    assert "състоянието при приемане" in detail["message"].lower()

    with session_factory() as session:
        repair = session.get(Repair, repair_id)
        assert repair.status == "ACCEPTED"
        assert repair.reported_problem == "Първоначален контролен проблем"
        assert repair.inspection_completed_at is None
        assert len(repair.events) == 1


def test_participants_parts_attachments_and_timeline_persist_without_duplicate_rows(
    client, auth_headers, machine_ids, session_factory
):
    repair_id = _create_repair(client, auth_headers, machine_ids["10"])
    with session_factory() as session:
        administrator_id = session.scalar(
            select(User.id).where(User.email == "admin@assetcore.local")
        )

    participant = client.post(
        f"/api/repair-cases/{repair_id}/participants",
        headers=auth_headers,
        json={
            "user_id": administrator_id,
            "contribution": "Контролна проверка на ремонта",
            "minutes_worked": 25,
        },
    )
    assert participant.status_code == 201, participant.text
    duplicate = client.post(
        f"/api/repair-cases/{repair_id}/participants",
        headers=auth_headers,
        json={"user_id": administrator_id, "minutes_worked": 25},
    )
    assert duplicate.status_code == 409, duplicate.text
    assert duplicate.json()["detail"]["code"] == "repair_participant_already_exists"

    second_participant = client.post(
        f"/api/repair-cases/{repair_id}/participants",
        headers=auth_headers,
        json={
            "full_name": "Иван Иванов Иванов",
            "job_title": "Електротехник",
            "contribution": "Електрическа диагностика",
            "minutes_worked": 40,
        },
    )
    assert second_participant.status_code == 201, second_participant.text

    catalog = client.get(
        f"/api/catalog/parts?verified_only=true&machine_id={machine_ids['10']}",
        headers=auth_headers,
    )
    assert catalog.status_code == 200, catalog.text
    verified_part = catalog.json()[0]
    part = client.post(
        f"/api/repair-cases/{repair_id}/parts",
        headers=auth_headers,
        json={
            "catalog_part_id": verified_part["id"],
            "description": verified_part["description"],
            "quantity": 2,
        },
    )
    assert part.status_code == 201, part.text
    assert part.json()["part_number"] == verified_part["part_number"]

    png = base64.b64encode(
        base64.b64decode(
            "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAAC0lEQVR42mNk+A8AAQUBAScY42YAAAAASUVORK5CYII="
        )
    ).decode()
    attachment = client.post(
        f"/api/repair-cases/{repair_id}/attachments",
        headers=auth_headers,
        json={
            "filename": "repair-control.png",
            "media_type": "image/png",
            "content_base64": png,
            "stage": "ACCEPTED",
            "description": "Контролно приложение",
        },
    )
    assert attachment.status_code == 201, attachment.text

    transitioned = client.patch(
        f"/api/repair-cases/{repair_id}",
        headers=auth_headers,
        json={"status": "DIAGNOSIS", "inspection_complete": True},
    )
    assert transitioned.status_code == 200, transitioned.text

    reloaded = client.get(
        f"/api/repair-cases/{repair_id}", headers=auth_headers
    )
    assert reloaded.status_code == 200, reloaded.text
    body = reloaded.json()
    assert len(body["participants"]) == 2
    assert body["participant_total_minutes"] == 65
    assert {item["minutes_worked"] for item in body["participants"]} == {25, 40}
    assert len(body["parts_used"]) == 1
    assert body["parts_used"][0]["quantity"] == 2
    assert body["parts_used"][0]["catalog_part_id"] == verified_part["id"]
    assert len(body["attachments"]) == 1
    assert {
        "PARTICIPANT_ADDED",
        "PART_ADDED",
        "ATTACHMENT_ADDED",
    }.issubset({event["event_type"] for event in body["events"]})

    with session_factory() as session:
        assert session.scalar(
            select(RepairParticipant).where(RepairParticipant.repair_id == repair_id)
        ) is not None
        assert session.scalar(
            select(RepairPart).where(RepairPart.repair_id == repair_id)
        ) is not None
        assert session.scalar(
            select(RepairAttachment).where(RepairAttachment.repair_id == repair_id)
        ) is not None


def test_concurrent_duplicate_participant_submission_creates_exactly_one_row(
    client, auth_headers, machine_ids, session_factory
):
    repair_id = _create_repair(client, auth_headers, machine_ids["11"])
    with session_factory() as session:
        administrator_id = session.scalar(
            select(User.id).where(User.email == "admin@assetcore.local")
        )
    payload = {"user_id": administrator_id, "minutes_worked": 30}

    def submit() -> int:
        return client.post(
            f"/api/repair-cases/{repair_id}/participants",
            headers=auth_headers,
            json=payload,
        ).status_code

    with ThreadPoolExecutor(max_workers=2) as executor:
        statuses = sorted(executor.map(lambda _: submit(), range(2)))
    assert statuses == [201, 409]
    with session_factory() as session:
        assert len(
            list(
                session.scalars(
                    select(RepairParticipant).where(
                        RepairParticipant.repair_id == repair_id
                    )
                )
            )
        ) == 1


def test_completion_is_atomic_generates_three_part_protocol_and_restores_availability(
    client,
    auth_headers,
    machine_ids,
    issue_payload,
    session_factory,
):
    repair_id = _create_repair(client, auth_headers, machine_ids["12"])
    participant = client.post(
        f"/api/repair-cases/{repair_id}/participants",
        headers=auth_headers,
        json={
            "full_name": "Контролен Участник Ремонт",
            "job_title": "Контролна длъжност",
            "contribution": "Контролна ремонтна операция",
            "minutes_worked": 80,
        },
    )
    assert participant.status_code == 201, participant.text
    catalog = client.get(
        f"/api/catalog/parts?verified_only=true&machine_id={machine_ids['12']}",
        headers=auth_headers,
    )
    assert catalog.status_code == 200, catalog.text
    verified_part = catalog.json()[0]
    part = client.post(
        f"/api/repair-cases/{repair_id}/parts",
        headers=auth_headers,
        json={
            "catalog_part_id": verified_part["id"],
            "description": verified_part["description"],
            "quantity": 1,
        },
    )
    assert part.status_code == 201, part.text
    _advance_to_final_stage(client, auth_headers, repair_id)

    completed = client.patch(
        f"/api/repair-cases/{repair_id}",
        headers=auth_headers,
        json={
            "status": "COMPLETED",
            "test_passed": True,
            "test_method": "Контролен функционален тест",
            "test_details": "Контролният тест е успешен",
            "functional_test_result": "Нормална работа в контролираната проверка",
            "result": "Ремонтната операция е завършена",
            "condition_after": "Готова за работа след успешен тест",
            "testing_minutes": 25,
        },
    )
    assert completed.status_code == 200, completed.text
    body = completed.json()
    assert body["status"] == "COMPLETED"
    assert body["total_work_minutes"] == 130
    assert body["participant_total_minutes"] == 80
    assert body["approved_by_id"] is not None
    assert {item["format"] for item in body["generated_documents"]} == {
        "docx",
        "pdf",
    }

    with session_factory() as session:
        repair = session.scalar(
            select(Repair).where(Repair.id == repair_id)
        )
        workshop = session.scalar(
            select(Location).where(Location.name == "Цех", Location.is_active.is_(True))
        )
        machine = session.get(Machine, machine_ids["12"])
        assert repair.closed_at is not None
        assert machine.status == "READY"
        assert machine.location_id == workshop.id
        documents = list(
            session.scalars(
                select(GeneratedDocument).where(
                    GeneratedDocument.repair_id == repair_id
                )
            )
        )
        assert len(documents) == 2
        assert {
            "COMPLETED",
            "DOCUMENT_GENERATED",
        }.issubset(
            set(
                session.scalars(
                    select(RepairEvent.event_type).where(
                        RepairEvent.repair_id == repair_id
                    )
                )
            )
        )
        assert session.scalar(
            select(AuditLog.id).where(
                AuditLog.entity_type == "repair",
                AuditLog.entity_id == repair_id,
            )
        ) is not None
        docx_record = next(item for item in documents if item.format == "docx")
        pdf_record = next(item for item in documents if item.format == "pdf")

    assert docx_record.snapshot["reported_problem"] == "Контролен проблем в изолираната тестова база"
    assert docx_record.snapshot["diagnostic_cleaning"] == "Извършено контролно почистване при диагностиката"
    assert docx_record.snapshot["test_method"] == "Контролен функционален тест"
    assert docx_record.snapshot["participant_total_minutes"] == 80
    assert docx_record.snapshot["participants"][0]["minutes_worked"] == 80
    assert docx_record.snapshot["parts_used"][0]["part_number"] == verified_part["part_number"]

    document = Document(io.BytesIO(docx_record.content))
    docx_text = "\n".join(paragraph.text for paragraph in document.paragraphs)
    docx_text += "\n" + "\n".join(
        cell.text for table in document.tables for row in table.rows for cell in row.cells
    )
    for required_text in (
        "ПРОТОКОЛ ЗА ПРИЕМАНЕ НА ОБОРУДВАНЕ ЗА РЕМОНТ",
        "ПРОТОКОЛ ЗА ИЗВЪРШЕН РЕМОНТ",
        "Описание на реално отработеното време",
        "Контролният тест е успешен",
        "Извършено контролно почистване при диагностиката",
        "Контролен Участник Ремонт",
        "1 ч 20 мин",
        "35 мин",
        "1 ч 10 мин",
        "25 мин",
        verified_part["part_number"],
    ):
        assert required_text in docx_text
    with zipfile.ZipFile(io.BytesIO(docx_record.content)) as package:
        document_xml = package.read("word/document.xml")
    assert document_xml.count(b'w:type="page"') >= 2

    approved_header = (
        ROOT / "backend" / "resources" / "templates" / "transfer_issue-bg-v3.docx"
    ).read_bytes()
    assert _media_hashes(approved_header) == _media_hashes(docx_record.content)

    pdf = PdfReader(io.BytesIO(pdf_record.content))
    assert len(pdf.pages) == 3
    pdf_text = "\n".join(page.extract_text() or "" for page in pdf.pages)
    assert "Контролният тест е успешен" in pdf_text
    assert "12" in pdf_text

    issue = client.post(
        "/api/transfers/bulk-issue",
        headers=auth_headers,
        json=issue_payload(machine_ids["12"]),
    )
    assert issue.status_code == 201, issue.text


def test_repair_duration_validation_is_enforced_before_database_write(
    client, auth_headers, machine_ids, session_factory
):
    repair_id = _create_repair(client, auth_headers, machine_ids["13"])
    rejected = client.patch(
        f"/api/repair-cases/{repair_id}",
        headers=auth_headers,
        json={"diagnosis_minutes": -1},
    )
    assert rejected.status_code == 422, rejected.text
    with session_factory() as session:
        assert session.get(Repair, repair_id).diagnosis_minutes is None


def test_four_active_stages_reject_retired_statuses_and_final_draft_stays_repairing(
    client, auth_headers, machine_ids
):
    repair_id = _create_repair(client, auth_headers, machine_ids["14"])
    diagnosis = client.patch(
        f"/api/repair-cases/{repair_id}",
        headers=auth_headers,
        json={"status": "DIAGNOSIS"},
    )
    assert diagnosis.status_code == 200, diagnosis.text
    retired = client.patch(
        f"/api/repair-cases/{repair_id}",
        headers=auth_headers,
        json={"status": "WAITING_PARTS"},
    )
    assert retired.status_code == 409, retired.text
    assert retired.json()["detail"]["code"] == "invalid_repair_status_transition"

    repairing = client.patch(
        f"/api/repair-cases/{repair_id}",
        headers=auth_headers,
        json={
            "status": "REPAIRING",
            "diagnosis": "Контролна четиристепенна диагностика",
            "required_work": "Контролна необходима работа",
            "diagnosis_minutes": 30,
            "diagnostic_cleaning": "Контролно диагностично почистване",
        },
    )
    assert repairing.status_code == 200, repairing.text
    final_draft = client.patch(
        f"/api/repair-cases/{repair_id}",
        headers=auth_headers,
        json={
            "work_performed": "Контролно извършена работа",
            "repair_minutes": 45,
            "test_passed": True,
            "test_method": "Контролен метод",
            "test_details": "Контролният тест е успешен",
            "testing_minutes": 15,
            "condition_after": "Контролно крайно състояние",
            "result": "Контролен краен резултат",
        },
    )
    assert final_draft.status_code == 200, final_draft.text
    assert final_draft.json()["status"] == "REPAIRING"
    assert not final_draft.json()["generated_documents"]


def test_final_step_opens_only_after_explicit_advance_and_keeps_repairing_status(
    client, auth_headers, machine_ids
):
    repair_id = _create_repair(client, auth_headers, machine_ids["18"])
    for payload in (
        {"status": "DIAGNOSIS"},
        {
            "status": "REPAIRING",
            "diagnosis": "Контролна диагностика",
            "required_work": "Контролен ремонт",
            "diagnosis_minutes": 20,
        },
    ):
        response = client.patch(
            f"/api/repair-cases/{repair_id}", headers=auth_headers, json=payload
        )
        assert response.status_code == 200, response.text

    saved = client.patch(
        f"/api/repair-cases/{repair_id}",
        headers=auth_headers,
        json={"work_performed": "Контролна работа", "repair_minutes": 25},
    )
    assert saved.status_code == 200, saved.text
    assert saved.json()["status"] == "REPAIRING"
    assert not any(
        (event.get("structured_data") or {}).get("wizard_stage") == "COMPLETION"
        for event in saved.json()["events"]
    )

    advanced = client.patch(
        f"/api/repair-cases/{repair_id}",
        headers=auth_headers,
        json={"advance_to_final": True},
    )
    assert advanced.status_code == 200, advanced.text
    assert advanced.json()["status"] == "REPAIRING"
    assert any(
        (event.get("structured_data") or {}).get("wizard_stage") == "COMPLETION"
        for event in advanced.json()["events"]
    )


def test_editing_previous_stage_preserves_all_later_stage_data(
    client, auth_headers, machine_ids
):
    repair_id = _create_repair(client, auth_headers, machine_ids["15"])
    _advance_to_final_stage(client, auth_headers, repair_id)
    final_values = {
        "test_passed": True,
        "test_method": "Контролен метод за запазване",
        "test_details": "Контролни тестови подробности",
        "testing_minutes": 18,
        "condition_after": "Контролно състояние след ремонта",
        "result": "Контролен резултат след ремонта",
    }
    saved = client.patch(
        f"/api/repair-cases/{repair_id}", headers=auth_headers, json=final_values
    )
    assert saved.status_code == 200, saved.text
    edited = client.patch(
        f"/api/repair-cases/{repair_id}",
        headers=auth_headers,
        json={"diagnosis": "Редактирана контролна диагноза"},
    )
    assert edited.status_code == 200, edited.text
    reloaded = client.get(
        f"/api/repair-cases/{repair_id}", headers=auth_headers
    ).json()
    assert reloaded["diagnosis"] == "Редактирана контролна диагноза"
    for key, value in final_values.items():
        assert reloaded[key] == value
    assert reloaded["work_performed"] == "Изпълнена и проверена ремонтна операция"
    assert reloaded["repair_minutes"] == 70


def test_invalid_final_completion_rolls_back_status_machine_and_documents(
    client, auth_headers, machine_ids, session_factory
):
    repair_id = _create_repair(client, auth_headers, machine_ids["16"])
    _advance_to_final_stage(client, auth_headers, repair_id)
    rejected = client.patch(
        f"/api/repair-cases/{repair_id}",
        headers=auth_headers,
        json={
            "status": "COMPLETED",
            "test_passed": True,
            "test_details": "Има подробности, но липсва метод",
            "testing_minutes": 12,
            "condition_after": "Контролно крайно състояние",
            "result": "Контролен резултат",
        },
    )
    assert rejected.status_code == 409, rejected.text
    assert "test_method" in rejected.json()["detail"]["missing_fields"]
    with session_factory() as session:
        repair = session.get(Repair, repair_id)
        assert repair.status == "REPAIRING"
        assert repair.machine.status == "REPAIR"
        assert not list(
            session.scalars(
                select(GeneratedDocument).where(
                    GeneratedDocument.repair_id == repair_id
                )
            )
        )


def test_repair_part_rejects_verified_catalog_item_for_wrong_machine(
    client, auth_headers, machine_ids, session_factory
):
    repair_id = _create_repair(client, auth_headers, machine_ids["9"])
    with session_factory() as session:
        part = next((
            item
            for item in session.scalars(
                select(PartCatalog).where(PartCatalog.is_verified.is_(True))
            )
            if "9" not in {str(value) for value in (item.compatible_machine_numbers or [])}
        ), None)
        assert part is not None
        part_id = part.id
        description = part.description
    rejected = client.post(
        f"/api/repair-cases/{repair_id}/parts",
        headers=auth_headers,
        json={"catalog_part_id": part_id, "description": description, "quantity": 1},
    )
    assert rejected.status_code == 409, rejected.text
    assert rejected.json()["detail"]["code"] == "catalog_part_not_compatible_with_machine"
