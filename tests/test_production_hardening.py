from __future__ import annotations

import base64
import io
import tarfile
from datetime import timedelta

import pytest
from app.licensing import RIGHTSHOLDER, canonical_payload, evaluate_license
from app.models import (
    AuditLog,
    DocumentSignature,
    EmergencyAccessSession,
    InstallationOwnership,
    OfficialDocumentVersion,
    SoftwareLicense,
    User,
    UserRole,
    utcnow,
)
from app.security import hash_password
from app.seed import seed_database
from app.settings import Settings, settings
from cryptography.hazmat.primitives import serialization
from cryptography.hazmat.primitives.asymmetric.ed25519 import Ed25519PrivateKey
from docx import Document
from PIL import Image, ImageDraw
from reportlab.lib.pagesizes import A4
from reportlab.pdfgen import canvas
from sqlalchemy import func, select

from scripts.restore_assetcore import _safe_extract


def _login(client, email: str, password: str) -> dict[str, str]:
    response = client.post("/api/auth/login", json={"email": email, "password": password})
    assert response.status_code == 200, response.text
    return {"Authorization": f"Bearer {response.json()['access_token']}"}


def test_production_restart_does_not_require_reusing_bootstrap_password():
    configured = Settings(
        _env_file=None,
        production_mode=True,
        database_url="postgresql+psycopg://assetcore@database/assetcore",
        secret_key="release-test-secret-that-is-not-a-real-credential",
        owner_email="owner@example.invalid",
        owner_initial_password=None,
        signature_encryption_key="release-test-encryption-key",
        license_enforcement_enabled=False,
    )
    assert configured.owner_initial_password is None


def _licence_payload(*, valid_from, valid_until, max_users: int = 20, max_assets: int = 50) -> dict:
    return {
        "license_id": "TEST-LICENCE-001",
        "rightsholder": RIGHTSHOLDER,
        "client_name": "AssetCore automated test",
        "installation_id": "assetcore-test-installation",
        "modules": ["asset-management", "documents", "signatures"],
        "max_users": max_users,
        "max_assets": max_assets,
        "valid_from": valid_from.isoformat() + "Z",
        "valid_until": valid_until.isoformat() + "Z",
        "license_type": "ANNUAL",
        "environment": "development",
        "allowed_domains": [],
        "max_installations": 1,
        "grace_days": 3,
        "version": 1,
    }


def _signed_envelope(monkeypatch, payload: dict) -> dict:
    private_key = Ed25519PrivateKey.generate()
    public_key = private_key.public_key().public_bytes(
        serialization.Encoding.Raw,
        serialization.PublicFormat.Raw,
    )
    monkeypatch.setattr(settings, "license_public_key", base64.b64encode(public_key).decode())
    monkeypatch.setattr(settings, "installation_id", "assetcore-test-installation")
    monkeypatch.setattr(settings, "deployment_environment", "development")
    return {
        "payload": payload,
        "signature": base64.b64encode(private_key.sign(canonical_payload(payload))).decode(),
    }


def _external_signer(client, headers: dict[str, str], suffix: str) -> dict:
    response = client.post(
        "/api/external-signers",
        headers=headers,
        json={
            "first_name": "QA",
            "middle_name": "Test",
            "last_name": f"Signer-{suffix}",
            "job_title": "Test participant",
            "company": "Automated test",
            "participant_role": "Test signature participant",
            "note": "Record used only in an isolated automated test database.",
        },
    )
    assert response.status_code == 201, response.text
    return response.json()


def _document_files() -> tuple[str, str]:
    document = Document()
    document.add_heading("AssetCore QA protocol", level=1)
    document.add_paragraph("НЕПЪЛНО ПОДПИСАН")
    docx_output = io.BytesIO()
    document.save(docx_output)

    pdf_output = io.BytesIO()
    pdf = canvas.Canvas(pdf_output, pagesize=A4)
    pdf.drawString(50, 790, "AssetCore QA protocol")
    pdf.save()
    return (
        base64.b64encode(docx_output.getvalue()).decode(),
        base64.b64encode(pdf_output.getvalue()).decode(),
    )


def _signature_payload(consent_text: str, variant: int) -> dict:
    output = io.BytesIO()
    image = Image.new("RGBA", (320, 120), "white")
    ImageDraw.Draw(image).line(
        [(15, 80), (70, 30 + variant), (135, 88), (205, 25), (300, 75)],
        fill="black",
        width=5,
    )
    image.save(output, format="PNG")
    points = [
        {"x": 20 + index * 20, "y": 30 + (index % 2) * 25, "t": index * 10, "pressure": 0.5}
        for index in range(8)
    ]
    return {
        "consent_accepted": True,
        "consent_text": consent_text,
        "strokes": [points],
        "image_base64": base64.b64encode(output.getvalue()).decode(),
        "canvas_width": 320,
        "canvas_height": 120,
    }


def test_incomplete_profile_requires_full_identity_or_documented_exception(
    client, auth_headers, session_factory
):
    with session_factory() as db:
        user = User(
            email="profile-test@example.invalid",
            full_name="Incomplete test profile",
            password_hash=hash_password("StrongPass123!"),
            role=UserRole.OBSERVER.value,
            profile_status="PROFILE_INCOMPLETE",
        )
        db.add(user)
        db.commit()
        user_id = user.id
    headers = _login(client, "profile-test@example.invalid", "StrongPass123!")

    invalid = client.put(
        "/api/users/me/profile",
        headers=headers,
        json={"first_name": "QA", "last_name": "Profile", "job_title": "Test observer"},
    )
    assert invalid.status_code == 422

    self_exception = client.put(
        "/api/users/me/profile",
        headers=headers,
        json={
            "first_name": "QA",
            "middle_name": None,
            "last_name": "Profile",
            "job_title": "Test observer",
            "legal_name_exception": True,
            "legal_name_exception_reason": "Automated test of the documented legal-name exception.",
            "preferred_language": "ru",
        },
    )
    assert self_exception.status_code == 403
    assert self_exception.json()["detail"]["code"] == "legal_name_exception_requires_admin"

    completed = client.put(
        f"/api/users/{user_id}/profile",
        headers=auth_headers,
        json={
            "first_name": "QA",
            "middle_name": None,
            "last_name": "Profile",
            "job_title": "Test observer",
            "legal_name_exception": True,
            "legal_name_exception_reason": "Automated test of the documented legal-name exception.",
            "preferred_language": "ru",
        },
    )
    assert completed.status_code == 200, completed.text
    assert completed.json()["profile_status"] == "PROFILE_COMPLETE"
    with session_factory() as db:
        stored = db.get(User, user_id)
        assert stored.full_name == "QA Profile"
        assert stored.preferred_language == "ru"
        assert stored.legal_name_exception_approved_by_id is not None
        assert stored.legal_name_exception_approved_at is not None
        assert db.scalar(select(func.count(AuditLog.id)).where(AuditLog.entity_type == "user_profile")) == 2


def test_owner_transfer_is_protected_and_survives_reseeding(
    client, auth_headers, session_factory
):
    with session_factory() as db:
        target = User(
            email="next-owner@example.invalid",
            full_name="QA Next Owner",
            first_name="QA",
            middle_name="Next",
            last_name="Owner",
            job_title="Test administrator",
            profile_status="PROFILE_COMPLETE",
            password_hash=hash_password("StrongPass123!"),
            role=UserRole.ADMINISTRATOR.value,
            is_active=True,
        )
        db.add(target)
        db.commit()
        target_id = target.id

    rejected = client.post(
        "/api/owner/transfer",
        headers=auth_headers,
        json={
            "target_user_id": target_id,
            "current_password": "wrong-password",
            "reason": "Automated verification of protected ownership transfer.",
        },
    )
    assert rejected.status_code == 403

    transferred = client.post(
        "/api/owner/transfer",
        headers=auth_headers,
        json={
            "target_user_id": target_id,
            "current_password": "AssetCore123!",
            "reason": "Automated verification of protected ownership transfer.",
        },
    )
    assert transferred.status_code == 200, transferred.text
    assert transferred.json()["owner_user_id"] == target_id

    with session_factory() as db:
        seed_database(db)
        ownership = db.scalar(select(InstallationOwnership))
        assert ownership.owner_user_id == target_id
        assert db.scalar(select(func.count(User.id)).where(User.is_system_owner.is_(True))) == 1
        assert db.get(User, target_id).is_system_owner is True


def test_emergency_access_is_owner_only_time_limited_and_audited(
    client, auth_headers, session_factory
):
    with session_factory() as db:
        observer = User(
            email="emergency-observer@example.invalid",
            full_name="QA Emergency Observer",
            first_name="QA",
            middle_name="Emergency",
            last_name="Observer",
            job_title="Test observer",
            profile_status="PROFILE_COMPLETE",
            password_hash=hash_password("StrongPass123!"),
            role=UserRole.OBSERVER.value,
            is_active=True,
        )
        db.add(observer)
        db.commit()
    observer_headers = _login(
        client, "emergency-observer@example.invalid", "StrongPass123!"
    )
    payload = {
        "current_password": "StrongPass123!",
        "reason": "Automated verification of the controlled emergency procedure.",
        "duration_minutes": 5,
    }
    forbidden = client.post(
        "/api/emergency-access/start", headers=observer_headers, json=payload
    )
    assert forbidden.status_code == 403

    rejected = client.post(
        "/api/emergency-access/start",
        headers=auth_headers,
        json={**payload, "current_password": "wrong-password"},
    )
    assert rejected.status_code == 403
    assert rejected.json()["detail"]["code"] == "reauthentication_failed"

    started = client.post(
        "/api/emergency-access/start",
        headers=auth_headers,
        json={**payload, "current_password": "AssetCore123!"},
    )
    assert started.status_code == 201, started.text
    assert started.json()["active"] is True
    assert started.json()["mfa_verified"] is False
    session_id = started.json()["session_id"]

    duplicate = client.post(
        "/api/emergency-access/start",
        headers=auth_headers,
        json={**payload, "current_password": "AssetCore123!"},
    )
    assert duplicate.status_code == 409

    visible = client.get("/api/emergency-access/status", headers=observer_headers)
    assert visible.status_code == 200
    assert visible.json()["active"] is True
    assert visible.json()["session_id"] == session_id

    ended = client.post(
        f"/api/emergency-access/{session_id}/end",
        headers=auth_headers,
        json={
            "current_password": "AssetCore123!",
            "reason": "Automated verification completed; close the procedure.",
        },
    )
    assert ended.status_code == 200, ended.text
    assert ended.json()["active"] is False
    with session_factory() as db:
        stored = db.get(EmergencyAccessSession, session_id)
        assert stored.ended_at is not None
        assert stored.end_reason
        assert (
            db.scalar(
                select(func.count(AuditLog.id)).where(
                    AuditLog.entity_type == "emergency_access"
                )
            )
            >= 4
        )


def test_offline_ed25519_licence_active_grace_read_only_and_tamper_detection(
    client, auth_headers, session_factory, monkeypatch
):
    now = utcnow()
    payload = _licence_payload(valid_from=now - timedelta(days=1), valid_until=now + timedelta(days=2))
    envelope = _signed_envelope(monkeypatch, payload)
    installed = client.post("/api/license/install", headers=auth_headers, json=envelope)
    assert installed.status_code == 200, installed.text
    assert installed.json()["state"] == "ACTIVE"

    with session_factory() as db:
        assert evaluate_license(db, now=now + timedelta(days=3)).state == "GRACE_PERIOD"
        expired = evaluate_license(db, now=now + timedelta(days=6))
        assert expired.state == "READ_ONLY"
        assert expired.read_only is True
        item = db.scalar(select(SoftwareLicense))
        item.payload = {**item.payload, "max_users": 999}
        db.commit()
        assert evaluate_license(db).state == "INVALID"


def test_licence_rejects_invalid_signature_and_capacity_below_verified_inventory(
    client, auth_headers, session_factory, monkeypatch
):
    now = utcnow()
    payload = _licence_payload(
        valid_from=now - timedelta(days=1),
        valid_until=now + timedelta(days=1),
        max_assets=1,
    )
    envelope = _signed_envelope(monkeypatch, payload)
    invalid = client.post(
        "/api/license/install",
        headers=auth_headers,
        json={**envelope, "signature": base64.b64encode(b"x" * 64).decode()},
    )
    assert invalid.status_code == 409
    assert invalid.json()["detail"]["code"] == "invalid_license_signature"

    capacity = client.post("/api/license/install", headers=auth_headers, json=envelope)
    assert capacity.status_code == 409
    assert capacity.json()["detail"]["code"] == "license_capacity_below_current_usage"
    with session_factory() as db:
        assert db.scalar(select(func.count(SoftwareLicense.id))) == 0
        assert db.scalar(select(func.count(AuditLog.id)).where(AuditLog.entity_type == "software_license")) == 2


def test_sequential_mobile_signatures_finalize_immutable_version_and_allow_superseding(
    client, auth_headers, session_factory
):
    first = _external_signer(client, auth_headers, "A")
    second = _external_signer(client, auth_headers, "B")
    docx, pdf = _document_files()
    created = client.post(
        "/api/official-documents",
        headers=auth_headers,
        json={
            "document_number": "QA-SIGN-001",
            "document_type": "PART_REQUEST",
            "language": "bg",
            "snapshot": {"purpose": "isolated automated signature test"},
            "docx_base64": docx,
            "pdf_base64": pdf,
            "participants": [
                {"slot_code": "REQUESTED_BY", "operation_role": "Requested by", "external_signer_id": first["id"]},
                {"slot_code": "APPROVED_BY", "operation_role": "Approved by", "external_signer_id": second["id"]},
            ],
        },
    )
    assert created.status_code == 201, created.text
    body = created.json()
    participants = {item["slot_code"]: item for item in body["participants"]}

    blocked = client.post(
        "/api/signatures/sessions",
        headers=auth_headers,
        json={"participant_id": participants["APPROVED_BY"]["id"]},
    )
    assert blocked.status_code == 409
    assert blocked.json()["detail"]["code"] == "signature_sequence_blocked"

    first_session = client.post(
        "/api/signatures/sessions",
        headers=auth_headers,
        json={"participant_id": participants["REQUESTED_BY"]["id"]},
    )
    assert first_session.status_code == 201
    first_token = first_session.json()["signing_token"]
    first_summary = client.get(f"/api/signing/{first_token}")
    assert first_summary.status_code == 200
    assert client.post(
        f"/api/signing/{first_token}",
        json=_signature_payload(first_summary.json()["consent_notice"], 1),
    ).status_code == 201
    first_confirmed = client.post(f"/api/signing/{first_token}/confirm")
    assert first_confirmed.status_code == 200, first_confirmed.text
    assert first_confirmed.json()["document_status"] == "PARTIALLY_SIGNED"
    assert client.post(f"/api/signing/{first_token}/confirm").status_code == 410

    second_session = client.post(
        "/api/signatures/sessions",
        headers=auth_headers,
        json={"participant_id": participants["APPROVED_BY"]["id"]},
    )
    second_token = second_session.json()["signing_token"]
    second_summary = client.get(f"/api/signing/{second_token}")
    assert client.post(
        f"/api/signing/{second_token}",
        json=_signature_payload(second_summary.json()["consent_notice"], 2),
    ).status_code == 201
    second_confirmed = client.post(f"/api/signing/{second_token}/confirm")
    assert second_confirmed.status_code == 200, second_confirmed.text
    assert second_confirmed.json()["document_status"] == "SIGNED"

    signed = client.get(f"/api/official-documents/{body['id']}", headers=auth_headers).json()
    assert signed["signed_count"] == signed["required_count"] == 2
    assert signed["current_version"]["docx_sha256"] != body["current_version"]["docx_sha256"]
    assert signed["current_version"]["pdf_sha256"] != body["current_version"]["pdf_sha256"]
    assert client.get(
        f"/api/official-documents/{body['id']}/versions/1/download/docx",
        headers=auth_headers,
    ).status_code == 200

    old_docx_hash = signed["current_version"]["docx_sha256"]
    old_pdf_hash = signed["current_version"]["pdf_sha256"]
    superseded = client.post(
        f"/api/official-documents/{body['id']}/supersede",
        headers=auth_headers,
        json={
            "reason": "Automated correction creates a new immutable version.",
            "snapshot": {"purpose": "isolated automated correction test"},
            "docx_base64": docx,
            "pdf_base64": pdf,
            "participants": [],
        },
    )
    assert superseded.status_code == 201, superseded.text
    assert superseded.json()["current_version"]["version"] == 2
    with session_factory() as db:
        old = db.scalar(
            select(OfficialDocumentVersion).where(
                OfficialDocumentVersion.document_id == body["id"],
                OfficialDocumentVersion.version == 1,
            )
        )
        assert old.status == "SUPERSEDED"
        assert old.docx_sha256 == old_docx_hash
        assert old.pdf_sha256 == old_pdf_hash
        assert db.scalar(select(func.count(DocumentSignature.id))) == 2


def test_restore_rejects_archive_path_traversal(tmp_path):
    archive_path = tmp_path / "unsafe.tar"
    with tarfile.open(archive_path, "w") as bundle:
        payload = b"unsafe"
        member = tarfile.TarInfo("../outside.txt")
        member.size = len(payload)
        bundle.addfile(member, io.BytesIO(payload))
    with tarfile.open(archive_path, "r") as bundle:
        with pytest.raises(SystemExit, match="unsafe path"):
            _safe_extract(bundle, tmp_path / "destination")
    assert not (tmp_path / "outside.txt").exists()
