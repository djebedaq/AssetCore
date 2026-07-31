from __future__ import annotations

import io
import json
import zipfile
from concurrent.futures import ThreadPoolExecutor
from threading import Barrier

from app.models import (
    AuditLog,
    Machine,
    ProtocolDocument,
    TransferBatch,
    TransferProtocol,
)
from docx import Document
from fastapi.testclient import TestClient
from sqlalchemy import func, inspect, select
from sqlalchemy.dialects import postgresql
from sqlalchemy.schema import CreateIndex


def issue(client: TestClient, headers: dict, payload: dict):
    return client.post("/api/transfers/bulk-issue", headers=headers, json=payload)


def return_payload(*transfers: dict) -> dict:
    return {
        "items": [
            {
                "transfer_id": transfer["transfer_id"],
                "machine_id": transfer["machine_id"],
                "condition_text": f"Състояние за тест №{transfer['machine_number']}",
                "result_text": "Насочена към преглед",
                "notes": "",
                "returned_by": "",
                "accepted_by": "",
                "next_status": "INSPECTION",
            }
            for transfer in transfers
        ]
    }


def test_issue_one_available_machine_successfully(
    client, auth_headers, machine_ids, issue_payload
):
    response = issue(client, auth_headers, issue_payload(machine_ids["4"]))
    assert response.status_code == 201, response.text
    body = response.json()
    assert body["batch_reference"].startswith("HPWJ-B-")
    assert [item["machine_number"] for item in body["transfers"]] == ["4"]
    assert len(body["transfers"][0]["documents"]) == 2


def test_reject_already_issued_machine_with_structured_bulgarian_conflict(
    client, auth_headers, machine_ids, issue_payload
):
    first = issue(client, auth_headers, issue_payload(machine_ids["7"]))
    second = issue(client, auth_headers, issue_payload(machine_ids["7"]))
    assert first.status_code == 201
    assert second.status_code == 409
    detail = second.json()["detail"]
    assert detail["code"] == "issue_conflict"
    assert "Машина №7" in detail["message"]
    assert detail["conflicts"][0]["protocol_number"].startswith("HPWJ-")
    assert detail["conflicts"][0]["status"] == "ISSUED"
    assert detail["conflicts"][0]["status_label"] == "Издадена"


def test_bulk_issue_is_atomic_when_all_machines_are_available(
    client, auth_headers, machine_ids, issue_payload, session_factory
):
    response = issue(
        client,
        auth_headers,
        issue_payload(machine_ids["4"], machine_ids["5"], machine_ids["7"]),
    )
    assert response.status_code == 201
    with session_factory() as session:
        assert session.scalar(select(func.count(TransferProtocol.id))) == 3
        statuses = session.scalars(
            select(Machine.status).where(
                Machine.id.in_([machine_ids["4"], machine_ids["5"], machine_ids["7"]])
            )
        ).all()
        assert statuses == ["ISSUED", "ISSUED", "ISSUED"]


def test_bulk_issue_rejects_everything_if_one_machine_is_unavailable(
    client, auth_headers, machine_ids, issue_payload, session_factory
):
    issued = issue(client, auth_headers, issue_payload(machine_ids["7"]))
    assert issued.status_code == 201
    rejected = issue(
        client,
        auth_headers,
        issue_payload(machine_ids["4"], machine_ids["7"]),
    )
    assert rejected.status_code == 409
    with session_factory() as session:
        assert session.scalar(select(func.count(TransferProtocol.id))) == 1
        assert session.get(Machine, machine_ids["4"]).status == "READY"


def test_duplicate_machine_identifiers_are_rejected_before_writes(
    client, auth_headers, machine_ids, issue_payload, session_factory
):
    response = issue(
        client,
        auth_headers,
        issue_payload(machine_ids["4"], machine_ids["4"]),
    )
    assert response.status_code == 422
    with session_factory() as session:
        assert session.scalar(select(func.count(TransferProtocol.id))) == 0


def test_two_simultaneous_issue_attempts_create_only_one_active_transfer(
    client, auth_headers, machine_ids, issue_payload, session_factory
):
    barrier = Barrier(2)

    def attempt():
        barrier.wait()
        return issue(client, auth_headers, issue_payload(machine_ids["4"])).status_code

    with ThreadPoolExecutor(max_workers=2) as executor:
        statuses = sorted(executor.map(lambda _: attempt(), range(2)))
    assert statuses == [201, 409]
    with session_factory() as session:
        active = session.scalar(
            select(func.count(TransferProtocol.id)).where(
                TransferProtocol.machine_id == machine_ids["4"],
                TransferProtocol.is_active.is_(True),
            )
        )
        assert active == 1


def test_full_batch_return_closes_every_individual_transfer(
    client, auth_headers, machine_ids, issue_payload
):
    created = issue(
        client, auth_headers, issue_payload(machine_ids["4"], machine_ids["5"])
    ).json()
    response = client.post(
        "/api/transfers/bulk-return",
        headers=auth_headers,
        json=return_payload(*created["transfers"]),
    )
    assert response.status_code == 200, response.text
    progress = response.json()["batches"][0]
    assert progress["returned_machines"] == 2
    assert progress["still_issued_machines"] == 0
    assert progress["status"] == "RETURNED"


def test_partial_batch_return_keeps_remaining_machine_issued(
    client, auth_headers, machine_ids, issue_payload, session_factory
):
    created = issue(
        client,
        auth_headers,
        issue_payload(machine_ids["4"], machine_ids["5"], machine_ids["7"]),
    ).json()
    response = client.post(
        "/api/transfers/bulk-return",
        headers=auth_headers,
        json=return_payload(created["transfers"][0]),
    )
    assert response.status_code == 200
    progress = response.json()["batches"][0]
    assert progress["status"] == "PARTIALLY_RETURNED"
    assert progress["returned_machines"] == 1
    assert progress["still_issued_machines"] == 2
    with session_factory() as session:
        assert session.get(Machine, created["transfers"][0]["machine_id"]).status == "INSPECTION"
        assert session.get(Machine, created["transfers"][1]["machine_id"]).status == "ISSUED"


def test_mixed_batch_return_updates_each_batch_and_scopes_its_audit(
    client, auth_headers, machine_ids, issue_payload, session_factory
):
    first = issue(client, auth_headers, issue_payload(machine_ids["4"])).json()
    second = issue(client, auth_headers, issue_payload(machine_ids["5"])).json()
    first_transfer = first["transfers"][0]
    second_transfer = second["transfers"][0]

    response = client.post(
        "/api/transfers/bulk-return",
        headers=auth_headers,
        json=return_payload(first_transfer, second_transfer),
    )
    assert response.status_code == 200, response.text
    assert {batch["batch_id"] for batch in response.json()["batches"]} == {
        first["batch_id"],
        second["batch_id"],
    }
    assert all(batch["still_issued_machines"] == 0 for batch in response.json()["batches"])

    expected = {
        first["batch_reference"]: (
            first_transfer["transfer_id"],
            first_transfer["machine_number"],
        ),
        second["batch_reference"]: (
            second_transfer["transfer_id"],
            second_transfer["machine_number"],
        ),
    }
    with session_factory() as session:
        logs = session.scalars(
            select(AuditLog).where(
                AuditLog.entity_type == "transfer_batch",
                AuditLog.action == "Актуализирано връщане на партида",
            )
        ).all()
        assert len(logs) == 2
        for log in logs:
            details = json.loads(log.details)
            transfer_id, machine_number = expected[log.operation_reference]
            assert details["returned_transfer_ids"] == [transfer_id]
            assert details["returned_machine_numbers"] == [machine_number]


def test_return_without_active_issue_and_double_return_are_rejected(
    client, auth_headers, machine_ids, issue_payload
):
    created = issue(client, auth_headers, issue_payload(machine_ids["4"])).json()
    payload = return_payload(created["transfers"][0])
    assert client.post("/api/transfers/bulk-return", headers=auth_headers, json=payload).status_code == 200
    repeated = client.post("/api/transfers/bulk-return", headers=auth_headers, json=payload)
    assert repeated.status_code == 409
    assert repeated.json()["detail"]["code"] == "return_conflict"


def test_return_through_wrong_transfer_is_rejected_atomically(
    client, auth_headers, machine_ids, issue_payload, session_factory
):
    created = issue(
        client, auth_headers, issue_payload(machine_ids["4"], machine_ids["5"])
    ).json()["transfers"]
    wrong = return_payload(created[0])
    wrong["items"][0]["machine_id"] = created[1]["machine_id"]
    response = client.post("/api/transfers/bulk-return", headers=auth_headers, json=wrong)
    assert response.status_code == 409
    with session_factory() as session:
        assert session.scalar(
            select(func.count(TransferProtocol.id)).where(TransferProtocol.is_active.is_(True))
        ) == 2


def test_audit_records_success_and_rejected_conflict(
    client, auth_headers, machine_ids, issue_payload, session_factory
):
    issue(client, auth_headers, issue_payload(machine_ids["4"]))
    issue(client, auth_headers, issue_payload(machine_ids["4"]))
    with session_factory() as session:
        actions = session.scalars(select(AuditLog.action).order_by(AuditLog.id)).all()
        assert "Издадена машина" in actions
        assert "Групово издаване" in actions
        assert "Отказано групово издаване" in actions
        rejected = session.scalar(
            select(AuditLog).where(AuditLog.action == "Отказано групово издаване")
        )
        assert rejected.user_name == "Администратор"
        assert "конфликти" in rejected.details


def test_one_immutable_docx_and_pdf_are_stored_per_machine(
    client, auth_headers, machine_ids, issue_payload, session_factory
):
    created = issue(
        client, auth_headers, issue_payload(machine_ids["4"], machine_ids["5"])
    ).json()
    assert all(
        sorted(document["format"] for document in transfer["documents"]) == ["docx", "pdf"]
        for transfer in created["transfers"]
    )
    with session_factory() as session:
        documents = session.scalars(select(ProtocolDocument)).all()
        assert len(documents) == 4
        assert all(document.content and len(document.sha256) == 64 for document in documents)
        docx = next(document for document in documents if document.format == "docx")
        parsed = Document(io.BytesIO(docx.content))
        section = parsed.sections[0]
        assert round(section.page_width.mm) == 210
        assert round(section.page_height.mm) == 297
        all_text = "\n".join(
            [paragraph.text for paragraph in parsed.paragraphs]
            + [cell.text for table in parsed.tables for row in table.rows for cell in row.cells]
        )
        assert created["batch_reference"] in all_text


def test_bulk_zip_exposes_every_generated_protocol(
    client, auth_headers, machine_ids, issue_payload
):
    created = issue(
        client, auth_headers, issue_payload(machine_ids["4"], machine_ids["5"])
    ).json()
    response = client.get(created["zip_download_endpoint"], headers=auth_headers)
    assert response.status_code == 200
    with zipfile.ZipFile(io.BytesIO(response.content)) as archive:
        names = archive.namelist()
        assert len(names) == 4
        assert len(names) == len(set(names))
        assert all("/" not in name and "\\" not in name for name in names)


def test_document_and_batch_references_are_deterministic_and_traceable(
    client, auth_headers, machine_ids, issue_payload
):
    created = issue(
        client, auth_headers, issue_payload(machine_ids["4"], machine_ids["5"])
    ).json()
    assert created["batch_reference"].endswith(f"-{created['batch_id']:06d}")
    for transfer in created["transfers"]:
        assert transfer["protocol_number"].endswith(f"-{transfer['transfer_id']:06d}")


def test_document_generation_failure_rolls_back_entire_issue(
    client, auth_headers, machine_ids, issue_payload, session_factory, monkeypatch
):
    import app.transfer_service as service

    original = service.make_protocol_documents
    calls = 0

    def fail_on_second(*args, **kwargs):
        nonlocal calls
        calls += 1
        if calls == 2:
            raise RuntimeError("controlled document failure")
        return original(*args, **kwargs)

    monkeypatch.setattr(service, "make_protocol_documents", fail_on_second)
    response = issue(
        client, auth_headers, issue_payload(machine_ids["4"], machine_ids["5"])
    )
    assert response.status_code == 500
    with session_factory() as session:
        assert session.scalar(select(func.count(TransferProtocol.id))) == 0
        assert session.scalar(select(func.count(TransferBatch.id))) == 0
        assert session.scalar(select(func.count(ProtocolDocument.id))) == 0
        assert session.get(Machine, machine_ids["4"]).status == "READY"
        assert session.get(Machine, machine_ids["5"]).status == "READY"


def test_sqlite_partial_unique_index_is_present(session_factory):
    with session_factory() as session:
        indexes = {index["name"]: index for index in inspect(session.bind).get_indexes("transfer_protocols")}
        assert indexes["uq_transfer_protocols_active_machine"]["unique"] == 1
        assert indexes["uq_transfer_protocols_active_machine"]["dialect_options"]["sqlite_where"] is not None


def test_active_transfer_index_compiles_for_postgresql():
    index = next(
        index
        for index in TransferProtocol.__table__.indexes
        if index.name == "uq_transfer_protocols_active_machine"
    )
    ddl = str(CreateIndex(index).compile(dialect=postgresql.dialect()))
    assert "UNIQUE INDEX" in ddl
    assert "WHERE is_active IS TRUE" in ddl


def test_new_transfer_operations_require_authentication_and_admin_role(
    client, viewer_headers, machine_ids, issue_payload
):
    assert issue(client, {}, issue_payload(machine_ids["4"])).status_code == 401
    forbidden = issue(client, viewer_headers, issue_payload(machine_ids["4"]))
    assert forbidden.status_code == 403
    assert forbidden.json()["detail"]["code"] == "transfer_permission_denied"


def test_batch_details_show_individual_partial_progress(
    client, auth_headers, machine_ids, issue_payload
):
    created = issue(
        client, auth_headers, issue_payload(machine_ids["4"], machine_ids["5"])
    ).json()
    client.post(
        "/api/transfers/bulk-return",
        headers=auth_headers,
        json=return_payload(created["transfers"][0]),
    )
    details = client.get(
        f"/api/transfer-batches/{created['batch_id']}", headers=auth_headers
    )
    assert details.status_code == 200
    body = details.json()
    assert body["returned_machines"] == 1
    assert body["still_issued_machines"] == 1
    assert sorted(transfer["is_active"] for transfer in body["transfers"]) == [False, True]


def test_duplicate_selection_returns_structured_bulgarian_validation_error(
    client, auth_headers, machine_ids, issue_payload
):
    machine_id = machine_ids["4"]
    response = issue(client, auth_headers, issue_payload(machine_id, machine_id))
    assert response.status_code == 422
    detail = response.json()["detail"]
    assert detail["code"] == "validation_error"
    assert "повече от веднъж" in detail["message"]
