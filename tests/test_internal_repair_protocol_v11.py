from __future__ import annotations

import io
import logging
import re
from pathlib import Path

from app.models import (
    DocumentTemplate,
    DocumentTemplateVersion,
    GeneratedDocument,
    Repair,
)
from app.repairs import service as repair_service
from docx import Document
from sqlalchemy import select


def _advance_to_completed(client, headers, repair_id: int):
    payloads = [
        {"status": "DIAGNOSIS", "inspection_complete": True, "diagnosis": "Диагностика на помпа", "required_work": "Подмяна на уплътнения", "diagnosis_minutes": 30},
        {"status": "REPAIRING", "work_performed": "Подменени уплътнения и извършено сглобяване", "repair_minutes": 75, "result": "Машината е възстановена"},
        {"test_passed": True, "test_method": "Функционален тест", "test_details": "Тестът е успешен", "functional_test_result": "Работи нормално", "condition_after": "Добро", "testing_minutes": 20},
        {"status": "COMPLETED", "test_passed": True, "test_method": "Функционален тест", "testing_minutes": 20, "work_performed": "Подменени уплътнения и извършено сглобяване", "result": "Машината е възстановена", "test_details": "Тестът е успешен", "condition_after": "Добро"},
    ]
    response = None
    for payload in payloads:
        response = client.patch(f"/api/repair-cases/{repair_id}", headers=headers, json=payload)
        assert response.status_code == 200, response.text
    return response


def test_internal_repair_protocol_is_bg_contains_participants_and_is_idempotent(client, auth_headers, machine_ids, session_factory):
    created = client.post(
        "/api/repair-cases", headers=auth_headers,
        json={"machine_id": machine_ids["4"], "reported_problem": "Теч от помпата", "condition_before": "Неизправно"},
    )
    assert created.status_code == 201, created.text
    repair_id = created.json()["id"]
    participant = client.post(
        f"/api/repair-cases/{repair_id}/participants", headers=auth_headers,
        json={"full_name": "Иван Петров Иванов", "job_title": "Механик", "contribution": "Съдействие при сглобяването", "minutes_worked": 55},
    )
    assert participant.status_code == 201, participant.text

    completed = _advance_to_completed(client, auth_headers, repair_id)
    assert completed.json()["status"] == "COMPLETED"
    assert completed.json().get("document_generation_warning") is None

    # UI language requests are intentionally ignored for internal repair protocols.
    generated = client.post(f"/api/repair-cases/{repair_id}/documents?language=en", headers=auth_headers)
    assert generated.status_code == 201, generated.text
    assert generated.json()["language"] == "bg"
    assert generated.json()["requested_language"] == "en"

    with session_factory() as db:
        docs = list(db.scalars(select(GeneratedDocument).where(GeneratedDocument.repair_id == repair_id)))
        assert len(docs) == 2
        assert {item.language for item in docs} == {"bg"}
        docx = next(item for item in docs if item.format == "docx")
        snapshot = docx.snapshot
        assert snapshot["participants"][0]["full_name"] == "Иван Петров Иванов"
        assert snapshot["responsible_user"]["job_title"]
        document = Document(io.BytesIO(docx.content))
        text = "\n".join(p.text for p in document.paragraphs)
        for table in document.tables:
            text += "\n" + "\n".join(cell.text for row in table.rows for cell in row.cells)
        assert "Иван Петров Иванов" in text
        assert "ПРОТОКОЛ ЗА ПРИЕМАНЕ НА ОБОРУДВАНЕ ЗА РЕМОНТ" in text
        assert "ПРОТОКОЛ ЗА ИЗВЪРШЕН РЕМОНТ" in text

    repeated = client.patch(f"/api/repair-cases/{repair_id}", headers=auth_headers, json={"status": "COMPLETED"})
    assert repeated.status_code == 200, repeated.text
    with session_factory() as db:
        assert len(list(db.scalars(select(GeneratedDocument).where(GeneratedDocument.repair_id == repair_id)))) == 2


def test_missing_bg_template_rolls_back_completed_repair(client, auth_headers, machine_ids, session_factory):
    created = client.post(
        "/api/repair-cases", headers=auth_headers,
        json={"machine_id": machine_ids["5"], "reported_problem": "Повреден клапан", "condition_before": "Приета за диагностика"},
    )
    assert created.status_code == 201
    repair_id = created.json()["id"]
    for payload in [
        {"status": "DIAGNOSIS", "inspection_complete": True, "diagnosis": "Клапанът е блокирал", "required_work": "Подмяна на клапана", "diagnosis_minutes": 20},
        {"status": "REPAIRING", "work_performed": "Клапанът е сменен", "repair_minutes": 45, "result": "Ремонтът е извършен"},
        {"test_passed": True, "test_method": "Функционален тест", "test_details": "Успешен тест", "testing_minutes": 15, "condition_after": "Изправна"},
    ]:
        response = client.patch(f"/api/repair-cases/{repair_id}", headers=auth_headers, json=payload)
        assert response.status_code == 200, response.text
    with session_factory() as db:
        template = db.scalar(select(DocumentTemplate).where(DocumentTemplate.document_type == "REPAIR_PROTOCOL"))
        versions = list(db.scalars(select(DocumentTemplateVersion).where(DocumentTemplateVersion.template_id == template.id, DocumentTemplateVersion.language == "bg")))
        for version in versions:
            version.is_published = False
        db.commit()
    completed = client.patch(
        f"/api/repair-cases/{repair_id}", headers=auth_headers,
        json={"status": "COMPLETED", "test_passed": True, "test_method": "Функционален тест", "testing_minutes": 15, "work_performed": "Клапанът е сменен", "result": "Ремонтът е извършен", "condition_after": "Изправна", "test_details": "Успешен тест"},
    )
    assert completed.status_code == 409, completed.text
    body = completed.json()
    assert body["detail"]["code"] == "repair_protocol_template_unavailable"
    with session_factory() as db:
        repair = db.get(Repair, repair_id)
        assert repair.status == "REPAIRING"
        assert repair.machine.status == "REPAIR"
        assert not list(db.scalars(select(GeneratedDocument).where(GeneratedDocument.repair_id == repair_id)))


def test_unexpected_repair_document_error_is_traceable_and_rolls_back(
    client,
    auth_headers,
    machine_ids,
    session_factory,
    monkeypatch,
    caplog,
):
    created = client.post(
        "/api/repair-cases",
        headers=auth_headers,
        json={
            "machine_id": machine_ids["5"],
            "reported_problem": "Проверка на атомарно приключване",
            "condition_before": "Приета за диагностика",
        },
    )
    assert created.status_code == 201, created.text
    repair_id = created.json()["id"]
    for payload in [
        {
            "status": "DIAGNOSIS",
            "inspection_complete": True,
            "diagnosis": "Извършена диагностика",
            "required_work": "Контролирана ремонтна операция",
            "diagnosis_minutes": 20,
        },
        {
            "status": "REPAIRING",
            "work_performed": "Извършена контролирана ремонтна операция",
            "repair_minutes": 40,
            "result": "Възстановена работа",
        },
        {
            "test_passed": True,
            "test_method": "Функционален тест",
            "testing_minutes": 15,
            "test_details": "Успешен тест",
            "condition_after": "Изправна",
        },
    ]:
        response = client.patch(
            f"/api/repair-cases/{repair_id}", headers=auth_headers, json=payload
        )
        assert response.status_code == 200, response.text

    def fail_document_generation(*_args, **_kwargs):
        raise RuntimeError("controlled document generator failure")

    monkeypatch.setattr(
        repair_service, "make_repair_documents", fail_document_generation
    )
    caplog.set_level(logging.ERROR, logger="uvicorn.error")
    completed = client.patch(
        f"/api/repair-cases/{repair_id}",
        headers=auth_headers,
        json={
            "status": "COMPLETED",
            "test_passed": True,
            "test_method": "Функционален тест",
            "testing_minutes": 15,
            "work_performed": "Извършена контролирана ремонтна операция",
            "result": "Възстановена работа",
            "condition_after": "Изправна",
            "test_details": "Успешен тест",
        },
    )

    assert completed.status_code == 500, completed.text
    detail = completed.json()["detail"]
    assert detail["code"] == "repair_protocol_generation_failed"
    assert detail["operation"] == "repair_completion"
    assert detail["stage"] == "document_generation"
    assert re.fullmatch(r"REPERR-[0-9A-F]{12}", detail["diagnostic_id"])
    assert detail["diagnostic_id"] in detail["message"]
    with session_factory() as db:
        repair = db.get(Repair, repair_id)
        assert repair.status == "REPAIRING"
        assert repair.machine.status == "REPAIR"
        assert not list(
            db.scalars(
                select(GeneratedDocument).where(
                    GeneratedDocument.repair_id == repair_id
                )
            )
        )
    record = next(
        item for item in caplog.records if item.message == "assetcore_workflow_error"
    )
    assert {
        key: record.assetcore[key]
        for key in (
            "code",
            "operation",
            "stage",
            "diagnostic_id",
            "repair_id",
            "machine_id",
        )
    } == {
        "code": "repair_protocol_generation_failed",
        "operation": "repair_completion",
        "stage": "document_generation",
        "diagnostic_id": detail["diagnostic_id"],
        "repair_id": repair_id,
        "machine_id": machine_ids["5"],
    }
    assert isinstance(record.assetcore["user_id"], int)


def test_compatibility_repair_close_generates_required_protocol_atomically(
    client, auth_headers, machine_ids, session_factory
):
    created = client.post(
        "/api/repair-cases",
        headers=auth_headers,
        json={"machine_id": machine_ids["7"], "reported_problem": "Проверка на съвместимия маршрут", "condition_before": "Приета за диагностика"},
    )
    assert created.status_code == 201, created.text
    repair_id = created.json()["id"]
    for payload in [
        {"status": "DIAGNOSIS", "inspection_complete": True, "diagnosis": "Извършена диагностика", "required_work": "Извършване на ремонтна операция", "diagnosis_minutes": 25},
        {"status": "REPAIRING", "work_performed": "Извършена ремонтна операция", "repair_minutes": 50, "result": "Възстановена работа"},
        {"test_passed": True, "test_method": "Функционален тест", "testing_minutes": 15, "test_details": "Успешен функционален тест", "condition_after": "Изправна"},
    ]:
        response = client.patch(
            f"/api/repair-cases/{repair_id}", headers=auth_headers, json=payload
        )
        assert response.status_code == 200, response.text

    completed = client.patch(
        f"/api/repairs/{repair_id}", headers=auth_headers, json={"close": True}
    )
    assert completed.status_code == 200, completed.text
    assert completed.json()["status"] == "COMPLETED"
    assert completed.json()["machine"]["status"] == "READY"
    with session_factory() as db:
        documents = list(
            db.scalars(
                select(GeneratedDocument).where(
                    GeneratedDocument.repair_id == repair_id
                )
            )
        )
        assert {document.format for document in documents} == {"docx", "pdf"}


def test_repair_participant_cannot_change_after_completion(client, auth_headers, machine_ids):
    created = client.post("/api/repair-cases", headers=auth_headers, json={"machine_id": machine_ids["7"], "reported_problem": "Тест", "condition_before": "Приета за диагностика"})
    repair_id = created.json()["id"]
    added = client.post(f"/api/repair-cases/{repair_id}/participants", headers=auth_headers, json={"full_name": "Петър Иванов Петров", "minutes_worked": 20})
    assert added.status_code == 201
    _advance_to_completed(client, auth_headers, repair_id)
    blocked = client.delete(f"/api/repair-cases/{repair_id}/participants/{added.json()['id']}", headers=auth_headers)
    assert blocked.status_code == 409
    assert blocked.json()["detail"]["code"] == "completed_repair_is_locked"


def test_repair_frontend_contract_is_internal_and_bulgarian():
    frontend = Path(__file__).resolve().parents[1] / "frontend" / "src"
    source = "\n".join(
        path.read_text(encoding="utf-8")
        for path in (
            frontend / "IndustrialPlatform.tsx",
            frontend / "features" / "repairs" / "IndustrialRepairs.tsx",
        )
    )
    translations = (Path(__file__).resolve().parents[1] / "frontend" / "src" / "i18n.tsx").read_text(encoding="utf-8")
    assert "/documents?language=${locale}" not in source
    assert "repairCase.generateProtocolBg" in source
    assert "repairCase.section.participants" in source
    assert "Нов вътрешен ремонт" in translations
    assert "Приета → Диагностика → В ремонт → Завършена" in translations
